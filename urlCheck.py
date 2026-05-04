import argparse, ctypes, datetime, html, io, json, os, pathlib, platform, re, shutil, signal, struct, subprocess, sys, tempfile, time, traceback, urllib.error, urllib.parse, urllib.request

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from playwright.sync_api import sync_playwright

# pythonnet (the `clr` module that bridges into the .NET Framework) is imported
# lazily inside showGuiDialog and showFinalGuiMessage so the CLI path has no
# GUI cost or hard runtime dependency on it. The .NET Framework 4.8 ships
# with every supported Windows 10 (since 1903) and Windows 11; no extra
# runtime install is required.


# --- Constants ---

bDefaultFullPage = True
bDefaultHeadless = False
bDefaultIgnoreHttpsErrors = True

iCdnTimeoutSec = 30
iCsvMaxHtmlLen = 32000
iDefaultNavTimeoutMs = 60000
iDefaultPostLoadDelayMs = 1500
iDefaultViewportHeight = 1440
iDefaultViewportWidth = 1600
iLayoutButtonHeight = 26
iLayoutButtonWidth = 130
iLayoutFormWidth = 600
iLayoutGap = 7
iLayoutLabelWidth = 130
iLayoutLeft = 12
iLayoutRight = 12
iLayoutRowGap = 11
iLayoutTextHeight = 23
iLayoutTop = 12
iMaxTitleLen = 80
iNetworkIdleTimeoutMs = 8000
# Extra delay applied AFTER the user confirms the --authenticate
# prompt, on top of networkidle and the default post-load delay.
# Post-auth UIs on sites like Facebook, WhatsApp Web, and Slack
# render their main content asynchronously after networkidle fires,
# so without this extra wait the accessibility scan can run against
# a half-rendered DOM and produce a misleadingly empty report.
iAuthPostConfirmSettleDelayMs = 4000

sAccessibilityInsightsUrl = "https://accessibilityinsights.io/docs/web/overview/"
sAccessibilityYamlName = "page.yaml"
sAcrDocxName = "ACR.docx"
sAcrWorkbookName = "ACR.xlsx"
sBrowserChannel = "msedge"
sConfigDirName = "urlCheck"
sConfigFileName = "urlCheck.ini"
sDequeRuleUrlBase = "https://dequeuniversity.com/rules/axe"
sFallbackAxeVersion = "4.10"
sFallbackTitle = "untitled-page"
sJsonName = "results.json"
sLogFileName = "urlCheck.log"
sMsAccessibilityUrl = "https://learn.microsoft.com/accessibility/"
sProgramName = "urlCheck"
sProgramVersion = "1.11.0"
sReportName = "report.htm"
sReportWorkbookName = "report.xlsx"
sScreenshotName = "page.png"
sSourceName = "page.htm"
sUsage = "Usage: urlCheck [options] <url, domain, local html file, or url-list text file>"
sUserAgent = "urlCheck/1.11.0 (+Playwright Python + axe-core)"
sWcagBaseUrl = "https://www.w3.org/WAI/WCAG22/Understanding/"
sWcagQuickRefBase = "https://www.w3.org/WAI/WCAG22/quickref/"


# --- Data tables ---

aAllowedLocalExtensions = [".htm", ".html", ".xhtml"]

aLocalFileExts = {
    ".asp", ".aspx", ".css", ".csv", ".docx", ".gif", ".htm", ".html",
    ".ico", ".jpeg", ".jpg", ".js", ".json", ".jsx", ".md", ".pdf",
    ".php", ".png", ".pptx", ".py", ".svg", ".ts", ".txt", ".xhtml",
    ".xlsm", ".xlsx", ".xml", ".zip",
}

reDomain = re.compile(
    r"^([a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}(:[0-9]+)?(/[^\s]*)?$"
)
reIpAddress = re.compile(r"^(\d{1,3}\.){3}\d{1,3}(:\d+)?(/[^\s]*)?$")

aAxeCdnUrls = [
    "https://cdn.jsdelivr.net/npm/axe-core@4.11.0/axe.min.js",
    "https://unpkg.com/axe-core@4.11.0/axe.min.js",
]

aAxeRunOptions = {"resultTypes": ["violations", "incomplete", "passes", "inapplicable"]}

aGlossaryRows = [
    ["axe-core", "Deque Systems accessibility testing engine that runs rules against the browser DOM."],
    ["impact", "Axe severity estimate for a rule or node result, such as critical or serious."],
    ["inapplicable", "A rule that did not apply to this page or document."],
    ["incomplete", "A result that needs manual review because automated certainty was not possible."],
    ["node", "A single DOM target reported by axe-core under a rule result."],
    ["outcome", "The top-level axe result bucket: violations, incomplete, passes, or inapplicable."],
    ["pass", "A rule that applied and passed automated checks."],
    ["rule", "An axe-core test identified by a rule ID, help text, description, tags, and node results."],
    ["target", "A CSS selector or selector path returned by axe-core for a specific node result."],
    ["violation", "A confirmed automated accessibility issue found by axe-core."],
    ["WCAG reference", "A success criterion reference derived from axe-core tags such as wcag143, rendered as 1.4.3 where possible."],
]

aProcedures = [
    "Open the supplied url or local HTML file in the installed Microsoft Edge browser through the synchronous Playwright API.",
    "Wait for the page to finish loading, then pause briefly so late DOM updates are more likely to settle.",
    "Load axe-core from reliable public CDNs without requiring Node.js on the user system.",
    "Run axe-core with its default behavior and the minimal resultTypes option set used to return violations, incomplete, passes, and inapplicable results.",
    "Capture the page title, final url, user agent, browser version, screenshot, page HTML, and structured accessibility findings.",
    "Write a static HTML report and a structured Excel workbook.",
]

aRelevantTagPrefixes = ["EN-301-549", "best-practice", "cat.", "section508", "TT", "wcag", "wcag2", "wcag21", "wcag22"]

aOutputSections = ["violations"]
aReportSections = ["violations", "incomplete", "passes", "inapplicable"]

dImpactEmoji = {"critical": "🔴", "serious": "🟠", "moderate": "🟡", "minor": "⚪"}

dImpactRank = {"critical": 1, "serious": 2, "moderate": 3, "minor": 4, "": 5, None: 5}

dOutcomeRank = {"violations": 1, "incomplete": 2, "passes": 3, "inapplicable": 4}

dSectionHeadings = {
    "violations": "Violations",
    "incomplete": "Needs Review",
    "passes": "Passes",
    "inapplicable": "Inapplicable",
}

# WCAG 2.0 / 2.1 / 2.2 success criteria: SC number -> (short name, level, principle)
dWcagSc = {
    "1.1.1": ("Non-text Content",                         "A",   "Perceivable"),
    "1.2.1": ("Audio-only and Video-only (Prerecorded)",  "A",   "Perceivable"),
    "1.2.2": ("Captions (Prerecorded)",                   "A",   "Perceivable"),
    "1.2.3": ("Audio Description or Media Alternative",   "A",   "Perceivable"),
    "1.2.4": ("Captions (Live)",                          "AA",  "Perceivable"),
    "1.2.5": ("Audio Description (Prerecorded)",          "AA",  "Perceivable"),
    "1.2.6": ("Sign Language (Prerecorded)",              "AAA", "Perceivable"),
    "1.2.7": ("Extended Audio Description (Prerecorded)", "AAA", "Perceivable"),
    "1.2.8": ("Media Alternative (Prerecorded)",          "AAA", "Perceivable"),
    "1.2.9": ("Audio-only (Live)",                        "AAA", "Perceivable"),
    "1.3.1": ("Info and Relationships",                   "A",   "Perceivable"),
    "1.3.2": ("Meaningful Sequence",                      "A",   "Perceivable"),
    "1.3.3": ("Sensory Characteristics",                  "A",   "Perceivable"),
    "1.3.4": ("Orientation",                              "AA",  "Perceivable"),
    "1.3.5": ("Identify Input Purpose",                   "AA",  "Perceivable"),
    "1.3.6": ("Identify Purpose",                         "AAA", "Perceivable"),
    "1.4.1": ("Use of Color",                             "A",   "Perceivable"),
    "1.4.2": ("Audio Control",                            "A",   "Perceivable"),
    "1.4.3": ("Contrast (Minimum)",                       "AA",  "Perceivable"),
    "1.4.4": ("Resize Text",                              "AA",  "Perceivable"),
    "1.4.5": ("Images of Text",                           "AA",  "Perceivable"),
    "1.4.6": ("Contrast (Enhanced)",                      "AAA", "Perceivable"),
    "1.4.7": ("Low or No Background Audio",               "AAA", "Perceivable"),
    "1.4.8": ("Visual Presentation",                      "AAA", "Perceivable"),
    "1.4.9": ("Images of Text (No Exception)",            "AAA", "Perceivable"),
    "1.4.10": ("Reflow",                                  "AA",  "Perceivable"),
    "1.4.11": ("Non-text Contrast",                       "AA",  "Perceivable"),
    "1.4.12": ("Text Spacing",                            "AA",  "Perceivable"),
    "1.4.13": ("Content on Hover or Focus",               "AA",  "Perceivable"),
    "2.1.1": ("Keyboard",                                 "A",   "Operable"),
    "2.1.2": ("No Keyboard Trap",                         "A",   "Operable"),
    "2.1.3": ("Keyboard (No Exception)",                  "AAA", "Operable"),
    "2.1.4": ("Character Key Shortcuts",                  "A",   "Operable"),
    "2.2.1": ("Timing Adjustable",                        "A",   "Operable"),
    "2.2.2": ("Pause, Stop, Hide",                        "A",   "Operable"),
    "2.2.3": ("No Timing",                                "AAA", "Operable"),
    "2.2.4": ("Interruptions",                            "AAA", "Operable"),
    "2.2.5": ("Re-authenticating",                        "AAA", "Operable"),
    "2.2.6": ("Timeouts",                                 "AAA", "Operable"),
    "2.3.1": ("Three Flashes or Below Threshold",         "A",   "Operable"),
    "2.3.2": ("Three Flashes",                            "AAA", "Operable"),
    "2.3.3": ("Animation from Interactions",              "AAA", "Operable"),
    "2.4.1": ("Bypass Blocks",                            "A",   "Operable"),
    "2.4.2": ("Page Titled",                              "A",   "Operable"),
    "2.4.3": ("Focus Order",                              "A",   "Operable"),
    "2.4.4": ("Link Purpose (In Context)",                "A",   "Operable"),
    "2.4.5": ("Multiple Ways",                            "AA",  "Operable"),
    "2.4.6": ("Headings and Labels",                      "AA",  "Operable"),
    "2.4.7": ("Focus Visible",                            "AA",  "Operable"),
    "2.4.8": ("Location",                                 "AAA", "Operable"),
    "2.4.9": ("Link Purpose (Link Only)",                 "AAA", "Operable"),
    "2.4.10": ("Section Headings",                        "AAA", "Operable"),
    "2.4.11": ("Focus Not Obscured (Minimum)",            "AA",  "Operable"),
    "2.4.12": ("Focus Not Obscured (Enhanced)",           "AAA", "Operable"),
    "2.4.13": ("Focus Appearance",                        "AA",  "Operable"),
    "2.5.1": ("Pointer Gestures",                         "A",   "Operable"),
    "2.5.2": ("Pointer Cancellation",                     "A",   "Operable"),
    "2.5.3": ("Label in Name",                            "A",   "Operable"),
    "2.5.4": ("Motion Actuation",                         "A",   "Operable"),
    "2.5.5": ("Target Size (Enhanced)",                   "AAA", "Operable"),
    "2.5.6": ("Concurrent Input Mechanisms",              "AAA", "Operable"),
    "2.5.7": ("Dragging Movements",                       "AA",  "Operable"),
    "2.5.8": ("Target Size (Minimum)",                    "AA",  "Operable"),
    "3.1.1": ("Language of Page",                         "A",   "Understandable"),
    "3.1.2": ("Language of Parts",                        "AA",  "Understandable"),
    "3.1.3": ("Unusual Words",                            "AAA", "Understandable"),
    "3.1.4": ("Abbreviations",                            "AAA", "Understandable"),
    "3.1.5": ("Reading Level",                            "AAA", "Understandable"),
    "3.1.6": ("Pronunciation",                            "AAA", "Understandable"),
    "3.2.1": ("On Focus",                                 "A",   "Understandable"),
    "3.2.2": ("On Input",                                 "A",   "Understandable"),
    "3.2.3": ("Consistent Navigation",                    "AA",  "Understandable"),
    "3.2.4": ("Consistent Identification",                "AA",  "Understandable"),
    "3.2.5": ("Change on Request",                        "AAA", "Understandable"),
    "3.2.6": ("Consistent Help",                          "A",   "Understandable"),
    "3.3.1": ("Error Identification",                     "A",   "Understandable"),
    "3.3.2": ("Labels or Instructions",                   "A",   "Understandable"),
    "3.3.3": ("Error Suggestion",                         "AA",  "Understandable"),
    "3.3.4": ("Error Prevention (Legal, Financial, Data)","AA",  "Understandable"),
    "3.3.5": ("Help",                                     "AAA", "Understandable"),
    "3.3.6": ("Error Prevention (All)",                   "AAA", "Understandable"),
    "3.3.7": ("Redundant Entry",                          "A",   "Understandable"),
    "3.3.8": ("Accessible Authentication (Minimum)",      "AA",  "Understandable"),
    "3.3.9": ("Accessible Authentication (Enhanced)",     "AAA", "Understandable"),
    "4.1.1": ("Parsing",                                  "A",   "Robust"),
    "4.1.2": ("Name, Role, Value",                        "A",   "Robust"),
    "4.1.3": ("Status Messages",                          "AA",  "Robust"),
}

# Supplementary WCAG references for best-practice rules that have no wcagXXX axe tags.
# These are the closest related WCAG success criteria, included to give users context
# about why the issue matters. They are marked as advisory rather than strict failures.
dBpWcagRefs = {
    "accesskeys":                          ["2.1.4"],
    "aria-allowed-role":                   ["4.1.2"],
    "aria-dialog-name":                    ["4.1.2"],
    "avoid-inline-spacing":                ["1.4.12"],
    "css-orientation-lock":                ["1.3.4"],
    "empty-heading":                       ["2.4.6"],
    "empty-table-header":                  ["1.3.1"],
    "focus-trap":                          ["2.1.2"],
    "frame-tested":                        ["4.1.2"],
    "heading-order":                       ["2.4.6"],
    "hidden-content":                      ["4.1.2"],
    "identical-links-same-purpose":        ["2.4.9"],
    "label-title-only":                    ["3.3.2"],
    "landmark-banner-is-top-level":        ["1.3.1"],
    "landmark-complementary-is-top-level": ["1.3.1"],
    "landmark-contentinfo-is-top-level":   ["1.3.1"],
    "landmark-main-is-top-level":          ["1.3.1"],
    "landmark-no-duplicate-banner":        ["1.3.1"],
    "landmark-no-duplicate-contentinfo":   ["1.3.1"],
    "landmark-no-duplicate-main":          ["1.3.1"],
    "landmark-one-main":                   ["2.4.1"],
    "landmark-unique":                     ["1.3.1"],
    "meta-viewport":                       ["1.4.4"],
    "p-as-heading":                        ["1.3.1"],
    "page-has-heading-one":                ["2.4.6"],
    "presentation-role-conflict":          ["4.1.2"],
    "region":                              ["1.3.1", "2.4.1"],
    "scope-attr-valid":                    ["1.3.1"],
    "scrollable-region-focusable":         ["2.1.1"],
    "select-name":                         ["4.1.2"],
    "server-side-image-map":               ["2.1.1"],
    "skip-link":                           ["2.4.1"],
    "summary-name":                        ["4.1.2"],
    "tabindex":                            ["2.4.3"],
    "table-duplicate-name":                ["1.3.1"],
    "table-fake-caption":                  ["1.3.1"],
    "target-size":                         ["2.5.8"],
}



# WCAG 2.2 success criteria with manual-test instructions for the
# Accessibility Conformance Report (ACR.xlsx) feature. 86 criteria
# total: 31 at Level A, 24 at AA, 31 at AAA. The obsolete 4.1.1
# Parsing is omitted (removed from WCAG 2.2). Manual-test text is
# in second-person imperative voice, 4-5 numbered steps per
# criterion. fragment is the URL fragment on the W3C Quick
# Reference at https://www.w3.org/WAI/WCAG22/quickref/.
dWcag22 = {
    "1.1.1": {
        "name": "Non-text Content",
        "level": "A",
        "fragment": "non-text-content",
        "checks": [
            "Identify each non-text element on the page (images, icons, charts, audio clips, video clips, CAPTCHA, form controls).",
            "For each element, decide whether it conveys information, is purely decorative, or has a specific role (CAPTCHA, test, sensory experience).",
            "Confirm that informational images have a text alternative that conveys the same purpose. Use a screen reader, inspect the alt attribute, or examine accessible-name properties.",
            "Confirm that decorative images are hidden from assistive technologies (empty alt, role=presentation, aria-hidden, or CSS background).",
            "For controls and inputs, confirm that an accessible name describes the purpose. For media, confirm that a text alternative identifies the content.",
        ],
    },
    "1.2.1": {
        "name": "Audio-only and Video-only (Prerecorded)",
        "level": "A",
        "fragment": "audio-only-and-video-only-prerecorded",
        "checks": [
            "Identify each prerecorded audio-only and video-only resource on the page.",
            "For prerecorded audio-only, confirm that a text transcript is provided that conveys the same information.",
            "For prerecorded video-only (no audio), confirm that either a text alternative or an audio track describing the visual content is provided.",
            "Verify that the alternative is clearly labeled, easily found near the media, and accessible to screen readers.",
        ],
    },
    "1.2.2": {
        "name": "Captions (Prerecorded)",
        "level": "A",
        "fragment": "captions-prerecorded",
        "checks": [
            "Identify each prerecorded video with synchronized audio (movies, recorded webinars, instructional videos).",
            "Confirm captions are available and synchronized with the audio track.",
            "Verify captions include all spoken dialogue plus important non-speech audio (music, sound effects, speaker identification).",
            "Confirm the user can turn captions on or off through standard player controls.",
        ],
    },
    "1.2.3": {
        "name": "Audio Description or Media Alternative (Prerecorded)",
        "level": "A",
        "fragment": "audio-description-or-media-alternative-prerecorded",
        "checks": [
            "Identify each prerecorded video that conveys information visually beyond what the audio track describes.",
            "Confirm that EITHER an audio description is provided (additional narration during natural pauses) OR a text alternative is provided that conveys all visual and auditory content.",
            "If an audio description is provided, listen with the audio track muted and verify that the description conveys the visual information.",
            "If a text alternative is provided, verify it is clearly labeled and located near the media.",
        ],
    },
    "1.2.4": {
        "name": "Captions (Live)",
        "level": "AA",
        "fragment": "captions-live",
        "checks": [
            "Identify each live or streaming video that includes audio (live broadcasts, webinars, conferences).",
            "Confirm that real-time captions are available and synchronized with the live audio.",
            "Verify captions cover all spoken dialogue and important non-speech audio.",
            "Test that captions are reliably delivered even when network conditions vary.",
        ],
    },
    "1.2.5": {
        "name": "Audio Description (Prerecorded)",
        "level": "AA",
        "fragment": "audio-description-prerecorded",
        "checks": [
            "Identify each prerecorded video with significant visual information not conveyed in its audio track.",
            "Confirm an audio description track is provided that narrates important visual details during natural audio pauses.",
            "Listen with the screen turned away and verify the audio description, combined with the original audio, conveys the full content.",
            "Confirm the audio description can be enabled and disabled.",
        ],
    },
    "1.2.6": {
        "name": "Sign Language (Prerecorded)",
        "level": "AAA",
        "fragment": "sign-language-prerecorded",
        "checks": [
            "Identify each prerecorded video with synchronized audio.",
            "Confirm that a sign language interpretation is provided for the audio content, either embedded in the video or available as a separate synchronized track.",
            "Verify that the sign language interpreter is clearly visible (well-lit, in the foreground, large enough to read facial expressions and hand shapes).",
            "Confirm the sign language matches the language community of the intended audience (e.g., ASL for U.S. content, BSL for U.K. content).",
        ],
    },
    "1.2.7": {
        "name": "Extended Audio Description (Prerecorded)",
        "level": "AAA",
        "fragment": "extended-audio-description-prerecorded",
        "checks": [
            "Identify prerecorded videos where audio pauses are too short to convey adequate audio descriptions.",
            "Confirm that an extended audio description is available, where the video pauses to allow longer description.",
            "Verify the extended description conveys all important visual information that the standard audio track does not cover.",
            "Confirm the extended description can be turned on or off, since pausing the video may disrupt users who do not need it.",
        ],
    },
    "1.2.8": {
        "name": "Media Alternative (Prerecorded)",
        "level": "AAA",
        "fragment": "media-alternative-prerecorded",
        "checks": [
            "Identify each prerecorded synchronized media (video with audio) and each prerecorded video-only resource.",
            "Confirm a complete text alternative (alternative for time-based media) is provided that conveys all visual and auditory information.",
            "Verify the text alternative is structured (with headings, time markers, or scene descriptions) so a reader can follow the content.",
            "Confirm the text alternative is clearly labeled and easy to find near the media.",
        ],
    },
    "1.2.9": {
        "name": "Audio-only (Live)",
        "level": "AAA",
        "fragment": "audio-only-live",
        "checks": [
            "Identify each live audio-only stream (live radio, live audio podcast, conference call).",
            "Confirm that a real-time text alternative is available, such as live captioning or a real-time transcript.",
            "Verify the text alternative captures all spoken content and important non-speech audio.",
            "Confirm the text alternative is reliably available throughout the broadcast.",
        ],
    },
    "1.3.1": {
        "name": "Info and Relationships",
        "level": "A",
        "fragment": "info-and-relationships",
        "checks": [
            "Inspect the page structure: headings, lists, tables, form labels, and groupings.",
            "Confirm visual structure is conveyed programmatically: headings use h1\u2013h6, lists use ul/ol/dl, tables use proper th/td and scope, form fields have associated label elements.",
            "Use a screen reader to verify the structure is announced (heading levels, list items, table headers when navigating cells).",
            "Confirm landmark roles (header, nav, main, footer) or HTML5 sectioning elements are present where appropriate.",
        ],
    },
    "1.3.2": {
        "name": "Meaningful Sequence",
        "level": "A",
        "fragment": "meaningful-sequence",
        "checks": [
            "Disable CSS or use a screen reader to read the page in DOM order.",
            "Confirm the reading order makes sense: headings precede their content, labels precede their inputs, related items are adjacent.",
            "Verify that visual reordering via CSS (flexbox order, grid placement, absolute positioning) has not made the DOM order incoherent.",
            "Test with the keyboard: tab order should match the meaningful reading order.",
        ],
    },
    "1.3.3": {
        "name": "Sensory Characteristics",
        "level": "A",
        "fragment": "sensory-characteristics",
        "checks": [
            "Search the page for instructions that rely on shape, size, color, position, or sound (\"click the round button\", \"the field on the right\", \"when you hear a beep\").",
            "Confirm each such instruction is supplemented with a non-sensory cue (a label, a heading, a name).",
            "Verify a screen reader user can follow the instruction without seeing the page or hearing audio cues.",
            "Verify a user with low vision or colorblindness can identify the referenced item.",
        ],
    },
    "1.3.4": {
        "name": "Orientation",
        "level": "AA",
        "fragment": "orientation",
        "checks": [
            "Open the page in both portrait and landscape orientations on a mobile device or rotated viewport.",
            "Confirm the content displays correctly in both orientations.",
            "Verify the page does not lock the user into a single orientation unless that orientation is essential (a check-deposit photo, a piano app).",
            "Confirm functionality is preserved across orientation changes.",
        ],
    },
    "1.3.5": {
        "name": "Identify Input Purpose",
        "level": "AA",
        "fragment": "identify-input-purpose",
        "checks": [
            "Identify each form input that collects information about the user (name, email, address, phone, payment, demographic data).",
            "Confirm each such input has the appropriate autocomplete attribute set to the standard token (\"name\", \"email\", \"street-address\", \"tel\", \"cc-number\", and so on).",
            "Test browser autofill: the field should be filled correctly when the user invokes autofill.",
            "Verify input types are appropriate (email, tel, url, number) so assistive technologies can adapt.",
        ],
    },
    "1.3.6": {
        "name": "Identify Purpose",
        "level": "AAA",
        "fragment": "identify-purpose",
        "checks": [
            "Identify user-interface components, icons, and regions on the page (navigation, search, contact, log-in).",
            "Confirm the purpose of each component is programmatically determinable through ARIA roles, microdata, schema.org annotations, or descriptive accessible names.",
            "Verify that personalization tools can recognize and adapt the components for users who need symbol substitution, simplified layouts, or content filtering.",
            "Test with at least one personalization or symbol-augmentation tool that depends on these annotations.",
        ],
    },
    "1.4.1": {
        "name": "Use of Color",
        "level": "A",
        "fragment": "use-of-color",
        "checks": [
            "Identify any place where color alone conveys information, indicates an action, prompts a response, or distinguishes a visual element (required-field highlighting, status indicators, error states, link styling).",
            "Confirm a non-color cue is also present: text, icon, underline, pattern, or shape.",
            "View the page with a colorblindness simulator or in grayscale and verify all information is still conveyed.",
            "Verify links within text are distinguishable from surrounding text by something other than color (underline, weight, icon).",
        ],
    },
    "1.4.2": {
        "name": "Audio Control",
        "level": "A",
        "fragment": "audio-control",
        "checks": [
            "Identify any audio that plays automatically on page load and lasts more than three seconds.",
            "Confirm a mechanism is provided to pause, stop, or independently control the volume of that audio.",
            "Verify the control is operable by keyboard and announced by screen readers.",
            "Test that the control is reachable before the auto-playing audio interferes with screen reader output.",
        ],
    },
    "1.4.3": {
        "name": "Contrast (Minimum)",
        "level": "AA",
        "fragment": "contrast-minimum",
        "checks": [
            "Identify each text element on the page (body text, headings, links, buttons, labels).",
            "For text below 18pt regular or 14pt bold, measure contrast and confirm it is at least 4.5:1 against its background.",
            "For larger text (18pt+ regular or 14pt+ bold), confirm contrast is at least 3:1.",
            "Use a contrast-checking tool or browser extension; check both the resting state and any hover, focus, or active states.",
        ],
    },
    "1.4.4": {
        "name": "Resize Text",
        "level": "AA",
        "fragment": "resize-text",
        "checks": [
            "Use the browser's text-zoom feature (or set zoom to 200%) and reload the page.",
            "Confirm all text remains readable, all functionality remains operable, and no content is lost or clipped.",
            "Verify text does not require horizontal scrolling within content blocks.",
            "Confirm assistive technology (zoom software, screen magnifiers) is not required to read the resized content.",
        ],
    },
    "1.4.5": {
        "name": "Images of Text",
        "level": "AA",
        "fragment": "images-of-text",
        "checks": [
            "Identify each image that contains text (banner graphics, infographics, scanned documents, decorative quotes).",
            "Confirm real text is used instead of images of text wherever the text could be presented with CSS styling.",
            "Verify exceptions are limited to logos, brand marks, and cases where a particular visual presentation is essential.",
            "For permitted images of text, confirm an accurate text alternative is provided.",
        ],
    },
    "1.4.6": {
        "name": "Contrast (Enhanced)",
        "level": "AAA",
        "fragment": "contrast-enhanced",
        "checks": [
            "Identify each text element on the page.",
            "For text below 18pt regular or 14pt bold, confirm contrast is at least 7:1 against its background.",
            "For larger text (18pt+ regular or 14pt+ bold), confirm contrast is at least 4.5:1.",
            "Use a contrast-checking tool; verify both resting state and any focus, hover, or active states meet the enhanced ratio.",
        ],
    },
    "1.4.7": {
        "name": "Low or No Background Audio",
        "level": "AAA",
        "fragment": "low-or-no-background-audio",
        "checks": [
            "Identify each prerecorded audio-only resource that is primarily speech and is not a CAPTCHA, song, or vocalization.",
            "Listen and confirm one of these is true: there is no background sound, the background sound can be turned off by the user, or the background sound is at least 20 decibels (4 times) quieter than the foreground speech.",
            "Verify any user-controllable background audio has accessible controls.",
            "Test playback in a noisy environment to confirm speech remains intelligible.",
        ],
    },
    "1.4.8": {
        "name": "Visual Presentation",
        "level": "AAA",
        "fragment": "visual-presentation",
        "checks": [
            "Verify a mechanism is available for the user to select foreground and background colors.",
            "Confirm that text blocks are no wider than 80 characters (or 40 for CJK languages).",
            "Verify text is not justified (aligned to both left and right edges).",
            "Confirm line spacing is at least 1.5 within paragraphs and paragraph spacing is at least 1.5 times the line spacing.",
            "Verify text resizes to 200% without requiring horizontal scrolling for paragraphs.",
        ],
    },
    "1.4.9": {
        "name": "Images of Text (No Exception)",
        "level": "AAA",
        "fragment": "images-of-text-no-exception",
        "checks": [
            "Identify each image that contains text on the page.",
            "Confirm real text is used in every case except for logos and brand marks.",
            "Verify there are no decorative images of text where styled text could have been used.",
            "Test with text resizing and high-contrast mode to confirm text remains usable.",
        ],
    },
    "1.4.10": {
        "name": "Reflow",
        "level": "AA",
        "fragment": "reflow",
        "checks": [
            "Resize the browser viewport to 320 CSS pixels wide (mobile width) or zoom to 400% in a 1280-pixel window.",
            "Confirm content reflows to fit without requiring horizontal scrolling, except for content that needs two-dimensional layout (data tables, maps, complex images, code).",
            "Verify all functionality remains operable at the reduced viewport.",
            "Confirm no content is hidden or clipped beyond a single scroll direction.",
        ],
    },
    "1.4.11": {
        "name": "Non-text Contrast",
        "level": "AA",
        "fragment": "non-text-contrast",
        "checks": [
            "Identify each user-interface component (buttons, form fields, links, icons indicating state).",
            "Confirm the visual indicators (boundaries, fill, focus rings) have at least 3:1 contrast against adjacent colors.",
            "Identify graphical objects necessary for understanding (chart segments, icons that convey meaning, infographic elements).",
            "Confirm those have at least 3:1 contrast against adjacent colors.",
        ],
    },
    "1.4.12": {
        "name": "Text Spacing",
        "level": "AA",
        "fragment": "text-spacing",
        "checks": [
            "Apply a user stylesheet or browser extension that sets line height to 1.5x font size, paragraph spacing to 2x font size, letter spacing to 0.12x font size, and word spacing to 0.16x font size.",
            "Confirm no content is lost or clipped, no text overlaps, and all functionality remains operable.",
            "Verify text containers grow to accommodate the spacing rather than truncating content.",
            "Test on key page templates (forms, navigation, content pages, interactive widgets).",
        ],
    },
    "1.4.13": {
        "name": "Content on Hover or Focus",
        "level": "AA",
        "fragment": "content-on-hover-or-focus",
        "checks": [
            "Identify content that appears on hover (tooltips, expanded menus) or on focus (popovers, hint text).",
            "Confirm the additional content is dismissible without moving the pointer or focus, typically by pressing Escape.",
            "Verify the user can move the pointer onto the additional content without it disappearing.",
            "Confirm the additional content remains visible until the user dismisses it, moves focus or pointer away, or its information is no longer valid.",
        ],
    },
    "2.1.1": {
        "name": "Keyboard",
        "level": "A",
        "fragment": "keyboard",
        "checks": [
            "Set the page focus to the top of the document.",
            "Press Tab repeatedly to move through every interactive element and confirm each is reachable.",
            "Activate each element using Enter, Space, or arrow keys as appropriate, and confirm the action completes.",
            "Verify drag-and-drop, gesture-based, and time-dependent interactions have keyboard equivalents.",
        ],
    },
    "2.1.2": {
        "name": "No Keyboard Trap",
        "level": "A",
        "fragment": "no-keyboard-trap",
        "checks": [
            "Tab through the entire page; confirm focus can move forward and backward through every interactive element.",
            "Within widgets that capture keyboard input (modal dialogs, custom controls, embedded media), confirm the user can leave the widget using Tab, Shift+Tab, Escape, or another documented key.",
            "Test embedded objects (media players, plugins, iframes) for trapping behavior.",
            "Verify any non-standard exit method is documented or announced to the user.",
        ],
    },
    "2.1.3": {
        "name": "Keyboard (No Exception)",
        "level": "AAA",
        "fragment": "keyboard-no-exception",
        "checks": [
            "Identify all functionality on the page, including features that depend on path-specific user input (free drawing, signature capture, gesture-based interactions).",
            "Confirm every function can be operated by keyboard alone, including those that 2.1.1 permits to be input-method-specific.",
            "Verify keyboard alternatives exist for every gesture, drag operation, and continuous input.",
            "Test with keyboard only and confirm no functionality is unreachable.",
        ],
    },
    "2.1.4": {
        "name": "Character Key Shortcuts",
        "level": "A",
        "fragment": "character-key-shortcuts",
        "checks": [
            "Identify any single-character keyboard shortcuts implemented by the page (typing \"j\" to jump, \"r\" to reply).",
            "Confirm one of these is true for each shortcut: it can be turned off, it can be remapped to use modifier keys (Ctrl, Alt), or it is active only when the relevant control has focus.",
            "Test shortcuts in form contexts to verify they do not interfere with text entry.",
            "Verify shortcuts are documented somewhere accessible (help page, keyboard-shortcut overlay).",
        ],
    },
    "2.2.1": {
        "name": "Timing Adjustable",
        "level": "A",
        "fragment": "timing-adjustable",
        "checks": [
            "Identify any time limit set by the content (session timeouts, response deadlines, auto-advance carousels).",
            "Confirm one of these is true for each: the user can turn it off, the user can adjust it to at least 10 times the default, or the user receives a warning at least 20 seconds before the limit and can extend it.",
            "Verify the option to extend is reachable by keyboard and announced to assistive technology.",
            "Confirm exceptions (real-time events, essential limits, 20-hour-plus limits) are genuinely unavoidable.",
        ],
    },
    "2.2.2": {
        "name": "Pause, Stop, Hide",
        "level": "A",
        "fragment": "pause-stop-hide",
        "checks": [
            "Identify any moving, blinking, or scrolling content lasting more than 5 seconds and presented in parallel with other content (carousels, marquees, animations, news tickers).",
            "Confirm a mechanism is provided to pause, stop, or hide the moving content.",
            "Identify any auto-updating content (live scoreboards, news feeds, stock tickers) and confirm a mechanism is provided to pause, stop, hide, or control the update frequency.",
            "Verify the controls are operable by keyboard and accessible to assistive technology.",
        ],
    },
    "2.2.3": {
        "name": "No Timing",
        "level": "AAA",
        "fragment": "no-timing",
        "checks": [
            "Identify any time limits in the page or content.",
            "Confirm there are no time limits except for real-time events (live auctions, real-time games) where timing is essential.",
            "Verify session timeouts and content advancement are not time-based.",
            "Confirm the user can take any amount of time to complete tasks.",
        ],
    },
    "2.2.4": {
        "name": "Interruptions",
        "level": "AAA",
        "fragment": "interruptions",
        "checks": [
            "Identify any content that interrupts the user (notifications, modal alerts, automatic refreshes, popovers).",
            "Confirm the user can postpone or suppress these interruptions, except for those involving emergencies.",
            "Verify a setting or mechanism is available before the interruption occurs.",
            "Test that suppressing interruptions does not prevent the user from accessing essential information later.",
        ],
    },
    "2.2.5": {
        "name": "Re-authenticating",
        "level": "AAA",
        "fragment": "re-authenticating",
        "checks": [
            "Identify any session that may expire while the user is performing a task.",
            "Simulate a session expiry (or wait for one) while the user has unsaved data.",
            "Confirm that after re-authentication, the user can continue without losing the data they had entered.",
            "Verify the data is preserved through the re-authentication flow without requiring the user to start over.",
        ],
    },
    "2.2.6": {
        "name": "Timeouts",
        "level": "AAA",
        "fragment": "timeouts",
        "checks": [
            "Identify any user inactivity that could cause data loss (form expiration, shopping cart clearance, session-based document storage).",
            "Confirm the user is warned of the duration of inactivity that would cause loss, before they begin the task.",
            "Verify the warning is presented in text and is reachable for screen readers.",
            "Confirm exceptions (data preserved for 20+ hours of inactivity) are reasonable.",
        ],
    },
    "2.3.1": {
        "name": "Three Flashes or Below Threshold",
        "level": "A",
        "fragment": "three-flashes-or-below-threshold",
        "checks": [
            "Identify any content that flashes (rapid changes in luminance or color).",
            "Confirm the flashing occurs no more than three times in any one-second period.",
            "If flashing occurs more than three times per second, confirm the flash area is small (less than ~25% of a 10-degree visual field) and below the general flash and red flash thresholds.",
            "Use the Photosensitive Epilepsy Analysis Tool (PEAT) or equivalent to measure flash content.",
        ],
    },
    "2.3.2": {
        "name": "Three Flashes",
        "level": "AAA",
        "fragment": "three-flashes",
        "checks": [
            "Identify any content that flashes on the page.",
            "Confirm there is no flashing more than three times per second, regardless of size or threshold.",
            "Test all media, animations, and transitions for compliance.",
            "If flashing exists, confirm it can be removed or disabled by the user.",
        ],
    },
    "2.3.3": {
        "name": "Animation from Interactions",
        "level": "AAA",
        "fragment": "animation-from-interactions",
        "checks": [
            "Identify motion animation triggered by user interactions (parallax scrolling, slide transitions, large pan-and-zoom effects).",
            "Confirm the animation can be disabled, except where it is essential to the functionality or information being conveyed.",
            "Test that the prefers-reduced-motion media query is respected if used.",
            "Verify a setting or toggle is exposed to disable non-essential motion.",
        ],
    },
    "2.4.1": {
        "name": "Bypass Blocks",
        "level": "A",
        "fragment": "bypass-blocks",
        "checks": [
            "Identify content that repeats across pages (header, navigation, footer).",
            "Confirm a mechanism is provided to skip past repeated content: a skip link, ARIA landmarks (navigation, main, etc.), proper heading structure, or an explicit \"skip to main content\" feature.",
            "Test the skip mechanism with the keyboard: pressing Tab early on the page should reach a skip link or landmark navigation.",
            "Verify the skip target is appropriate and focus moves there correctly.",
        ],
    },
    "2.4.2": {
        "name": "Page Titled",
        "level": "A",
        "fragment": "page-titled",
        "checks": [
            "View the page title (browser tab, window title bar, or inspect the title element).",
            "Confirm the title describes the page's topic or purpose.",
            "Verify the title differs across pages so each is uniquely identifiable.",
            "For applications, confirm the title updates as the user navigates through views.",
        ],
    },
    "2.4.3": {
        "name": "Focus Order",
        "level": "A",
        "fragment": "focus-order",
        "checks": [
            "Tab through every focusable element on the page from start to finish.",
            "Confirm the focus order matches the meaning and operation of the page (top to bottom, left to right in left-to-right languages, related items together).",
            "Verify modals trap and release focus appropriately when opened and closed.",
            "Confirm dynamic content (newly inserted elements) appears in a logical place in the focus order.",
        ],
    },
    "2.4.4": {
        "name": "Link Purpose (In Context)",
        "level": "A",
        "fragment": "link-purpose-in-context",
        "checks": [
            "Identify each link on the page.",
            "For each link, read the link text plus its surrounding context (sentence, list item, table cell, paragraph).",
            "Confirm the purpose of the link is determinable from the link text or its programmatically determined context.",
            "Avoid generic link text (\"click here\", \"read more\") unless the surrounding context fully clarifies the destination.",
        ],
    },
    "2.4.5": {
        "name": "Multiple Ways",
        "level": "AA",
        "fragment": "multiple-ways",
        "checks": [
            "Identify the navigation mechanisms available on the site.",
            "Confirm at least two of the following are present: site map, search, table of contents, links to all pages, related-page links, and primary navigation.",
            "Verify each mechanism is reachable from every page (or marked as part of a single process where exceptions apply).",
            "Test each mechanism with a keyboard and a screen reader.",
        ],
    },
    "2.4.6": {
        "name": "Headings and Labels",
        "level": "AA",
        "fragment": "headings-and-labels",
        "checks": [
            "Identify each heading on the page.",
            "Confirm each heading describes the content that follows it.",
            "Identify each form label.",
            "Confirm each label describes the purpose of the input it accompanies.",
        ],
    },
    "2.4.7": {
        "name": "Focus Visible",
        "level": "AA",
        "fragment": "focus-visible",
        "checks": [
            "Tab through every focusable element on the page.",
            "Confirm each element shows a clearly visible focus indicator (outline, border change, background change, underline).",
            "Verify the focus indicator is not removed by CSS (no outline:none without a replacement).",
            "Test in both light and dark color schemes if both are supported.",
        ],
    },
    "2.4.8": {
        "name": "Location",
        "level": "AAA",
        "fragment": "location",
        "checks": [
            "Identify the user's current location indicators on the page (breadcrumb trail, current navigation item highlighted, page title in heading).",
            "Confirm the user's location within the site structure is identifiable.",
            "Verify the location information is exposed to assistive technology (current page marked with aria-current, breadcrumbs in a nav with appropriate label).",
            "Test by navigating from the home page and confirming each visited page indicates where the user is.",
        ],
    },
    "2.4.9": {
        "name": "Link Purpose (Link Only)",
        "level": "AAA",
        "fragment": "link-purpose-link-only",
        "checks": [
            "Identify each link on the page.",
            "For each link, confirm the link text alone (without surrounding context) describes the destination.",
            "Avoid generic link text such as \"click here\", \"read more\", or \"learn more\" without further qualification.",
            "If contextual exceptions exist, confirm they meet the criterion's exemption (the purpose is genuinely ambiguous to all users).",
        ],
    },
    "2.4.10": {
        "name": "Section Headings",
        "level": "AAA",
        "fragment": "section-headings",
        "checks": [
            "Identify the major sections of long content (chapters, parts, distinct topics).",
            "Confirm each section begins with a heading element of appropriate level.",
            "Verify the heading structure forms a logical outline (no skipped levels, hierarchical nesting).",
            "Test with a screen reader to navigate by heading and confirm the structure is meaningful.",
        ],
    },
    "2.4.11": {
        "name": "Focus Not Obscured (Minimum)",
        "level": "AA",
        "fragment": "focus-not-obscured-minimum",
        "checks": [
            "Tab through every focusable element on the page.",
            "Confirm the focused element is not entirely hidden by other content (sticky headers, cookie banners, chat widgets, modal overlays).",
            "Verify the user can see at least part of the focused element.",
            "Test at multiple viewport sizes including narrow mobile widths.",
        ],
    },
    "2.4.12": {
        "name": "Focus Not Obscured (Enhanced)",
        "level": "AAA",
        "fragment": "focus-not-obscured-enhanced",
        "checks": [
            "Tab through every focusable element on the page.",
            "Confirm the focused element is fully visible \u2014 no part of it is obscured by other content.",
            "Verify sticky elements, overlays, and floating elements never cover any portion of the focused element.",
            "Test at multiple viewport sizes.",
        ],
    },
    "2.4.13": {
        "name": "Focus Appearance",
        "level": "AAA",
        "fragment": "focus-appearance",
        "checks": [
            "Tab to each focusable element and inspect the focus indicator.",
            "Confirm the focus indicator has an area at least as large as a 2-CSS-pixel-thick perimeter around the focused element, OR a 2-CSS-pixel-thick line along the shortest side and not less than the line thickness in the other direction.",
            "Confirm the contrast between the focused-state pixels and the unfocused-state pixels is at least 3:1.",
            "Verify the indicator is not fully obscured by other content.",
        ],
    },
    "2.5.1": {
        "name": "Pointer Gestures",
        "level": "A",
        "fragment": "pointer-gestures",
        "checks": [
            "Identify any functionality activated by multipoint gestures (two-finger pinch, two-finger swipe) or path-based gestures (swipe, drag along a path).",
            "Confirm the same functionality is available through a single-point pointer action (tap, click) without a path.",
            "Verify the simpler alternative is documented or visually evident.",
            "Test with assistive input devices (switch control, head tracker) to confirm the simpler alternative is operable.",
        ],
    },
    "2.5.2": {
        "name": "Pointer Cancellation",
        "level": "A",
        "fragment": "pointer-cancellation",
        "checks": [
            "Identify each interactive element with single-pointer activation.",
            "Confirm one of these is true for each: activation occurs on the up-event (mouseup, pointerup, click), the user can abort by moving away before releasing, completion of the action can be undone, or the down-event activation is essential.",
            "Test by pressing down on a control then dragging away to confirm the action does not fire.",
            "Verify accidental activations can be reversed.",
        ],
    },
    "2.5.3": {
        "name": "Label in Name",
        "level": "A",
        "fragment": "label-in-name",
        "checks": [
            "Identify each user-interface component with a visible text label.",
            "Confirm the accessible name (what assistive technology announces) contains the visible label text.",
            "Verify the visible label is presented first or in the same order within the accessible name.",
            "Test with a screen reader: speaking the visible label should activate the control via voice control software.",
        ],
    },
    "2.5.4": {
        "name": "Motion Actuation",
        "level": "A",
        "fragment": "motion-actuation",
        "checks": [
            "Identify any functionality triggered by device motion (shaking, tilting) or user motion (swiping, gesturing in space).",
            "Confirm the same functionality is available through a standard user-interface control.",
            "Verify the user can disable motion actuation to prevent accidental activation.",
            "Confirm exceptions (an accelerometer-driven step counter) are essential to the function.",
        ],
    },
    "2.5.5": {
        "name": "Target Size (Enhanced)",
        "level": "AAA",
        "fragment": "target-size-enhanced",
        "checks": [
            "Identify each pointer-actuated target on the page (links, buttons, form controls, icons).",
            "Measure the size of each target's hit area.",
            "Confirm each target is at least 44 by 44 CSS pixels.",
            "Verify exceptions (inline links in text, browser-controlled targets) are limited.",
        ],
    },
    "2.5.6": {
        "name": "Concurrent Input Mechanisms",
        "level": "AAA",
        "fragment": "concurrent-input-mechanisms",
        "checks": [
            "Identify the input mechanisms supported by the page (mouse, keyboard, touch, voice, switch).",
            "Confirm none are restricted unless essential for security or essential for the activity.",
            "Test using multiple input methods in the same session (a touchscreen with an external keyboard).",
            "Verify the user can switch between input methods at any time without losing context.",
        ],
    },
    "2.5.7": {
        "name": "Dragging Movements",
        "level": "AA",
        "fragment": "dragging-movements",
        "checks": [
            "Identify each functionality activated by dragging (slider thumbs, drag-and-drop, sortable lists, draggable map controls).",
            "Confirm a single-pointer alternative is available: tap-to-position, increment/decrement buttons, keyboard-driven movement.",
            "Verify the alternative completes the same action as dragging.",
            "Confirm exceptions (signature capture, drawing apps) are genuinely path-essential.",
        ],
    },
    "2.5.8": {
        "name": "Target Size (Minimum)",
        "level": "AA",
        "fragment": "target-size-minimum",
        "checks": [
            "Identify each pointer-actuated target on the page.",
            "Measure the size of each target's hit area.",
            "Confirm each target is at least 24 by 24 CSS pixels, OR has at least 24-pixel spacing to neighboring targets.",
            "Verify exceptions (inline links in sentences, default browser-styled targets, equivalents available elsewhere) apply correctly.",
        ],
    },
    "3.1.1": {
        "name": "Language of Page",
        "level": "A",
        "fragment": "language-of-page",
        "checks": [
            "Inspect the html element's lang attribute.",
            "Confirm the lang attribute is present and contains a valid language code (en, en-US, fr, ja).",
            "Verify the language code matches the primary human language of the page content.",
            "Test with a screen reader to confirm pronunciation matches the declared language.",
        ],
    },
    "3.1.2": {
        "name": "Language of Parts",
        "level": "AA",
        "fragment": "language-of-parts",
        "checks": [
            "Identify each passage or phrase on the page that is in a different language from the surrounding content.",
            "Confirm each such passage has a lang attribute on a containing element.",
            "Verify the lang value is the correct language code for the passage.",
            "Test with a screen reader to confirm pronunciation switches appropriately.",
        ],
    },
    "3.1.3": {
        "name": "Unusual Words",
        "level": "AAA",
        "fragment": "unusual-words",
        "checks": [
            "Identify words used in unusual or restricted ways: jargon, idioms, slang, technical terms.",
            "Confirm a mechanism is available to find definitions: a glossary, an inline definition, a tooltip, or a link to a dictionary.",
            "Verify the definitions are accessible to screen readers and keyboard users.",
            "Test by hovering or focusing on each unusual word and confirming the definition is reachable.",
        ],
    },
    "3.1.4": {
        "name": "Abbreviations",
        "level": "AAA",
        "fragment": "abbreviations",
        "checks": [
            "Identify each abbreviation, acronym, or initialism on the page.",
            "Confirm a mechanism is available for identifying the expanded form: an abbr element with a title, an inline expansion at first occurrence, or a glossary.",
            "Verify the expansion is accessible to assistive technology.",
            "Test with a screen reader to confirm expansions are reachable.",
        ],
    },
    "3.1.5": {
        "name": "Reading Level",
        "level": "AAA",
        "fragment": "reading-level",
        "checks": [
            "Identify text that requires reading ability beyond a lower-secondary education level.",
            "Confirm a supplemental version is available: a simpler-language summary, illustrations, audio narration, or a video alternative.",
            "Run the text through a readability tool (Flesch-Kincaid, SMOG) to confirm the original reading level.",
            "Verify the simpler version conveys the same essential information.",
        ],
    },
    "3.1.6": {
        "name": "Pronunciation",
        "level": "AAA",
        "fragment": "pronunciation",
        "checks": [
            "Identify words whose meaning depends on pronunciation in context (homographs, words with non-obvious pronunciation, names).",
            "Confirm a mechanism is available for the pronunciation: a phonetic spelling, a pronunciation audio file, a ruby annotation, or an inline indication.",
            "Verify the mechanism is exposed to assistive technology.",
            "Test by reaching the pronunciation cue with a screen reader.",
        ],
    },
    "3.2.1": {
        "name": "On Focus",
        "level": "A",
        "fragment": "on-focus",
        "checks": [
            "Tab to each focusable element on the page.",
            "Confirm focusing on the element does not initiate a change of context (page navigation, form submission, opened window, major UI shift).",
            "Verify focus changes do not produce unexpected behavior that could disorient screen reader users.",
            "Test that any focus-driven changes are limited to expected behaviors (showing tooltips, highlighting current item).",
        ],
    },
    "3.2.2": {
        "name": "On Input",
        "level": "A",
        "fragment": "on-input",
        "checks": [
            "Identify each form input, dropdown, or selection control.",
            "Change the value of each control.",
            "Confirm the change does not automatically initiate a change of context unless the user has been advised of the behavior beforehand.",
            "Verify auto-submit behavior, if used, is announced or documented.",
        ],
    },
    "3.2.3": {
        "name": "Consistent Navigation",
        "level": "AA",
        "fragment": "consistent-navigation",
        "checks": [
            "Identify the navigation mechanisms used across the site (main nav, footer nav, breadcrumbs).",
            "Confirm each navigation mechanism appears in the same relative order on each page where it is present.",
            "Verify the same items appear in the same sequence within each instance.",
            "Test by visiting several pages and comparing the navigation.",
        ],
    },
    "3.2.4": {
        "name": "Consistent Identification",
        "level": "AA",
        "fragment": "consistent-identification",
        "checks": [
            "Identify components that have the same functionality across the site (search button, login link, shopping cart icon).",
            "Confirm these components are identified consistently (same icon, same label, same accessible name).",
            "Verify the same naming applies in tooltips, alt text, and aria-label attributes.",
            "Test by comparing the same component across multiple pages.",
        ],
    },
    "3.2.5": {
        "name": "Change on Request",
        "level": "AAA",
        "fragment": "change-on-request",
        "checks": [
            "Identify any automatic context changes (page navigation, popups, automatic refresh, automatic content update).",
            "Confirm such changes occur only when initiated by the user, OR a mechanism is available to turn them off.",
            "Verify auto-refresh, auto-redirect, and auto-launching behavior can be disabled.",
            "Test that all context changes are predictable and user-initiated.",
        ],
    },
    "3.2.6": {
        "name": "Consistent Help",
        "level": "A",
        "fragment": "consistent-help",
        "checks": [
            "Identify any help mechanisms present on the site (contact details, help link, self-help features, contact form, chat).",
            "Confirm these help mechanisms appear in the same relative order on each page where they exist.",
            "Verify the user can locate help reliably from any page.",
            "Test by visiting several pages and confirming help options are consistently placed.",
        ],
    },
    "3.3.1": {
        "name": "Error Identification",
        "level": "A",
        "fragment": "error-identification",
        "checks": [
            "Submit a form with intentional errors (invalid email, missing required field, value out of range).",
            "Confirm errors are identified to the user in text.",
            "Verify the field in error is described, not just flagged generically.",
            "Test with a screen reader: errors should be announced when surfaced.",
        ],
    },
    "3.3.2": {
        "name": "Labels or Instructions",
        "level": "A",
        "fragment": "labels-or-instructions",
        "checks": [
            "Identify each input control on the page.",
            "Confirm each control has a label or instruction explaining what to enter.",
            "Verify required fields, format expectations (date format, password rules), and field examples are clearly indicated.",
            "Test the form with a screen reader to confirm labels are announced when fields are reached.",
        ],
    },
    "3.3.3": {
        "name": "Error Suggestion",
        "level": "AA",
        "fragment": "error-suggestion",
        "checks": [
            "Trigger form validation errors.",
            "Confirm error messages provide suggestions for correction (\"email must contain @\", \"password must be at least 8 characters\").",
            "Verify suggestions are specific and actionable.",
            "Confirm exceptions (security risk, inability to suggest) are limited and the original requirement is restated.",
        ],
    },
    "3.3.4": {
        "name": "Error Prevention (Legal, Financial, Data)",
        "level": "AA",
        "fragment": "error-prevention-legal-financial-data",
        "checks": [
            "Identify forms that cause legal commitments, financial transactions, or data modifications.",
            "Confirm at least one of these is provided: submissions are reversible, the user must confirm the data before final submission, or the data is checked for validity with errors corrected before final submission.",
            "Test by submitting a transaction-class form and confirming the safeguard appears.",
            "Verify the safeguard is reachable by keyboard and screen reader.",
        ],
    },
    "3.3.5": {
        "name": "Help",
        "level": "AAA",
        "fragment": "help",
        "checks": [
            "Identify forms that require user input.",
            "Confirm context-sensitive help is available throughout the form.",
            "Verify help is reachable from each input (linked help text, popover, glossary).",
            "Test that help is accessible to screen readers and keyboard users.",
        ],
    },
    "3.3.6": {
        "name": "Error Prevention (All)",
        "level": "AAA",
        "fragment": "error-prevention-all",
        "checks": [
            "Identify all forms that submit user data of any kind.",
            "Confirm at least one of these is provided: submissions are reversible, data is checked with errors corrected before submission, or a confirmation step is provided before final submission.",
            "Test that the safeguard applies to every submission, not just transactional ones.",
            "Verify the safeguard is accessible to all users.",
        ],
    },
    "3.3.7": {
        "name": "Redundant Entry",
        "level": "A",
        "fragment": "redundant-entry",
        "checks": [
            "Identify multi-step processes (sign-up flows, checkouts) where the user enters data across multiple steps.",
            "Confirm previously entered information is auto-populated or made selectable in subsequent steps.",
            "Verify exceptions (re-entry essential for security, information that has changed, password fields) are limited.",
            "Test by completing a multi-step flow and confirming earlier-entered values are not asked for again.",
        ],
    },
    "3.3.8": {
        "name": "Accessible Authentication (Minimum)",
        "level": "AA",
        "fragment": "accessible-authentication-minimum",
        "checks": [
            "Identify authentication processes (login, password reset, multi-factor authentication).",
            "Confirm at least one method does not require a cognitive function test (remembering passwords, transcribing characters from CAPTCHA, solving puzzles).",
            "Verify alternatives like password managers, biometric login, magic-link email, or OAuth providers are accepted.",
            "Test that exceptions (object-recognition, personal-content recognition) apply only where genuinely needed.",
        ],
    },
    "3.3.9": {
        "name": "Accessible Authentication (Enhanced)",
        "level": "AAA",
        "fragment": "accessible-authentication-enhanced",
        "checks": [
            "Identify authentication processes.",
            "Confirm no step relies on a cognitive function test, including object-recognition or personal-content recognition.",
            "Verify only methods like passkeys, biometrics, or pre-existing trusted devices are required.",
            "Test that all users, regardless of cognitive ability, can authenticate without memorization or transcription.",
        ],
    },
    "4.1.2": {
        "name": "Name, Role, Value",
        "level": "A",
        "fragment": "name-role-value",
        "checks": [
            "Identify each user-interface component (form controls, links, custom widgets).",
            "Confirm each has a programmatically determinable name (accessible name) that conveys its purpose.",
            "Confirm each has a programmatically determinable role (button, link, checkbox, etc.).",
            "Confirm each has a programmatically determinable state and value where relevant (checked, expanded, selected, current value).",
            "Test with a screen reader and an accessibility inspector to verify the trio is correct for each component.",
        ],
    },
    "4.1.3": {
        "name": "Status Messages",
        "level": "AA",
        "fragment": "status-messages",
        "checks": [
            "Identify each status message presented to the user (form-submission feedback, search results count, dynamic loading indicators, error notifications).",
            "Confirm each is programmatically determinable through ARIA role (status, alert, log) or live regions.",
            "Verify status messages are announced by screen readers without moving focus.",
            "Test by triggering each status message with a screen reader running.",
        ],
    },
}


# Single-sentence summaries of each WCAG 2.2 success criterion,
# used in the Summary column of the ACR rollup. Length target is
# under 50 characters; longer entries wrap inside the cell.
dWcag22Summary = {
    "1.1.1": "Provide text alternatives for non-text content.",
    "1.2.1": "Alternatives for prerecorded audio-only and video-only.",
    "1.2.2": "Captions for prerecorded synchronized media.",
    "1.2.3": "Audio description or text alternative for prerecorded video.",
    "1.2.4": "Captions for live synchronized media.",
    "1.2.5": "Audio description for prerecorded video.",
    "1.2.6": "Sign language for prerecorded synchronized media.",
    "1.2.7": "Extended audio description for prerecorded video.",
    "1.2.8": "Text alternative for prerecorded synchronized media.",
    "1.2.9": "Text alternative for live audio.",
    "1.3.1": "Encode info, structure, and relationships in markup.",
    "1.3.2": "Reading order is meaningful in the DOM.",
    "1.3.3": "Instructions go beyond shape, color, or sound cues.",
    "1.3.4": "Content works in any screen orientation.",
    "1.3.5": "Mark inputs with their purpose for autofill.",
    "1.3.6": "Mark UI components by purpose for personalization.",
    "1.4.1": "Don't rely on color alone to convey information.",
    "1.4.2": "Provide audio control for auto-playing audio.",
    "1.4.3": "Text contrast is at least 4.5:1 (3:1 for large text).",
    "1.4.4": "Text scales to 200 percent without loss of content.",
    "1.4.5": "Use real text instead of images of text.",
    "1.4.6": "Text contrast is at least 7:1 (4.5:1 for large text).",
    "1.4.7": "Background audio is low or silenceable.",
    "1.4.8": "Customizable text presentation: color, width, spacing.",
    "1.4.9": "No images of text except logos.",
    "1.4.10": "Content reflows at 320 px without horizontal scroll.",
    "1.4.11": "UI and graphical objects have 3:1 contrast.",
    "1.4.12": "Text spacing can be increased without loss.",
    "1.4.13": "Hover and focus content is dismissible and persistent.",
    "2.1.1": "All functionality is keyboard-operable.",
    "2.1.2": "Keyboard focus can leave any component.",
    "2.1.3": "All functionality is keyboard-operable, no exceptions.",
    "2.1.4": "Single-key shortcuts are remappable or off.",
    "2.2.1": "Time limits can be turned off, adjusted, or extended.",
    "2.2.2": "Moving or auto-updating content can be paused.",
    "2.2.3": "No timing required (except real-time events).",
    "2.2.4": "Interruptions can be postponed or suppressed.",
    "2.2.5": "Re-authentication preserves the user's data.",
    "2.2.6": "User is warned before timeouts cause data loss.",
    "2.3.1": "No flashing more than three times per second.",
    "2.3.2": "No flashing more than three times per second, ever.",
    "2.3.3": "Motion animation can be disabled.",
    "2.4.1": "Mechanism to skip past repeated content.",
    "2.4.2": "Each page has a descriptive title.",
    "2.4.3": "Focus moves in a meaningful order.",
    "2.4.4": "Link purpose is clear from text or context.",
    "2.4.5": "Multiple ways to find pages on the site.",
    "2.4.6": "Headings and labels describe their content.",
    "2.4.7": "Keyboard focus is visually indicated.",
    "2.4.8": "User can identify their location in the site.",
    "2.4.9": "Link purpose is clear from link text alone.",
    "2.4.10": "Section headings organize long content.",
    "2.4.11": "Focused element is at least partly visible.",
    "2.4.12": "Focused element is fully visible.",
    "2.4.13": "Focus indicator has sufficient size and contrast.",
    "2.5.1": "Multipoint and path gestures have simpler alternatives.",
    "2.5.2": "Pointer activation can be aborted before release.",
    "2.5.3": "Accessible name contains the visible label text.",
    "2.5.4": "Motion-triggered functions have a UI alternative.",
    "2.5.5": "Pointer targets are at least 44 by 44 pixels.",
    "2.5.6": "All input mechanisms work concurrently.",
    "2.5.7": "Dragging movements have a single-pointer alternative.",
    "2.5.8": "Pointer targets are at least 24 by 24 pixels.",
    "3.1.1": "Page language is programmatically set.",
    "3.1.2": "Language of each part is programmatically set.",
    "3.1.3": "Definitions for unusual words are available.",
    "3.1.4": "Expansions for abbreviations are available.",
    "3.1.5": "A simpler version is available for advanced text.",
    "3.1.6": "Pronunciation cues are provided where needed.",
    "3.2.1": "Focusing a component does not change context.",
    "3.2.2": "Changing an input does not change context.",
    "3.2.3": "Navigation is consistent across pages.",
    "3.2.4": "Components are identified consistently.",
    "3.2.5": "Context changes occur only on user request.",
    "3.2.6": "Help is in the same place across pages.",
    "3.3.1": "Errors are identified and described in text.",
    "3.3.2": "Inputs have labels or instructions.",
    "3.3.3": "Errors include suggestions for correction.",
    "3.3.4": "Critical submissions can be reviewed or reversed.",
    "3.3.5": "Context-sensitive help is available.",
    "3.3.6": "All submissions can be reviewed or reversed.",
    "3.3.7": "Don't ask for the same data twice.",
    "3.3.8": "Authentication does not require a cognitive test.",
    "3.3.9": "Authentication does not require any cognitive test.",
    "4.1.2": "Components have name, role, and value programmatically.",
    "4.1.3": "Status messages are announced via ARIA roles.",
}

# --- Functions (alphabetical) ---

def flattenTarget(vTarget):
    """Convert an axe target value to a display string.
    axe target entries are strings for normal elements, or lists of strings
    for elements inside iframes or shadow DOM (path through each context)."""
    lParts = []
    vItem = None

    if isinstance(vTarget, list):
        for vItem in vTarget:
            if isinstance(vItem, list): lParts.append(" > ".join([str(v) for v in vItem]))
            else: lParts.append(str(vItem))
        return " >> ".join(lParts)
    return str(vTarget)



def buildCsvRows(dResults, dMetadata):
    dNode = {}
    dRow = {}
    dRule = {}
    iNodeCount = 0
    iRuleNodeIndex = 0
    lRows = []
    lWcagRefs = []
    sFailureSummary = ""
    sHelp = ""
    sHelpUrl = ""
    sHtml = ""
    sImpact = ""
    sOutcome = ""
    sRuleId = ""
    sStandardsRefs = ""
    sTags = ""
    sTarget = ""
    sWcagRefs = ""

    for sOutcome in aOutputSections:
        for dRule in dResults.get(sOutcome, []):
            iNodeCount = len(dRule.get("nodes", []))
            lWcagRefs = getWcagRefs(dRule)
            sHelp = str(dRule.get("help") or "")
            sHelpUrl = str(dRule.get("helpUrl") or "")
            sRuleId = str(dRule.get("id") or "")
            sStandardsRefs = " | ".join(getStandardsRefs(dRule))
            sTags = " | ".join(dRule.get("tags", []))
            sWcagRefs = " | ".join(lWcagRefs)
            if iNodeCount == 0:
                dRow = buildRowDict(dMetadata, sOutcome, dRule, sRuleId, sHelp, sHelpUrl, sTags, sWcagRefs, sStandardsRefs, 0, 0, "", "", "")
                lRows.append(dRow)
                continue
            iRuleNodeIndex = 0
            for dNode in dRule.get("nodes", []):
                iRuleNodeIndex += 1
                sFailureSummary = " ".join([str(v).strip() for v in str(dNode.get("failureSummary") or "").splitlines() if str(v).strip()])
                sHtml = str(dNode.get("html") or "")[:iCsvMaxHtmlLen]
                sImpact = str(dNode.get("impact") or dRule.get("impact") or "")
                sTarget = " | ".join([flattenTarget(v) for v in dNode.get("target", [])])
                dRow = buildRowDict(dMetadata, sOutcome, dRule, sRuleId, sHelp, sHelpUrl, sTags, sWcagRefs, sStandardsRefs, iNodeCount, iRuleNodeIndex, sTarget, sHtml, sFailureSummary, sImpact)
                lRows.append(dRow)
    lRows.sort(key=lambda dR: (dOutcomeRank.get(dR.get("outcome"), 99), dImpactRank.get(dR.get("impact"), 99), str(dR.get("ruleId") or ""), str(dR.get("target") or "")))
    return lRows


# Accessibility failure rate metric.
#
# For each page, we compute an impact-weighted violation density:
#
#   numerator   = 1*minorCount + 2*moderateCount + 3*seriousCount + 4*criticalCount
#   denominator = number of bytes in page.htm
#   rate        = iAccessibilityRateScale * numerator / denominator
#
# Counts are by node instance (one DOM element flagged), not by
# distinct rule. The constant tunes the result so that typical
# real-world pages produce a value the user can read at a glance.
# We use 100000 (=10^5) so the result is a percent: typical pages
# land in roughly 0%..100% with most being well under 100%, and a
# truly egregious page can exceed 100%. The percent framing is a
# display convention, not a strict mathematical claim that the
# ratio is bounded -- the underlying quantity is "impact-weighted
# instances per byte of page source," scaled into a friendly range.
#
# Lower is better. The metric's purpose is to give the page owner a
# single number to track over time. The goal is to reduce the rate
# from one scan to the next: accessibility is a journey, not a
# destination.
#
# When aggregating across multiple pages (the ACR-level rate), we
# sum the per-page numerators and sum the per-page denominators
# before dividing -- this is mathematically equivalent to a size-
# weighted average of per-page rates, and it correctly reflects
# that bigger pages with more content carry more weight in an
# accessibility profile.

iAccessibilityRateScale = 100000


def computePageImpactNumerator(dResults):
    """Sum 1*minor + 2*moderate + 3*serious + 4*critical instance
    counts across every violation rule on a page. Returns 0 if no
    violations.

    Note that an instance is one flagged DOM node. So a page with
    5 image elements lacking alt text contributes 5 to the count
    for the relevant impact level (typically critical or serious
    for image-alt). Best-practice and unknown-impact violations
    are ignored: they do not have a defined severity and shouldn't
    inflate the metric."""
    dImpactWeight = {"minor": 1, "moderate": 2, "serious": 3, "critical": 4}
    iSum = 0
    for dRule in dResults.get("violations", []):
        sImpact = str(dRule.get("impact") or "").lower()
        iWeight = dImpactWeight.get(sImpact, 0)
        if iWeight == 0: continue
        iSum += iWeight * len(dRule.get("nodes", []))
    return iSum


def computeAccessibilityFailureRate(iNumerator, iPageBytes):
    """Combine the impact-weighted numerator with the page-size
    denominator to produce the per-page accessibility failure rate
    as a percent. Returns 0.0 if iPageBytes is non-positive (avoids
    divide-by-zero and gracefully handles missing page.htm files).
    The return value is already in percent units; display it with a
    '%' sign. Typical pages produce values in 0..100; truly broken
    pages can exceed 100. Lower is better."""
    if iPageBytes <= 0: return 0.0
    return iAccessibilityRateScale * iNumerator / iPageBytes


def computePageBytes(pathOutputDir):
    """Return the byte size of the page.htm in the given per-page
    output folder. 0 if missing or unreadable. Used as the denom-
    inator of the accessibility failure rate."""
    try:
        path = pathlib.Path(pathOutputDir) / sSourceName
        if path.is_file():
            return path.stat().st_size
    except Exception:
        pass
    return 0


def captureViolationImages(page, dResults, pathOutputDir):
    """Capture an element-level screenshot for each violation node
    on the page and save as PNG in a 'violations/' subfolder.

    Returns a dict keyed by (sRuleId, iRuleNodeIndex) mapping to a
    relative path string like 'violations/image-001.png'. The
    relative path is what gets written into the report.xlsx
    Image-column cell as a hyperlink, so when the user clicks the
    cell the OS resolves it against the workbook's actual location
    and opens the PNG in the default viewer. Relative paths
    survive moving or zipping the page folder.

    Each successful capture increments a per-page counter (001,
    002, ...). The counter is dense (no gaps) -- file numbers
    reflect "successfully captured" not "node index." The mapping
    associates each successful capture with its source rule + node
    so the workbook writer can connect the right row to the right
    image.

    Failure modes are all silent (no log entry above DEBUG, no
    warning printed):
      - selector resolves to no element
      - selector resolves to multiple elements
      - element is hidden or zero-size (Playwright errors)
      - screenshot times out
      - any other Playwright exception

    The 'violations/' folder is created lazily on the first
    successful capture; if no captures succeed, no folder is
    created. If a 'violations/' folder already exists from a prior
    --force run, it is fully removed before this run starts so
    stale PNGs don't accumulate.
    """
    dImages = {}
    pathViolDir = pathlib.Path(pathOutputDir) / "violations"

    # Pre-clean: if a previous run left a violations/ folder, remove
    # it so this run's PNGs don't mix with stale ones. The page-
    # folder cleanup that --force triggers happens before scanUrl
    # runs, but be defensive in case of mid-run reruns.
    try:
        if pathViolDir.exists():
            shutil.rmtree(pathViolDir, ignore_errors=True)
    except Exception:
        pass

    iSeq = 0
    for dRule in dResults.get("violations", []) or []:
        sRuleId = str(dRule.get("id") or "")
        if not sRuleId: continue
        iRuleNodeIndex = 0
        for dNode in dRule.get("nodes", []) or []:
            iRuleNodeIndex += 1
            aTargets = dNode.get("target") or []
            if not aTargets: continue
            # axe's target is an array; the first element is the
            # primary CSS selector. Subsequent elements are present
            # for shadow-DOM piercing (axe nests them); for our
            # purposes the primary selector is what Playwright can
            # resolve directly.
            sCss = aTargets[0]
            if isinstance(sCss, list):
                # Defensive: nested array form. Take the first.
                sCss = sCss[0] if sCss else ""
            sCss = str(sCss or "")
            if not sCss: continue

            # Lazily create the violations/ folder on first attempt
            # at capture (one folder per page; we don't want an empty
            # violations/ folder if no captures succeed).
            iSeqAttempt = iSeq + 1
            sFileName = f"image-{iSeqAttempt:03d}.png"
            pathPng = pathViolDir / sFileName
            try:
                pathViolDir.mkdir(parents=True, exist_ok=True)
                # Playwright element-screenshot. .first picks the
                # first match if the selector resolves to multiple.
                # timeout in ms -- 2 seconds per node keeps a
                # pathological page from stalling the run.
                page.locator(sCss).first.screenshot(
                    path=str(pathPng), type="png", timeout=2000)
            except Exception:
                # Skip silently per project policy. Clean up an
                # empty file if Playwright created one before the
                # exception.
                try:
                    if pathPng.exists() and pathPng.stat().st_size == 0:
                        pathPng.unlink()
                except Exception:
                    pass
                continue

            # Successful capture: register and advance counter.
            iSeq = iSeqAttempt
            sRelPath = f"violations/{sFileName}"
            dImages[(sRuleId, iRuleNodeIndex)] = sRelPath

    return dImages


def buildConsoleSummary(dResults, dMetadata, sOutputDir):
    dRule = {}
    iCritical = 0
    iMinor = 0
    iModerate = 0
    iSerious = 0
    iTotal = 0
    lLines = []
    sPageTitle = ""
    sPageUrl = ""

    sPageTitle = str(dMetadata.get("pageTitle") or "")
    sPageUrl = str(dMetadata.get("pageUrl") or "")
    for dRule in dResults.get("violations", []):
        for dNode in dRule.get("nodes", []):
            sImpact = str(dNode.get("impact") or dRule.get("impact") or "")
            iTotal += 1
            if sImpact == "critical": iCritical += 1
            elif sImpact == "serious": iSerious += 1
            elif sImpact == "moderate": iModerate += 1
            elif sImpact == "minor": iMinor += 1
    lLines.append("")
    lLines.append(f"Page:      {sPageTitle}")
    lLines.append(f"url:       {sPageUrl}")
    lLines.append(f"Output:    {sOutputDir}")
    lLines.append(f"Timestamp: {str(dMetadata.get('scanTimestampUtc') or '')}")
    lLines.append("")
    lLines.append(f"Violations: {iTotal} nodes across {len(dResults.get('violations', []))} rules")
    if iTotal > 0:
        lLines.append(f"  Critical: {iCritical}  Serious: {iSerious}  Moderate: {iModerate}  Minor: {iMinor}")
    return "\n".join(lLines)


def buildCheckSummaryHtml(dNode):
    """Build the How to fix HTML for a single node, following the official Deque
    pattern from doc/examples/html-handlebars.md:
      - node.any  -> Fix ANY ONE of the following (only one fix needed)
      - node.all + node.none -> Fix ALL of the following (every fix required)
    Each check may also carry relatedNodes and a data dict (e.g. contrast ratio).
    """
    dCheck = {}
    lAll = []
    lAny = []
    lNone = []
    lParts = []
    sDataExtra = ""
    sMsg = ""
    sRelated = ""

    lAny = dNode.get("any", [])
    lAll = list(dNode.get("all", [])) + list(dNode.get("none", []))

    if not lAny and not lAll:
        return ""

    def buildCheckItem(dCheck):
        sMsg = str(dCheck.get("message") or "")
        lItemParts = []
        lItemParts.append(f"<li>{html.escape(sMsg)}")
        # Render extra data inline (e.g. color-contrast ratio, expected ratio)
        vData = dCheck.get("data")
        if isinstance(vData, dict):
            lDataParts = []
            for sKey, vVal in vData.items():
                if vVal is not None and str(vVal).strip():
                    lDataParts.append(f"{html.escape(str(sKey))}: {html.escape(str(vVal))}")
            if lDataParts:
                lItemParts.append(f" <span class='check-data'>({', '.join(lDataParts)})</span>")
        elif vData is not None and str(vData).strip():
            lItemParts.append(f" <span class='check-data'>({html.escape(str(vData))})</span>")
        # Render relatedNodes if present
        lRelated = dCheck.get("relatedNodes") or []
        if lRelated:
            lItemParts.append("<ul class='related-nodes'><li><em>Related paths:</em></li>")
            for dRelNode in lRelated:
                sRelTarget = " | ".join([flattenTarget(v) for v in (dRelNode.get("target") or [])])
                sRelHtml = str(dRelNode.get("html") or "")
                if sRelTarget:
                    lItemParts.append(f"<li><code class='selector'>{html.escape(sRelTarget)}</code>")
                    if sRelHtml: lItemParts.append(f" — <code>{html.escape(sRelHtml[:200])}</code>")
                    lItemParts.append("</li>")
            lItemParts.append("</ul>")
        lItemParts.append("</li>")
        return "".join(lItemParts)

    lParts.append("<div class='fix-box'>")
    if lAny:
        lParts.append("<p class='fix-label'>Fix any one of the following:</p>")
        lParts.append("<ul>")
        for dCheck in lAny: lParts.append(buildCheckItem(dCheck))
        lParts.append("</ul>")
    if lAll:
        lParts.append("<p class='fix-label'>Fix all of the following:</p>")
        lParts.append("<ul>")
        for dCheck in lAll: lParts.append(buildCheckItem(dCheck))
        lParts.append("</ul>")
    lParts.append("</div>")
    return "".join(lParts)


def buildOutcomeSummaryRows(dResults):
    dRule = {}
    iNodeCount = 0
    lRows = []
    sOutcome = ""

    for sOutcome in aReportSections:
        iNodeCount = 0
        for dRule in dResults.get(sOutcome, []): iNodeCount += len(dRule.get("nodes", []))
        lRows.append([dSectionHeadings[sOutcome], len(dResults.get(sOutcome, [])), iNodeCount])
    return lRows



def buildNarrativeSummary(dResults, dMetadata):
    """Return an HTML string with a plain-language summary of findings at ~9th grade level."""
    dCount = {}
    iCritical = 0
    iMinor = 0
    iModerate = 0
    iSerious = 0
    iTotal = 0
    lLeadRules = []
    lNextSteps = []
    lParts = []
    lTopRules = []
    lWcagLevels = {}
    sImpact = ""
    sPageTitle = ""
    sRuleId = ""
    sWcagNote = ""

    sPageTitle = str(dMetadata.get("pageTitle") or dMetadata.get("pageUrl") or "this page")
    lViolations = dResults.get("violations", [])
    iTotal = sum(len(dRule.get("nodes", [])) for dRule in lViolations)
    iRuleCount = len(lViolations)

    if iTotal == 0:
        lParts.append("<h2 id=\"summary\">Summary and next steps</h2>")
        lParts.append("<p>No accessibility violations were detected automatically on this page. "
                      "That is a good sign, but automated tools can only find some types of problems. "
                      "A manual review by a screen reader user and a keyboard-only user will find "
                      "issues that automated scans cannot detect.</p>")
        return "\n".join(lParts)

    # Count by impact
    for dRule in lViolations:
        sImpact = str(dRule.get("impact") or "")
        n = len(dRule.get("nodes", []))
        dCount[sImpact] = dCount.get(sImpact, 0) + n
    iCritical = dCount.get("critical", 0)
    iSerious   = dCount.get("serious",  0)
    iModerate  = dCount.get("moderate", 0)
    iMinor     = dCount.get("minor",    0)

    # Top rules by instance count
    lTopRules = sorted(lViolations, key=lambda r: len(r.get("nodes", [])), reverse=True)[:3]

    # Collect unique WCAG levels touched
    for dRule in lViolations:
        for sRef in getWcagRefs(dRule):
            tInfo = getWcagScInfo(sRef)
            if tInfo:
                sLevel = tInfo[1]
                lWcagLevels[sLevel] = lWcagLevels.get(sLevel, 0) + 1

    lParts.append("<h2 id=\"summary\">Summary and next steps</h2>")

    # Opening sentence
    if iTotal == 1:
        sLead = f"One accessibility problem was found on {html.escape(sPageTitle)}."
    else:
        sLead = f"{iTotal} accessibility problems were found across {iRuleCount} rule{'s' if iRuleCount != 1 else ''} on {html.escape(sPageTitle)}."
    lParts.append(f"<p>{sLead}</p>")

    # Severity breakdown
    sSevParts = []
    if iCritical: sSevParts.append(f"<strong>{iCritical} critical</strong>")
    if iSerious:  sSevParts.append(f"<strong>{iSerious} serious</strong>")
    if iModerate: sSevParts.append(f"{iModerate} moderate")
    if iMinor:    sSevParts.append(f"{iMinor} minor")
    if sSevParts:
        lParts.append(f"<p>By severity: {', '.join(sSevParts)}. "
                      "Critical and serious problems block or significantly harm people who rely on assistive technology "
                      "such as screen readers, keyboard navigation, or voice control. Fix these first.</p>")

    # WCAG note
    if "A" in lWcagLevels:
        sWcagNote = ("Some of these problems relate to WCAG Level A criteria, "
                     "which are the minimum required for basic accessibility compliance. ")
    elif "AA" in lWcagLevels:
        sWcagNote = ("Some of these problems relate to WCAG Level AA criteria, "
                     "which most accessibility laws and policies require. ")
    if sWcagNote:
        lParts.append(f"<p>{sWcagNote}See the WCAG column in each violation card for details.</p>")

    # Top rules
    if lTopRules:
        lParts.append("<p>The most common problems found:</p><ul>")
        for dRule in lTopRules:
            sRuleId = html.escape(str(dRule.get("id") or ""))
            sHelp   = html.escape(str(dRule.get("help") or ""))
            n       = len(dRule.get("nodes", []))
            sBp     = " (best practice)" if "best-practice" in dRule.get("tags", []) else ""
            lParts.append(f"<li><a href=\"#rule-id-{sRuleId}\">{sRuleId}</a> — {sHelp}{sBp}: "
                          f"<strong>{n}</strong> instance{'s' if n != 1 else ''}</li>")
        lParts.append("</ul>")

    # Next steps
    lNextSteps = [
        "Start with critical and serious violations — these have the most impact on users with disabilities.",
        "Use the Path and Snippet details in each violation card to locate the exact element in your code.",
        "Follow the How to fix guidance for each instance, then re-run this tool to confirm the fix.",
        "After fixing automated violations, do a manual review with a screen reader (NVDA or JAWS on Windows, "
        "VoiceOver on iOS or Mac) and with keyboard-only navigation.",
        "Automated tools like this one find roughly 30 to 40 percent of accessibility issues. "
        "Manual testing and user feedback are essential for full coverage.",
    ]
    lParts.append("<p><strong>Recommended next steps:</strong></p><ol>")
    for sStep in lNextSteps: lParts.append(f"<li>{sStep}</li>")
    lParts.append("</ol>")

    return "\n".join(lParts)

def buildReportHtml(dResults, dMetadata, lRows, nPageRate=0.0, iPageBytes=0, iImpactNumer=0):
    dCheck = {}
    dNode = {}
    dRule = {}
    dSummary = {}
    iCritical = 0
    iMinor = 0
    iModerate = 0
    iNodeIndex = 0
    iSerious = 0
    iTotalNodes = 0
    lCheck = []
    lFiles = []
    lImpactRows = []
    lParts = []
    lRuleLinks = []
    lRulesByFrequency = []
    lSummary = []
    lTopWcagRefs = []
    sBpTag = ""
    sCheckSection = ""
    sFixHtml = ""
    sImpactEmoji = ""
    sOutcome = ""
    sPageTitle = ""
    sRuleAnchor = ""
    sUrl = ""

    dSummary = getSummaryData(dResults, lRows)
    lImpactRows = dSummary.get("impactRows", [])
    lRulesByFrequency = dSummary.get("rulesByFrequency", [])
    lTopWcagRefs = dSummary.get("wcagRows", [])
    sPageTitle = html.escape(str(dMetadata.get("pageTitle") or sFallbackTitle))
    sUrl = html.escape(str(dMetadata.get("pageUrl") or ""))

    # Count nodes by impact for the banner
    for dRule in dResults.get("violations", []):
        for dNode in dRule.get("nodes", []):
            sImpact = str(dNode.get("impact") or dRule.get("impact") or "")
            iTotalNodes += 1
            if sImpact == "critical": iCritical += 1
            elif sImpact == "serious": iSerious += 1
            elif sImpact == "moderate": iModerate += 1
            elif sImpact == "minor": iMinor += 1

    lFiles = [
        f"<li><a href=\"{html.escape(sReportWorkbookName)}\">{html.escape(sReportWorkbookName)}</a> — Excel workbook</li>",
        f"<li><a href=\"{html.escape(sJsonName)}\">{html.escape(sJsonName)}</a> — full raw data</li>",
        f"<li><a href=\"{html.escape(sAccessibilityYamlName)}\">{html.escape(sAccessibilityYamlName)}</a> — ARIA accessibility tree</li>",
        f"<li><a href=\"{html.escape(sSourceName)}\">{html.escape(sSourceName)}</a> — saved page source</li>",
        f"<li><a href=\"{html.escape(sScreenshotName)}\">{html.escape(sScreenshotName)}</a> — page screenshot</li>",
        f"<li><a href=\"violations/\">violations/</a> — per-violation element screenshots, linked from the Image column of the Excel workbook</li>",
    ]

    # CSS
    sStyle = (
        "body{font-family:Segoe UI,Arial,sans-serif;line-height:1.6;margin:0;padding:0;background:#f5f5f5;color:#1a1a1a}"
        "main{max-width:960px;margin:0 auto;padding:1.5rem}"
        "a{color:#0060b9;text-decoration:underline}"
        "a:hover{text-decoration:none}"
        "h1{font-size:1.6rem;margin:0 0 .25rem 0}"
        "h2{font-size:1.25rem;border-bottom:2px solid #0060b9;padding-bottom:.25rem;margin-top:2rem}"
        "h3{font-size:1.1rem;margin-top:1.5rem}"
        "h4{font-size:1rem;margin:.75rem 0 .25rem 0}"
        "h5{font-size:.95rem;margin:.5rem 0 .2rem 0;color:#444}"
        "h6{font-size:.9rem;margin:.4rem 0 .15rem 0;color:#555}"
        "code,pre{font-family:Consolas,'Courier New',monospace;font-size:.88em}"
        "pre{background:#1e1e1e;color:#d4d4d4;padding:.75rem 1rem;border-radius:4px;overflow-x:auto;white-space:pre-wrap;word-break:break-all}"
        "table{border-collapse:collapse;width:100%;margin:.5rem 0 1rem 0;background:#fff;border-radius:4px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,.1)}"
        "th,td{border:1px solid #ddd;padding:.5rem .75rem;text-align:left;vertical-align:top}"
        "th{background:#0060b9;color:#fff;font-weight:600}"
        "tr:nth-child(even){background:#f9f9f9}"
        "nav ul{padding-left:1.25rem;margin:.25rem 0}"
        "nav li{margin:.15rem 0}"
        ".banner{background:#0060b9;color:#fff;padding:1rem 1.5rem;margin-bottom:1rem}"
        ".banner h1{color:#fff}"
        ".banner p{margin:.2rem 0;font-size:.95rem;opacity:.9}"
        ".impact-bar{display:flex;gap:.5rem;flex-wrap:wrap;margin:.5rem 0}"
        ".badge{display:inline-block;padding:.2rem .6rem;border-radius:3px;font-weight:700;font-size:.85rem;color:#fff}"
        ".badge-critical{background:#c00}"
        ".badge-serious{background:#c55000}"
        ".badge-moderate{background:#856404}"
        ".badge-minor{background:#555}"
        ".badge-pass{background:#1a7a1a}"
        ".rule-card{background:#fff;border-left:5px solid #999;border-radius:4px;margin:1rem 0;padding:1rem 1.25rem;box-shadow:0 1px 3px rgba(0,0,0,.08)}"
        ".rule-card.critical{border-color:#c00}"
        ".rule-card.serious{border-color:#c55000}"
        ".rule-card.moderate{border-color:#856404}"
        ".rule-card.minor{border-color:#555}"
        ".rule-meta{font-size:.88rem;color:#555;margin:.2rem 0}"
        ".node-card{background:#f8f8f8;border:1px solid #ddd;border-radius:3px;margin:.75rem 0;padding:.75rem 1rem}"
        ".node-label{font-size:.8rem;font-weight:700;text-transform:uppercase;letter-spacing:.05em;color:#666;margin-bottom:.2rem}"
        ".selector{font-family:Consolas,'Courier New',monospace;font-size:.85em;background:#eee;padding:.15rem .4rem;border-radius:2px;word-break:break-all}"
        ".fix-box{background:#fff8e1;border:1px solid #f0c040;border-radius:3px;padding:.5rem .75rem;margin:.5rem 0}"
        ".fix-box p{margin:.2rem 0;font-size:.9rem}"
        ".fix-label{font-weight:700;color:#6b4c00;font-size:.85rem;text-transform:uppercase;letter-spacing:.05em}"
        ".back-link{font-size:.85rem;float:right;margin-top:.2rem}"
        ".check-data{color:#555;font-size:.85em}"
        ".related-nodes{margin:.3rem 0 0 .5rem;font-size:.85em;list-style:disc;padding-left:1.2rem}"
        ".muted{color:#666;font-size:.9rem}"
        "dl.meta-grid{display:grid;grid-template-columns:max-content 1fr;gap:.2rem .75rem;font-size:.9rem}"
        "dt{font-weight:600}"
        "dd{margin:0 0 .2rem 0;word-break:break-all}"
        "#skip-link{position:absolute;left:-999px}"
        "#skip-link:focus{left:1rem;top:1rem;z-index:9999;background:#fff;padding:.5rem 1rem;border:2px solid #0060b9}"
        "details{margin:.5rem 0}"
        "summary{cursor:pointer;font-weight:600;padding:.4rem .5rem;background:#f0f0f0;border-radius:3px;user-select:none}"
        "summary:hover{background:#e0e0e0}"
        ".badge-bp{background:#5a3e8a}"
        ".tag-bp{display:inline-block;padding:.1rem .45rem;border-radius:3px;font-size:.78rem;font-weight:700;background:#ede7f6;color:#4a235a;margin-left:.4rem;vertical-align:middle}"
    )

    lParts.append("<!doctype html>")
    lParts.append("<html lang=\"en\">")
    lParts.append("<head>")
    lParts.append("<meta charset=\"utf-8\">")
    lParts.append("<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">")
    lParts.append(f"<title>Accessibility report — {sPageTitle}</title>")
    lParts.append(f"<style>{sStyle}</style>")
    lParts.append("</head>")
    lParts.append("<body>")
    lParts.append("<a href=\"#main-content\" id=\"skip-link\">Skip to main content</a>")

    # Banner
    lParts.append("<header class=\"banner\">")
    lParts.append(f"<h1>Accessibility report — {sPageTitle}</h1>")
    lParts.append(f"<p><a href=\"{sUrl}\" style=\"color:#cde\">{sUrl}</a></p>")
    lParts.append(f"<p>Scanned: {html.escape(str(dMetadata.get('scanTimestampUtc') or ''))} &nbsp;|&nbsp; {html.escape(str(dMetadata.get('browserVersion') or ''))}</p>")
    lParts.append("</header>")

    lParts.append("<main id=\"main-content\">")

    # Quick summary box
    iViolRules = len(dResults.get("violations", []))
    lParts.append("<section aria-label=\"Quick summary\">")
    lParts.append(f"<p><strong>{iTotalNodes} failed instance{'' if iTotalNodes==1 else 's'}</strong> across <strong>{iViolRules} rule{'' if iViolRules==1 else 's'}</strong> with accessibility violations.</p>")
    if iTotalNodes > 0:
        lParts.append("<div class=\"impact-bar\" aria-label=\"Counts by impact\">")
        if iCritical: lParts.append(f"<span class=\"badge badge-critical\">Critical: {iCritical}</span>")
        if iSerious:  lParts.append(f"<span class=\"badge badge-serious\">Serious: {iSerious}</span>")
        if iModerate: lParts.append(f"<span class=\"badge badge-moderate\">Moderate: {iModerate}</span>")
        if iMinor:    lParts.append(f"<span class=\"badge badge-minor\">Minor: {iMinor}</span>")
        lParts.append("</div>")
        lParts.append("<p class=\"muted\"><strong>Critical</strong> — blocks access completely. <strong>Serious</strong> — very hard to use. <strong>Moderate</strong> — causes difficulty. <strong>Minor</strong> — small problem with a workaround.</p>")
    # Accessibility failure rate. Lower is better; the goal is to
    # reduce this number from one scan to the next. Accessibility is
    # a journey, not a destination. Displayed as a percent so the
    # number is intuitive: a clean page is well under 1%, a typical
    # problematic page is double-digit %, and a truly bad page can
    # exceed 100%.
    lParts.append(f"<p><strong>Accessibility failure rate: {nPageRate:.1f}%</strong></p>")
    lParts.append(
        f"<p class=\"muted\">Computed as "
        f"<code>{iAccessibilityRateScale} &times; "
        f"impact-weighted instance count / page bytes</code> = "
        f"<code>{iAccessibilityRateScale} &times; {iImpactNumer} / "
        f"{iPageBytes}</code>. Impact weights: minor 1, moderate 2, "
        f"serious 3, critical 4. The constant {iAccessibilityRateScale} "
        f"scales the result into a percent for easy reading; values can "
        f"exceed 100% on heavily problematic pages. Lower is better. "
        f"Track this number over time; the goal is to reduce it on "
        f"every scan. Accessibility is a journey, not a destination.</p>"
    )
    lParts.append("</section>")

    # Table of contents
    lParts.append("<h2 id=\"table-of-contents\">Table of contents</h2>")
    lParts.append("<nav aria-label=\"Table of contents\"><ul>")
    lParts.append("<li><a href=\"#scan-details\">Scan details</a></li>")
    lParts.append("<li><a href=\"#patterns\">Patterns</a></li>")
    lParts.append("<li><a href=\"#violations\">Violations</a><ul>")
    for iRuleIdx, dRule in enumerate(dResults.get("violations", []), 1):
        sRuleAnchor = f"rule-{iRuleIdx}"
        sRuleLabel = html.escape(str(dRule.get("id") or ""))
        sRuleHelp = html.escape(str(dRule.get("help") or ""))
        sImpact = html.escape(str(dRule.get("impact") or ""))
        sImpactEmoji = dImpactEmoji.get(str(dRule.get("impact") or ""), "")
        sBpTagToc = " <span class=\"tag-bp\">Best Practice</span>" if "best-practice" in dRule.get("tags", []) else ""
        lParts.append(f"<li><a href=\"#{html.escape(sRuleAnchor)}\">{sImpactEmoji} {sRuleLabel}: {sRuleHelp}</a> <span class=\"badge badge-{sImpact}\" style=\"font-size:.75rem\">{sImpact}</span>{sBpTagToc}</li>")
    lParts.append("</ul></li>")
    lParts.append("<li><a href=\"#summary\">Summary and next steps</a></li>")
    lParts.append("<li><a href=\"#output-files\">Output files</a></li>")
    lParts.append("<li><a href=\"#glossary\">Glossary</a></li>")
    lParts.append("<li><a href=\"#resources\">Resources</a></li>")
    lParts.append("</ul></nav>")

    # Scan details
    lParts.append("<h2 id=\"scan-details\">Scan details</h2>")
    lParts.append("<dl class=\"meta-grid\">")
    for sLabel, sValue in [
        ["url scanned", str(dMetadata.get("pageUrl") or "")],
        ["Page title", str(dMetadata.get("pageTitle") or "")],
        ["Scan time (UTC)", str(dMetadata.get("scanTimestampUtc") or "")],
        ["Browser", str(dMetadata.get("browserVersion") or "")],
        ["Viewport", f"{dMetadata.get('viewportWidth') or 0} \u00d7 {dMetadata.get('viewportHeight') or 0} pixels"],
        ["Testing engine", str(dMetadata.get("axeSource") or "")],
        ["User agent", str(dMetadata.get("userAgent") or "")],
    ]:
        lParts.append(f"<dt>{html.escape(sLabel)}</dt><dd>{html.escape(sValue)}</dd>")
    lParts.append("</dl>")
    lParts.append("<h3>All outcome counts</h3>")
    lParts.append("<table><thead><tr><th>Outcome</th><th>Rules</th><th>Instances</th><th>What it means</th></tr></thead><tbody>")
    dOutcomeExplain = {
        "Violations": "Confirmed accessibility problems that need to be fixed.",
        "Needs Review": "Possible problems that need a person to check manually.",
        "Passes": "Rules that this page passed automatically.",
        "Inapplicable": "Rules that do not apply to this page.",
    }
    for lRow in buildOutcomeSummaryRows(dResults):
        sLabel = str(lRow[0])
        lParts.append(f"<tr><td><strong>{html.escape(sLabel)}</strong></td><td>{html.escape(str(lRow[1]))}</td><td>{html.escape(str(lRow[2]))}</td><td class=\"muted\">{html.escape(dOutcomeExplain.get(sLabel,''))}</td></tr>")
    lParts.append("</tbody></table>")

    # Patterns
    lParts.append("<h2 id=\"patterns\">Patterns</h2>")
    lParts.append("<p class=\"muted\">These tables show which problems appear most often, to help you decide where to start.</p>")
    lParts.append("<h3>Failed instances by impact</h3>")
    lParts.append("<table><thead><tr><th>Impact</th><th>Instances</th><th>What it means</th></tr></thead><tbody>")
    dImpactExplain = {
        "critical": "Blocks some users completely. Fix immediately.",
        "serious": "Very hard for some users to work around. High priority.",
        "moderate": "Causes real difficulty. Fix when possible.",
        "minor": "Small issue. A workaround usually exists.",
    }
    for lRow in lImpactRows:
        sImp = str(lRow[0])
        lParts.append(f"<tr><td><span class=\"badge badge-{html.escape(sImp)}\">{html.escape(sImp)}</span></td><td>{html.escape(str(lRow[1]))}</td><td class=\"muted\">{html.escape(dImpactExplain.get(sImp,''))}</td></tr>")
    if not lImpactRows: lParts.append("<tr><td colspan=\"3\">No violations found.</td></tr>")
    lParts.append("</tbody></table>")
    lParts.append("<h3>Most common rules by instance count</h3>")
    if lRulesByFrequency:
        lParts.append("<table><thead><tr><th>Rule</th><th>Failing elements</th></tr></thead><tbody>")
        for lRow in lRulesByFrequency:
            lParts.append(f"<tr><td><a href=\"#rule-id-{html.escape(str(lRow[0]))}\">{html.escape(str(lRow[0]))}</a></td><td>{html.escape(str(lRow[1]))}</td></tr>")
        lParts.append("</tbody></table>")
    else:
        lParts.append("<p>No violations found.</p>")
    lParts.append("<h3>Most common WCAG criteria</h3>")
    if lTopWcagRefs:
        lParts.append("<table><thead><tr><th>SC</th><th>Name</th><th>Level</th><th>Principle</th><th>Instances</th></tr></thead><tbody>")
        for lRow in lTopWcagRefs:
            sRef = str(lRow[0])
            tInfo = getWcagScInfo(sRef)
            sName = html.escape(tInfo[0]) if tInfo else ""
            sLevel = html.escape(tInfo[1]) if tInfo else ""
            sPrinciple = html.escape(tInfo[2]) if tInfo else ""
            lParts.append(f"<tr><td><a href=\"{html.escape(sWcagBaseUrl)}{html.escape(sRef)}/\">{html.escape(sRef)}</a></td><td>{sName}</td><td>{sLevel}</td><td>{sPrinciple}</td><td>{html.escape(str(lRow[1]))}</td></tr>")
        lParts.append("</tbody></table>")
    else:
        lParts.append("<p class=\"muted\">No violations with WCAG SC references were detected. Best-practice violations may still relate to WCAG criteria — see the per-rule WCAG notes in the Violations section below.</p>")

    # Violations
    lParts.append("<h2 id=\"violations\">Violations</h2>")
    lParts.append("<p class=\"muted\">Each section below is one failing rule. Inside each rule, every element on the page that failed is listed with its HTML, CSS path, and specific fix guidance.</p>")
    if not dResults.get("violations"):
        lParts.append("<p>No violations were detected on this page.</p>")
    else:
        for iRuleIdx, dRule in enumerate(dResults.get("violations", []), 1):
            sRuleAnchor = f"rule-{iRuleIdx}"
            sRuleId = str(dRule.get("id") or "")
            sImpact = str(dRule.get("impact") or "")
            iNodeCount = len(dRule.get("nodes", []))
            lRuleLinks = getRuleLinks(dRule)
            lWcagRefs = getWcagRefs(dRule)
            lParts.append(f"<section class=\"rule-card {html.escape(sImpact)}\" id=\"{html.escape(sRuleAnchor)}\" aria-label=\"Rule: {html.escape(sRuleId)}\">")
            lParts.append(f"<a href=\"#table-of-contents\" class=\"back-link\" aria-label=\"Back to table of contents\">&uarr; Contents</a>")
            sImpactEmoji = dImpactEmoji.get(sImpact, "")
            lParts.append(f"<h3 id=\"rule-id-{html.escape(sRuleId)}\">{sImpactEmoji} {html.escape(sRuleId)}: {html.escape(str(dRule.get('help') or ''))}</h3>")
            sBpTag = "<span class=\"tag-bp\">Best Practice</span>" if "best-practice" in dRule.get("tags", []) else ""
            lParts.append(f"<p class=\"rule-meta\"><span class=\"badge badge-{html.escape(sImpact)}\">{html.escape(sImpact)}</span>{sBpTag} &nbsp; {iNodeCount} failed instance{'' if iNodeCount==1 else 's'}</p>")
            lParts.append(f"<p>{html.escape(str(dRule.get('description') or ''))}</p>")
            if lWcagRefs:
                bIsBp = "best-practice" in dRule.get("tags", []) and not any(getWcagRef(t) for t in dRule.get("tags", []))
                lWcagParts = []
                for sRef in lWcagRefs:
                    tInfo = getWcagScInfo(sRef)
                    sName = f" — {html.escape(tInfo[0])} (Level {html.escape(tInfo[1])})" if tInfo else ""
                    lWcagParts.append(f"<a href=\"{html.escape(sWcagBaseUrl)}{html.escape(sRef)}/\">{html.escape(sRef)}</a>{sName}")
                sWcagLinks = "<br>".join(lWcagParts)
                sAdvisory = " <span class=\"muted\" style=\"font-style:italic\">(advisory — related criteria for this best-practice rule)</span>" if bIsBp else ""
                lParts.append(f"<p class=\"rule-meta\"><strong>WCAG:</strong>{sAdvisory}<br>{sWcagLinks} &nbsp;|&nbsp; <a href=\"{html.escape(str(dRule.get('helpUrl') or ''))}\">Deque rule documentation</a></p>")
            else:
                lParts.append(f"<p class=\"rule-meta\"><a href=\"{html.escape(str(dRule.get('helpUrl') or ''))}\">Deque rule documentation</a></p>")

            # Failing elements
            iNodeIndex = 0
            lParts.append(f"<details open><summary>{iNodeCount} failed instance{"" if iNodeCount==1 else "s"} — select to expand or collapse</summary>")
            for dNode in dRule.get("nodes", []):
                iNodeIndex += 1
                sNodeImpact = str(dNode.get("impact") or sImpact)
                sTarget = " | ".join([flattenTarget(v) for v in dNode.get("target", [])])
                sHtmlSnippet = str(dNode.get("html") or "")
                sFailSummary = str(dNode.get("failureSummary") or "")

                lParts.append("<div class=\"node-card\">")
                lParts.append(f"<p class=\"node-label\">Instance {iNodeIndex} of {iNodeCount}</p>")

                # CSS path
                if sTarget:
                    lParts.append("<p class=\"muted\" style=\"margin:.2rem 0 0\"><strong>Path</strong> — CSS selector for this element:</p>")
                    lParts.append(f"<p><code class=\"selector\">{html.escape(sTarget)}</code></p>")

                # HTML snippet
                if sHtmlSnippet:
                    lParts.append("<p class=\"muted\" style=\"margin:.4rem 0 0\"><strong>Snippet</strong> — HTML of this element:</p>")
                    lParts.append(f"<pre><code>{html.escape(sHtmlSnippet)}</code></pre>")

                # How to fix — use proper any/all/none semantics per Deque handlebars example
                sFixHtml = buildCheckSummaryHtml(dNode)
                if sFixHtml:
                    lParts.append(sFixHtml)
                elif sFailSummary:
                    lParts.append("<div class=\"fix-box\">")
                    lParts.append("<p class=\"fix-label\">How to fix</p>")
                    lParts.append(f"<pre>{html.escape(sFailSummary)}</pre>")
                    lParts.append("</div>")

                # Impact at node level if different from rule
                if sNodeImpact and sNodeImpact != sImpact:
                    lParts.append(f"<p class=\"muted\">Element impact: <span class=\"badge badge-{html.escape(sNodeImpact)}\">{html.escape(sNodeImpact)}</span></p>")

                lParts.append("</div>")
            lParts.append("</details>")
            lParts.append("</section>")

    # Narrative summary
    lParts.append(buildNarrativeSummary(dResults, dMetadata))

    # Output files
    lParts.append("<h2 id=\"output-files\">Output files</h2>")
    lParts.append("<p class=\"muted\">All files are saved in the same folder as this report.</p>")
    lParts.append(f"<ul>{''.join(lFiles)}</ul>")

    # Glossary
    lParts.append("<h2 id=\"glossary\">Glossary</h2>")
    lParts.append("<dl>")
    for lRow in aGlossaryRows: lParts.append(f"<dt><strong>{html.escape(str(lRow[0]))}</strong></dt><dd>{html.escape(str(lRow[1]))}</dd>")
    lParts.append("</dl>")

    # Resources
    lParts.append("<h2 id=\"resources\">Resources</h2>")
    lParts.append("<ul>")
    lParts.append(f"<li><a href=\"{html.escape(sWcagBaseUrl)}\">WCAG 2.2 Understanding — what each criterion means and how to meet it</a></li>")
    lParts.append(f"<li><a href=\"{html.escape(sAccessibilityInsightsUrl)}\">Accessibility Insights — free browser extension for manual testing</a></li>")
    lParts.append(f"<li><a href=\"{html.escape(sMsAccessibilityUrl)}\">Microsoft Accessibility — guidance and tools from Microsoft</a></li>")
    lParts.append("<li><a href=\"https://dequeuniversity.com/rules/axe/4.11/\">Deque — full list of axe-core rules with explanations</a></li>")
    lParts.append("</ul>")

    lParts.append("</main>")
    lParts.append("</body>")
    lParts.append("</html>")
    return "\n".join(lParts)


def buildRowDict(dMetadata, sOutcome, dRule, sRuleId, sHelp, sHelpUrl, sTags, sWcagRefs, sStandardsRefs, iNodeCount, iNodeIndex, sTarget, sHtml, sFailureSummary, sImpact=""):
    return {
        "axeSource": str(dMetadata.get("axeSource") or ""),
        "browserVersion": str(dMetadata.get("browserVersion") or ""),
        "description": str(dRule.get("description") or ""),
        "failureSummary": sFailureSummary,
        "help": sHelp,
        "helpUrl": sHelpUrl,
        "html": sHtml,
        "impact": sImpact or str(dRule.get("impact") or ""),
        "outcome": sOutcome,
        "pageTitle": str(dMetadata.get("pageTitle") or ""),
        "pageUrl": str(dMetadata.get("pageUrl") or ""),
        "ruleId": sRuleId,
        "ruleNodeCount": iNodeCount,
        "ruleNodeIndex": iNodeIndex,
        "scanTimestampUtc": str(dMetadata.get("scanTimestampUtc") or ""),
        "standardsRefs": sStandardsRefs,
        "tags": sTags,
        "target": sTarget,
        "wcagRefs": sWcagRefs,
    }


def cleanPreviousTempDirs():
    """Remove _MEI* temporary directories in %TEMP% left by previous runs of
    this program. The current run's own directory (sys._MEIPASS) is skipped.
    Directories belonging to currently running PyInstaller applications cannot
    be deleted because their DLLs are locked in memory; shutil.rmtree will
    raise an exception on the first locked file and the whole directory is left
    intact. Only fully-exited runs leave unlocked directories, so this is safe
    to run against all _MEI* siblings regardless of which program created them.
    Has no effect when running from source (sys._MEIPASS is absent)."""
    import shutil
    pathDir = None
    pathTemp = None
    sCurrentMei = ""

    sCurrentMei = getattr(sys, "_MEIPASS", "")
    if not sCurrentMei: return
    pathTemp = pathlib.Path(sCurrentMei).parent
    for pathDir in pathTemp.glob("_MEI*"):
        if pathDir == pathlib.Path(sCurrentMei): continue
        try:
            shutil.rmtree(str(pathDir))
        except Exception:
            pass


def chooseOutputDir(pathBaseDir, sPageTitle, bForce=False):
    """Decide the per-page output folder.

    Returns the Path of the folder to write into, OR None if the folder
    already exists and bForce is False (caller should skip this url).

    When the folder does not exist, it is created.
    When the folder exists and bForce is True, its contents are deleted
    so the run starts from a clean slate; the folder itself is reused.
    """
    pathCandidate = None
    sBaseName = ""

    sBaseName = getSafeTitle(sPageTitle)
    pathCandidate = pathBaseDir / sBaseName
    if not pathCandidate.exists():
        pathCandidate.mkdir(parents=True, exist_ok=False)
        return pathCandidate
    if bForce:
        # Existing folder, --force is on: empty its contents and reuse
        # the folder. Delete files and subdirectories defensively; if
        # any single entry can't be removed (e.g. a file is open in
        # another process), log it and continue -- the new run will
        # overwrite what it can.
        for child in pathCandidate.iterdir():
            try:
                if child.is_dir() and not child.is_symlink():
                    shutil.rmtree(child)
                else:
                    child.unlink()
            except Exception as ex:
                logger.info(f"Could not remove {child} while emptying "
                    f"{pathCandidate} for --force: {ex}")
        return pathCandidate
    # Existing folder, no --force: caller should skip.
    return None


def ensureSuccess(bCondition, sMessage):
    if not bCondition: raise RuntimeError(sMessage)
    return True


def waitForDevToolsPort(sUserDataDir, iTimeoutSeconds=30):
    """Wait for Edge to write DevToolsActivePort into its user-data
    directory and return the port number as a string.

    When Edge is launched with --remote-debugging-port=0, it picks a
    free port and writes the chosen port number (followed by a
    websocket URL on a second line) to a file named
    DevToolsActivePort inside its --user-data-dir. We poll for that
    file and return the first line.

    Returns "" on timeout. Caller should check.
    """
    sPortFile = os.path.join(sUserDataDir, "DevToolsActivePort")
    nDeadline = time.time() + iTimeoutSeconds
    while time.time() < nDeadline:
        if os.path.isfile(sPortFile):
            try:
                with open(sPortFile, "r", encoding="utf-8") as fIn:
                    sFirstLine = (fIn.readline() or "").strip()
                if sFirstLine and sFirstLine.isdigit():
                    return sFirstLine
            except Exception:
                pass
        time.sleep(0.2)
    return ""


def getEdgeExecutablePath():
    """Return the absolute path of msedge.exe on the system, or None
    if not found. Tried in this order:
      1. %ProgramFiles(x86)%\\Microsoft\\Edge\\Application\\msedge.exe
      2. %ProgramFiles%\\Microsoft\\Edge\\Application\\msedge.exe
      3. %LocalAppData%\\Microsoft\\Edge\\Application\\msedge.exe
    On Windows 10/11, Edge is installed system-wide at the first
    location for x64 installs (Edge is a 32-bit-folder install
    historically). We deliberately do NOT walk the registry --
    file probing is good enough for the common case and avoids
    pulling in winreg.
    """
    lCandidates = []
    sProgFiles86 = os.environ.get("ProgramFiles(x86)", "")
    sProgFiles = os.environ.get("ProgramFiles", "")
    sLocalAppData = os.environ.get("LOCALAPPDATA", "")
    if sProgFiles86: lCandidates.append(os.path.join(sProgFiles86, "Microsoft", "Edge", "Application", "msedge.exe"))
    if sProgFiles: lCandidates.append(os.path.join(sProgFiles, "Microsoft", "Edge", "Application", "msedge.exe"))
    if sLocalAppData: lCandidates.append(os.path.join(sLocalAppData, "Microsoft", "Edge", "Application", "msedge.exe"))
    for sCandidate in lCandidates:
        if os.path.isfile(sCandidate): return sCandidate
    return None


def getEdgeUserDataDir():
    """Return the path of the user's default Microsoft Edge user-data
    directory on Windows, or None if it cannot be determined.

    On Windows, Edge stores its user data (cookies, history, saved
    passwords, extensions, profile preferences) under
    %LOCALAPPDATA%\\Microsoft\\Edge\\User Data. This is the parent
    directory; specific profiles (Default, Profile 1, etc.) are
    subdirectories. Playwright's launch_persistent_context() expects
    this parent directory and selects "Default" by default; other
    profiles can be picked via --profile-directory=... in args.
    """
    sLocalAppData = os.environ.get("LOCALAPPDATA", "")
    if not sLocalAppData: return None
    sCandidate = os.path.join(sLocalAppData, "Microsoft", "Edge", "User Data")
    if not os.path.isdir(sCandidate): return None
    return sCandidate


def isEdgeRunning():
    """Return True if msedge.exe appears to be running, False otherwise.

    Used to fail fast with a clear message before urlCheck attempts
    to launch a persistent-context Edge against a user-data directory
    that the existing Edge process holds locked. Without this check,
    Playwright's launch fails with an opaque error that's hard for a
    blind user to interpret.

    Implementation: shell out to tasklist.exe (built into Windows)
    with a filter for msedge.exe. CREATE_NO_WINDOW so the cmd
    window doesn't flash for users in CLI mode.
    """
    try:
        iCreateNoWindow = 0x08000000
        result = subprocess.run(
            ["tasklist.exe", "/FI", "IMAGENAME eq msedge.exe", "/NH"],
            capture_output=True, text=True, timeout=5,
            creationflags=iCreateNoWindow)
        sOut = (result.stdout or "").lower()
        return "msedge.exe" in sOut
    except Exception:
        # If we can't tell, assume not running and let the launch
        # attempt produce its own error if there's a conflict. This
        # is safer than blocking the user erroneously.
        return False


def applyWebdriverOverride(context):
    """Add the navigator.webdriver=undefined init script to a context.

    Used both at initial launch and on every reconnect during -t
    (temp-profile + disconnect/reconnect) runs. The override runs on
    every new document before any site JavaScript can read it,
    producing navigator.webdriver=undefined without using a command-
    line flag (which would itself trigger the "unsupported flag"
    warning bar).

    The 'configurable: true' allows later test pages to redefine the
    property without TypeError.
    """
    try:
        context.add_init_script(
            "Object.defineProperty(navigator, 'webdriver', "
            "{get: () => undefined, configurable: true});")
    except Exception as ex:
        logger.warn(f"Could not add navigator.webdriver init script: {ex}")


def fetchText(sUrl):
    dHeaders = {}
    request = None
    response = None
    sText = ""

    dHeaders = {"User-Agent": sUserAgent}
    request = urllib.request.Request(sUrl, headers=dHeaders)
    with urllib.request.urlopen(request, timeout=iCdnTimeoutSec) as response: sText = response.read().decode("utf-8")
    return sText


def getAxeScript(page, sPreFetchedContent=""):
    ex = None
    sAxeScript = ""
    sUrl = ""

    # If we have pre-fetched content, inject it directly — this bypasses CDN
    # reachability issues and avoids CSP blocks on external script urls.
    if sPreFetchedContent:
        for sUrl in aAxeCdnUrls:
            try:
                page.add_script_tag(content=sPreFetchedContent)
                return sUrl
            except Exception as ex:
                break
    # Fall back to url injection then content fetch per CDN url
    for sUrl in aAxeCdnUrls:
        try:
            page.add_script_tag(url=sUrl)
            return sUrl
        except Exception as ex:
            continue
    for sUrl in aAxeCdnUrls:
        try:
            sAxeScript = fetchText(sUrl)
            page.add_script_tag(content=sAxeScript)
            return sUrl
        except Exception as ex:
            continue
    raise RuntimeError("Unable to load axe-core from the configured CDN urls.")


def getImpactRows(lRows):
    dCounts = {}
    dRow = {}
    lOutput = []
    sImpact = ""

    for dRow in lRows:
        sImpact = str(dRow.get("impact") or "")
        if not sImpact: continue
        dCounts[sImpact] = int(dCounts.get(sImpact, 0)) + 1
    for sImpact in sorted(dCounts.keys(), key=lambda sK: dImpactRank.get(sK, 99)): lOutput.append([sImpact, dCounts[sImpact]])
    return lOutput


def firstLine(sText):
    """Trim and shorten a message for inline display next to a url.

    Exception messages can span multiple lines or be very long; for
    inline rendering ("https://x.com: <reason>") we want a single
    short string. Returns "" if the input is empty/None.
    """
    iMaxLen = 120
    if not sText: return ""
    iNl = -1
    for sNl in ("\r", "\n"):
        i = sText.find(sNl)
        if i >= 0 and (iNl < 0 or i < iNl): iNl = i
    if iNl >= 0: sText = sText[:iNl]
    sText = sText.strip()
    if len(sText) > iMaxLen: sText = sText[:iMaxLen - 3] + "..."
    return sText


def writeTextFile(pathFile, sContent):
    """Write a text file in the canonical CRLF + UTF-8 BOM format
    optimized for Windows users. Use this for any file the user is
    expected to view or edit: .htm, .json, .yaml, .ini, .log, .md,
    .csv, .txt. Binary formats (.xlsx, .docx, .pdf, .png, .jpg) are
    written by their respective libraries and are not affected.

    The function ensures parent directories exist, normalizes any
    LF or CR-only line endings in the content to CRLF, and writes
    with utf-8-sig (which prepends the BOM bytes 0xEF 0xBB 0xBF).
    The path can be a string or pathlib.Path.

    Returns the resolved string path, for symmetry with prior
    write_text calls.
    """
    path = pathlib.Path(pathFile)
    try: path.parent.mkdir(parents=True, exist_ok=True)
    except Exception: pass
    # Normalize line endings: convert any \r\n or \r alone to \n
    # first, then write with newline="\r\n" so Python's text mode
    # handles the conversion. This handles mixed-ending content
    # that might come from network sources (page.htm etc.).
    sNormalized = sContent.replace("\r\n", "\n").replace("\r", "\n")
    with path.open("w", encoding="utf-8-sig", newline="\r\n") as fOut:
        fOut.write(sNormalized)
    return str(path)


def getOsVersion():
    """Return a succinct OS version string for display in reports.

    On Windows, returns "Windows 11" or "Windows 10" using
    platform.win32_ver() to disambiguate (platform.release() can
    report '10' even on Windows 11 depending on the Python build,
    because the kernel-version major didn't change between the
    two). On non-Windows systems, returns a similarly compact form
    like "macOS 14.6" or "Linux 6.5". On any failure, returns an
    empty string -- the caller treats that as "unknown" and
    suppresses the field from output."""
    try:
        sSystem = platform.system() or ""
        if sSystem == "Windows":
            # win32_ver returns (release, version, csd, ptype). The
            # version field is the build number string; if it
            # starts with "10.0.22000" or higher, it's Windows 11.
            try:
                aWin = platform.win32_ver()
                sBuild = (aWin[1] or "") if len(aWin) > 1 else ""
                # Build numbers >= 22000 = Windows 11; lower = Windows 10
                aParts = sBuild.split(".")
                iBuild = int(aParts[2]) if len(aParts) >= 3 and aParts[2].isdigit() else 0
                if iBuild >= 22000: return "Windows 11"
                if iBuild > 0: return "Windows 10"
            except Exception:
                pass
            # Fallback: use whatever release reports
            sRel = platform.release() or ""
            return f"Windows {sRel}".strip() or "Windows"
        if sSystem == "Darwin":
            sRel = platform.mac_ver()[0] or platform.release() or ""
            return f"macOS {sRel}".strip() or "macOS"
        if sSystem == "Linux":
            sRel = platform.release() or ""
            return f"Linux {sRel}".strip() or "Linux"
        return sSystem or ""
    except Exception:
        return ""


def getNormalizedUrl(sInput):
    """Convert a user-supplied input to a fully qualified url or local file URI.
    Accepts bare domains (microsoft.com), IP addresses, paths with or without
    scheme, and local files. Any existing local file is treated as HTML to be
    loaded in the browser, regardless of its extension; the user is
    responsible for supplying an HTML-renderable file. Unrecognised inputs
    are returned unchanged."""
    path = None
    sCandidate = ""
    sSuffix = ""
    sLower = ""

    sCandidate = sInput.strip().strip('"')
    sLower = sCandidate.lower()
    # Already has a scheme — pass through as-is
    if "://" in sLower: return sCandidate
    # Existing local file -- treat as HTML regardless of extension. urlCheck's
    # contract leaves it to the user to ensure the file is HTML-renderable.
    try:
        path = pathlib.Path(sCandidate).expanduser().resolve()
        if path.exists() and path.is_file():
            return path.as_uri()
    except Exception:
        pass
    # Has a recognised local file extension and no path separator — treat as local filename
    sSuffix = pathlib.Path(sCandidate.split("?")[0].split("#")[0]).suffix.lower()
    if sSuffix in aLocalFileExts and "/" not in sCandidate: return sCandidate
    # Looks like an IP address (optional port/path)
    if reIpAddress.match(sCandidate): return f"https://{sCandidate}"
    # Looks like a domain name (optional port/path)
    if reDomain.match(sCandidate): return f"https://{sCandidate}"
    return sCandidate


def getPageSnapshot(page, sPageUrl):
    dCss = {}
    dScript = {}
    sBaseUrl = ""
    sHref = ""
    sHtml = ""
    sInlineHtml = ""
    sSrc = ""

    sHtml = str(page.content())
    try:
        sBaseUrl = str(page.url or sPageUrl)
        for sHref in page.eval_on_selector_all("link[rel='stylesheet'][href]", "elements => elements.map(e => e.getAttribute('href'))"):
            try: dCss[urllib.parse.urljoin(sBaseUrl, str(sHref))] = fetchText(urllib.parse.urljoin(sBaseUrl, str(sHref)))
            except Exception: continue
        for sSrc in page.eval_on_selector_all("script[src]", "elements => elements.map(e => e.getAttribute('src'))"):
            try: dScript[urllib.parse.urljoin(sBaseUrl, str(sSrc))] = fetchText(urllib.parse.urljoin(sBaseUrl, str(sSrc)))
            except Exception: continue
    except Exception:
        return sHtml
    sInlineHtml = sHtml
    for sHref in list(dCss.keys()):
        sCssContent = f'<style data-pagecheck-source="{html.escape(sHref)}">\n{dCss[sHref]}\n</style>'
        sInlineHtml = re.sub(rf'<link\b[^>]*href=["\']{re.escape(sHref)}["\'][^>]*>', lambda m, s=sCssContent: s, sInlineHtml, flags=re.IGNORECASE)
    for sSrc in list(dScript.keys()):
        sScriptContent = f'<script data-pagecheck-source="{html.escape(sSrc)}">\n{dScript[sSrc]}\n</script>'
        sInlineHtml = re.sub(rf'<script\b[^>]*src=["\']{re.escape(sSrc)}["\'][^>]*>\s*</script>', lambda m, s=sScriptContent: s, sInlineHtml, flags=re.IGNORECASE)
    return sInlineHtml


def getRuleFrequencyRows(lRows):
    dCounts = {}
    dRow = {}
    lOutput = []
    sRuleId = ""

    for dRow in lRows:
        sRuleId = str(dRow.get("ruleId") or "")
        if not sRuleId: continue
        dCounts[sRuleId] = int(dCounts.get(sRuleId, 0)) + 1
    for sRuleId in sorted(dCounts.keys(), key=lambda sK: (-int(dCounts.get(sK, 0)), sK)): lOutput.append([sRuleId, dCounts[sRuleId]])
    return lOutput[:20]


def getRuleLinks(dRule):
    lLinks = []
    sRuleId = ""
    sTag = ""
    sWcagRef = ""

    sRuleId = html.escape(str(dRule.get("id") or ""))
    lLinks.append(f"<li><a href=\"{sMsAccessibilityUrl}\">Microsoft Accessibility guidance</a> for rule <code>{sRuleId}</code></li>")
    for sTag in dRule.get("tags", []):
        if not any(str(sTag).startswith(sPrefix) for sPrefix in aRelevantTagPrefixes): continue
        sWcagRef = getWcagRef(str(sTag))
        if sWcagRef:
            lLinks.append(f"<li><a href=\"{sWcagBaseUrl}{html.escape(sWcagRef)}/\">WCAG Understanding: {html.escape(sWcagRef)}</a></li>")
            continue
        lLinks.append(f"<li>{html.escape(str(sTag))}</li>")
    if len(lLinks) == 1: lLinks.append(f"<li><a href=\"{sAccessibilityInsightsUrl}\">Accessibility Insights overview</a></li>")
    return lLinks


def getSafeTitle(sTitle):
    """Sanitize a page title for use as a Windows folder name.

    Goals:
      - Preserve the title's natural readability when seen in File
        Explorer (where the user actually views output). That means
        keeping original capitalization and keeping spaces between
        words rather than replacing them with dashes. Word-separator
        style chosen by the original page title is preserved.
      - Strip only the characters Windows forbids in folder names:
        < > : " / \\ | ? * and control characters. Trailing dots and
        spaces are also illegal at the end of a folder name.
      - Avoid Windows reserved device names (CON, PRN, AUX, NUL,
        COM1..COM9, LPT1..LPT9).
      - Cap length at iMaxTitleLen.

    Returns a string. Always returns a non-empty name; falls back to
    sFallbackTitle for empty or pathologically-bad input.
    """
    sName = ""
    sBase = ""

    sBase = str(sTitle or sFallbackTitle).strip()
    # Remove characters that Windows forbids in file/folder names.
    # Keep spaces, alphanumerics, and most punctuation including dashes,
    # underscores, dots, parentheses, brackets that are not in the
    # forbidden set, etc. Original capitalization is preserved.
    sName = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", sBase)
    # Collapse runs of whitespace to a single space (so a title with
    # internal newlines or tabs doesn't produce multi-space gaps).
    sName = re.sub(r"\s+", " ", sName).strip()
    # Strip trailing dots and spaces -- both illegal in Windows folder
    # names. Keep stripping until clean.
    while sName.endswith(".") or sName.endswith(" "):
        sName = sName[:-1]
    # Reserved device-name check (case-insensitive). If the base name
    # (before any extension-like suffix) matches, prepend an underscore
    # to defang it.
    sCheck = sName.upper().split(".")[0]
    aReserved = {"CON","PRN","AUX","NUL",
                 "COM1","COM2","COM3","COM4","COM5","COM6","COM7","COM8","COM9",
                 "LPT1","LPT2","LPT3","LPT4","LPT5","LPT6","LPT7","LPT8","LPT9"}
    if sCheck in aReserved: sName = "_" + sName
    if not sName: sName = sFallbackTitle
    # Length cap. Strip any new trailing dots/spaces caused by the cut.
    sName = sName[:iMaxTitleLen]
    while sName.endswith(".") or sName.endswith(" "):
        sName = sName[:-1]
    if not sName: sName = sFallbackTitle
    return sName


def getInitialBrowseDir(sFieldText):
    """Return a directory to use as the initial location of a file or folder
    picker.

    The strategy follows Microsoft's guidance: start at the user's most
    recent / current choice when one exists, otherwise fall back to the
    user's Documents folder.

    sFieldText is the current text of the source-input or output-directory
    field. The function:

      - returns Documents if sFieldText is empty or looks like a url or
        domain (urlCheck source field can contain those, which are not
        filesystem paths)
      - looks at the first space-separated token if sFieldText has many
      - returns sFieldText itself if it points to an existing directory
      - returns the parent of sFieldText if it points to an existing file
      - returns the parent of sFieldText if only the parent exists
      - falls back to Documents otherwise

    Returns a string. Always returns a non-empty path that the OS knows.
    """
    sCandidate = ""
    sFirstToken = ""
    sParent = ""

    sFieldText = (sFieldText or "").strip()
    if not sFieldText:
        return getDocumentsDir()
    # Source field may contain space-separated tokens. Inspect the first.
    sFirstToken = sFieldText.split()[0] if sFieldText.split() else ""
    if not sFirstToken:
        return getDocumentsDir()
    # Strip surrounding quotes a user may have typed.
    if len(sFirstToken) >= 2 and sFirstToken[0] == '"' and sFirstToken[-1] == '"':
        sFirstToken = sFirstToken[1:-1]
    # url or domain -- not a filesystem location. Use Documents.
    if re.match(r"^[a-z][a-z0-9+.-]*://", sFirstToken, re.IGNORECASE):
        return getDocumentsDir()
    if re.match(r"^[a-z0-9-]+(\.[a-z0-9-]+)+$", sFirstToken, re.IGNORECASE) and "/" not in sFirstToken and "\\" not in sFirstToken:
        # Bare-domain heuristic: contains a dot, no slashes. Treat as url.
        return getDocumentsDir()
    # Try as a path. If it has wildcards, strip the basename and inspect
    # the parent directory.
    sCandidate = sFirstToken
    try:
        if any(c in sCandidate for c in "*?["):
            sCandidate = os.path.dirname(sCandidate)
        if sCandidate and os.path.isdir(sCandidate):
            return os.path.abspath(sCandidate)
        if sCandidate and os.path.isfile(sCandidate):
            return os.path.dirname(os.path.abspath(sCandidate))
        sParent = os.path.dirname(sCandidate) if sCandidate else ""
        if sParent and os.path.isdir(sParent):
            return os.path.abspath(sParent)
    except Exception:
        pass
    return getDocumentsDir()


def getDocumentsDir():
    """Return the user's Documents folder path as a string.

    On Windows, this resolves to the per-user Documents folder via
    Environment.SpecialFolder.MyDocuments (the SHGetKnownFolderPath
    KNOWNFOLDERID_Documents). On other platforms (developer machines
    only, urlCheck is Windows-only at runtime), falls back to the
    user's home directory.

    Always returns a non-empty string. If the Documents folder cannot be
    resolved for any reason, returns the home directory.
    """
    sPath = ""
    try:
        # The standard pythonic way: os.path.expanduser handles Windows
        # via the USERPROFILE environment variable, then we append
        # "Documents". But the more correct way on Windows is to query
        # the shell folder, since the user may have redirected
        # Documents to OneDrive or another non-default location.
        # We try SHGetFolderPathW first; if that fails, fall back to
        # %USERPROFILE%\Documents.
        if sys.platform == "win32":
            try:
                bufPath = ctypes.create_unicode_buffer(260)
                # CSIDL_PERSONAL = 0x0005 (Documents). Flags = 0 (current
                # location, not default). HResult 0 (S_OK) on success.
                iCsidlPersonal = 0x0005
                iHr = ctypes.windll.shell32.SHGetFolderPathW(
                    None, iCsidlPersonal, None, 0, bufPath)
                if iHr == 0 and bufPath.value and os.path.isdir(bufPath.value):
                    return bufPath.value
            except Exception:
                pass
        # Cross-platform fallback.
        sPath = os.path.join(os.path.expanduser("~"), "Documents")
        if os.path.isdir(sPath):
            return sPath
        return os.path.expanduser("~")
    except Exception:
        return os.path.expanduser("~")


def getStandardsRefs(dRule):
    lRefs = []
    sTag = ""

    for sTag in dRule.get("tags", []):
        if not any(str(sTag).startswith(sPrefix) for sPrefix in aRelevantTagPrefixes): continue
        if str(sTag).startswith("wcag"): continue
        lRefs.append(str(sTag))
    return lRefs


def getSummaryData(dResults, lRows):
    dSummary = {}

    dSummary = {
        "impactRows": getImpactRows(lRows),
        "outcomeRows": buildOutcomeSummaryRows(dResults),
        "rulesByFrequency": getRuleFrequencyRows(lRows),
        "wcagRows": getWcagFrequencyRows(lRows),
    }
    return dSummary


def getUrlsFromFile(sInput):
    """Read a plain text file and return a list of non-blank urls/paths.

    Each line is stripped; lines that are blank or start with # are ignored.
    The file's contents are sniffed up-front: if the leading bytes look
    binary (NUL bytes or a high ratio of non-printable bytes), we raise a
    clear error rather than producing garbage urls. Lines that don't look
    url-like are also rejected with a precise line-number error message.

    Raises ValueError on parse failure with a message suitable for the
    user. Raises OSError (FileNotFoundError, PermissionError) if the file
    cannot be opened.
    """
    bLikelyBinary = False
    bytesHead = b""
    iBlankLines = 0
    iLineNo = 0
    iSampleLen = 4096
    lUrls = []
    path = None
    sLine = ""
    sNormalized = ""

    path = pathlib.Path(sInput).expanduser().resolve()

    # Sniff the first few KB to catch binary files. NUL bytes are a strong
    # signal; otherwise any chunk where >30% of the bytes are outside the
    # printable-ASCII / common-control range is treated as binary.
    try:
        with path.open("rb") as fileRaw:
            bytesHead = fileRaw.read(iSampleLen)
    except Exception as ex:
        raise OSError(f"Could not read {sInput}: {ex}")
    if b"\x00" in bytesHead:
        bLikelyBinary = True
    elif bytesHead:
        iPrintable = sum(1 for b in bytesHead
            if (32 <= b < 127) or b in (9, 10, 13))
        if iPrintable / len(bytesHead) < 0.7:
            bLikelyBinary = True
    if bLikelyBinary:
        raise ValueError(
            f"File does not appear to be plain text: {sInput}\n"
            "urlCheck expects a plain text file with one url per line. "
            "If this is a Word, Excel, PDF, or other binary file, export "
            "or save it as plain text first.")

    # Now read as text. Use UTF-8 with strict error handling to surface
    # encoding problems explicitly; UTF-8-with-BOM is also accepted.
    try:
        with path.open("r", encoding="utf-8-sig", errors="strict") as file:
            iLineNo = 0
            for sLine in file:
                iLineNo += 1
                sLine = sLine.strip()
                if not sLine:
                    iBlankLines += 1
                    continue
                if sLine.startswith("#"): continue
                # Cheap sanity check: each non-comment line should look
                # like a url, a domain, or a local file path.
                if not _looksLikeUrlOrPath(sLine):
                    raise ValueError(
                        f"{sInput}: line {iLineNo} does not look like a "
                        f"url, domain, or file path: {sLine!r}\n"
                        "Each non-blank, non-comment line should be one "
                        "url, one domain name, or one local HTML file path.")
                lUrls.append(sLine)
    except UnicodeDecodeError as ex:
        raise ValueError(
            f"File is not valid UTF-8 text: {sInput}\n"
            f"Decode error: {ex}\n"
            "Re-save the file as plain UTF-8 text and try again.")
    if not lUrls:
        raise ValueError(f"No urls found in file: {sInput}")
    return lUrls


def _looksLikeUrlOrPath(sLine):
    """Heuristic: does sLine look like a url, a bare domain, or a file path?

    Used by getUrlsFromFile to catch obviously-wrong lines (e.g. random
    English text from a misclassified document). Generous on purpose --
    we'd rather pass through one bad line and have Playwright reject it
    with a clear error than refuse a legitimate url because we got the
    pattern wrong. Rejects only lines that contain whitespace or that
    are pure ASCII text with no dot, slash, or colon.
    """
    # Lines with embedded whitespace are never valid (urls and paths
    # don't contain bare whitespace; if they did, the user should
    # quote/encode).
    if any(c.isspace() for c in sLine): return False
    # A url has a colon (after the scheme), a path has slashes or a drive
    # letter, a domain has at least one dot. Anything with at least one of
    # these structural characters is plausible.
    return ("://" in sLine) or ("." in sLine) or ("/" in sLine) or ("\\" in sLine)


def getWcagFrequencyRows(lRows):
    dCounts = {}
    dRow = {}
    lOutput = []
    sPart = ""
    sWcagRefs = ""

    for dRow in lRows:
        sWcagRefs = str(dRow.get("wcagRefs") or "")
        if not sWcagRefs: continue
        for sPart in [sPart.strip() for sPart in sWcagRefs.split("|") if sPart.strip()]:
            dCounts[sPart] = int(dCounts.get(sPart, 0)) + 1
    for sPart in sorted(dCounts.keys(), key=lambda sK: (-int(dCounts.get(sK, 0)), sK)): lOutput.append([sPart, dCounts[sPart]])
    return lOutput[:20]


def getWcagRef(sTag):
    """Return a numeric WCAG SC ref like '1.4.3' from an axe tag like 'wcag143',
    or empty string if the tag is not a specific SC reference."""
    sDigits = ""
    sRef = ""

    if not str(sTag).startswith("wcag"): return ""
    if re.fullmatch(r"wcag\d+[a-z]+", str(sTag)): return ""
    sDigits = re.sub(r"[^0-9]", "", str(sTag))
    if len(sDigits) < 3: return ""
    sRef = f"{sDigits[0]}.{sDigits[1]}.{sDigits[2]}"
    return sRef


def getWcagRefs(dRule):
    lRefs = []
    sRef = ""
    sRuleId = ""
    sTag = ""

    # Primary: derive refs from axe wcagXXX tags
    for sTag in dRule.get("tags", []):
        sRef = getWcagRef(str(sTag))
        if sRef and sRef not in lRefs: lRefs.append(sRef)
    # Supplementary: if rule is best-practice with no wcag tags, use advisory map
    if not lRefs and "best-practice" in dRule.get("tags", []):
        sRuleId = str(dRule.get("id") or "")
        for sRef in dBpWcagRefs.get(sRuleId, []):
            if sRef not in lRefs: lRefs.append(sRef)
    return lRefs


def getWcagScInfo(sRef):
    """Return (shortName, level, principle) for a WCAG SC number, or None if unknown."""
    return dWcagSc.get(sRef)


def isUrlListFile(sInput):
    """Return True if sInput is a path to an existing file (any extension).

    The user is responsible for supplying a plain text file. If a file is
    given that turns out to be binary or non-url-like, getUrlsFromFile will
    fail with a clear error.
    """
    path = None
    if not sInput: return False
    try:
        path = pathlib.Path(sInput).expanduser()
    except Exception:
        return False
    try:
        if not path.is_file(): return False
    except Exception:
        return False
    return True


def classifyInput(sInput):
    """Classify the source input string.

    Friendlier parsing rules:
      1. Strip leading/trailing whitespace.
      2. Strip a single layer of surrounding double quotes (so the user
         can paste a Windows path with spaces verbatim, with or without
         the quotes File Explorer adds when copying a path).
      3. Test the entire string as a file path FIRST. This means a path
         containing internal spaces (e.g. C:\\My Documents\\urls.txt) is
         recognized as a file even without quotes.
      4. If it is not a file, fall back to space-tokenization, treating
         the result as one or more urls or domains.

    The user only needs to use quotes when supplying multiple specs and
    at least one contains a space. For a single path, no quotes are
    needed -- the entire trimmed input is tested as one path.

    Returns one of:
      ('listfile', sPath)  -- sInput is a path to an existing file (any
                              extension), to be parsed as plain text with
                              one url or local file path per line.
      ('urls', lUrls)      -- sInput is one or more space-separated urls/
                              domains. lUrls is the list of tokens, with
                              quotes around individual tokens stripped.
      ('error', sReason)   -- sInput is invalid (currently only the empty
                              case; bad file contents are surfaced later
                              by getUrlsFromFile with a precise message).
    """
    sStripped = ""
    sUnquoted = ""
    sStripped = (sInput or "").strip()
    if not sStripped:
        return ("error", "No input provided.")
    # Strip a single pair of surrounding double quotes.
    sUnquoted = sStripped
    if len(sUnquoted) >= 2 and sUnquoted[0] == '"' and sUnquoted[-1] == '"':
        sUnquoted = sUnquoted[1:-1].strip()
    # Test the entire (unquoted) input as a file path first. This catches
    # paths with embedded spaces without requiring the user to add quotes.
    if isUrlListFile(sUnquoted):
        return ("listfile", sUnquoted)
    # Not a single file. Fall back to space-tokenization, but honor
    # quoted segments so the user can supply a mix of paths-with-spaces
    # and individual urls.
    lTokens = parseSpaceSeparated(sStripped)
    return ("urls", lTokens)


def parseSpaceSeparated(sText):
    """Tokenize sText on whitespace, treating "..." as a single token.

    Used when the user has supplied multiple specs in a single string and
    at least one of them needs to be quoted because it contains spaces.

    Examples:
      parseSpaceSeparated('a b c')                    -> ['a', 'b', 'c']
      parseSpaceSeparated('"a b" c')                  -> ['a b', 'c']
      parseSpaceSeparated('"C:\\My Docs\\file.txt"')  -> ['C:\\My Docs\\file.txt']
      parseSpaceSeparated('a "b c" d')                -> ['a', 'b c', 'd']

    Backslashes are NOT treated as escapes (Windows path-friendly).
    Mismatched quotes: any leading-quote without a matching trailer
    closes at end of string.
    """
    lTokens = []
    sCurrent = ""
    bInQuote = False
    sChar = ""

    for sChar in sText:
        if bInQuote:
            if sChar == '"':
                bInQuote = False
                continue
            sCurrent += sChar
        else:
            if sChar == '"':
                bInQuote = True
                continue
            if sChar.isspace():
                if sCurrent:
                    lTokens.append(sCurrent)
                    sCurrent = ""
                continue
            sCurrent += sChar
    if sCurrent: lTokens.append(sCurrent)
    return lTokens


def parseArguments():
    argParser = None

    argParser = argparse.ArgumentParser(
        prog=sProgramName,
        description=(
            "Check one or more web pages for accessibility problems and save "
            "a set of output files in a folder named after each page title. "
            "Pass urls as separate arguments, or pass the path to a single "
            "plain text file that lists urls, domains, or local file paths -- "
            "one per line. The list file may have any extension; urlCheck "
            "verifies it is plain text by inspecting its contents."
        ),
        epilog=(
            f"Single url:    {sProgramName} https://example.com\n"
            f"Domain only:   {sProgramName} microsoft.com\n"
            f"Several urls:  {sProgramName} https://a.com https://b.com https://c.com\n"
            f"url list file: {sProgramName} urls.txt\n"
            f"GUI dialog:    {sProgramName} -g\n"
            f"\n"
            f"Output files (in a folder named after each page title):\n"
            f"  report.htm   Accessibility report with headings and links\n"
            f"  report.xlsx  Excel workbook with summary and full results\n"
            f"  results.json Full raw scan data including all metadata\n"
            f"  page.yaml    ARIA accessibility tree of the page\n"
            f"  page.htm     Saved page source with styles inlined\n"
            f"  page.png     Full-page screenshot\n"
            f"  violations\\  Per-violation element screenshots, linked\n"
            f"               from the Image column of report.xlsx"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    argParser.add_argument(
        "sSource",
        nargs="*",
        help=(
            "One or more urls (or domain names) separated by spaces, or the "
            "path to a single plain text file that lists urls, domains, or "
            "local file paths one per line. The list file may have any "
            "extension. Files referenced inside the list are loaded as HTML "
            "in the browser regardless of extension. Blank lines and lines "
            "starting with # are ignored."
        ),
    )
    argParser.add_argument("-v", "--version", action="version", version=f"%(prog)s {sProgramVersion}")
    argParser.add_argument("-g", "--gui-mode", dest="bGuiMode", action="store_true",
        help="Show the parameter dialog. GUI mode is also entered automatically when urlCheck is launched without arguments from a GUI shell (File Explorer, Start menu, desktop hotkey).")
    argParser.add_argument("-o", "--output-folder", dest="sOutputDir", default="",
        help="Parent folder under which the per-scan output folder is created. Defaults to the current working folder. Created if it does not exist. The per-scan folder is always uniquely named based on the page title.")
    argParser.add_argument("--view-output", dest="bViewOutput", action="store_true",
        help="After all scans complete, open the parent output folder (the -o folder, or the current working folder) in File Explorer.")
    argParser.add_argument("-u", "--use-configuration", dest="bUseConfig", action="store_true",
        help="Load saved settings from %%LOCALAPPDATA%%\\urlCheck\\urlCheck.ini at startup, and write them back on OK in GUI mode. Without this flag urlCheck leaves no filesystem footprint of its own.")
    argParser.add_argument("-l", "--log", dest="bLog", action="store_true",
        help="Write detailed diagnostics to urlCheck.log in the current working folder (UTF-8 with BOM). Appends across runs by default; combine with -f / --force to replace the prior log instead.")
    argParser.add_argument("-f", "--force", dest="bForce", action="store_true",
        help="Reuse an existing per-page output folder by emptying its contents and writing a fresh set of files. Without this flag urlCheck skips a url whose per-page output folder already exists, so previous scans are preserved.")
    argParser.add_argument("-i", "--invisible", dest="bInvisible", action="store_true",
        help="Run Microsoft Edge invisibly (the headless browser mode): no visible browser window during the scan.")
    argParser.add_argument("-a", "--authenticate", dest="bAuthenticate", action="store_true",
        help="When a url's domain is encountered for the first time in this run, pause after the page loads and prompt the user to authenticate (sign in, accept cookies, dismiss popups, etc.) and press Enter (or click OK in GUI mode) to resume. Within the same run, subsequent urls on the same domain reuse the established session without prompting. Without --main-profile, urlCheck disconnects its automation channel from Edge while the user authenticates and reconnects after, which improves the chance of success against sites that block automated browsers (e.g. WhatsApp Web). With --main-profile, the disconnect pattern is not used (browser security forbids it against your real profile). Forces visible browser mode (overrides --invisible).")
    argParser.add_argument("-m", "--main-profile", dest="bMainProfile", action="store_true",
        help="Launch Edge with your real (default) Edge user profile so saved logins, cookies, and session state are available. Without -m, urlCheck launches Edge with a fresh ephemeral profile so the scan is anonymous and your real profile is not exposed to the scanned site. Requires that no Microsoft Edge process is already running, since Edge cannot share a profile across two processes; urlCheck checks at startup and exits gracefully with a message if Edge is running.")
    return argParser.parse_args()


# Module-level state for the --authenticate feature. Tracks the set of
# url hostnames already seen in this run so the user is prompted at
# most once per registrable domain. Cleared implicitly when the
# program exits.
setSeenDomains = set()

# Two-letter labels that, when followed by a 2-letter country TLD,
# form a public 2-level suffix (so the registrable domain needs three
# labels rather than two). For example "co.uk", "com.au", "ac.uk",
# "co.jp" -- in those cases bbc.co.uk is the registrable domain, not
# co.uk. This heuristic covers the common cases without depending on
# the full Mozilla Public Suffix List. If the user encounters a real
# website whose registrable domain isn't extracted correctly, the
# only consequence is being prompted twice for the same site -- not
# data loss.
sSecondLevelPublicSuffixes = {
    "co", "com", "org", "net", "gov", "edu", "ac", "mil",
    "or", "ne", "info", "biz", "int", "nom", "go",
}


def getRegistrableDomain(sHost):
    """Reduce a hostname to its registrable domain.

    Examples:
      www.facebook.com         -> facebook.com
      m.facebook.com           -> facebook.com
      Accounts.Google.COM      -> google.com   (caller lowercases)
      bbc.co.uk                -> bbc.co.uk
      news.bbc.co.uk           -> bbc.co.uk
      example.com.au           -> example.com.au
      192.168.1.1              -> 192.168.1.1   (IPv4 stays as-is)
      [::1]                    -> ::1            (IPv6 stays as-is)
      localhost                -> localhost     (single-label)

    The result is always lowercased.
    """
    if not sHost: return ""
    sHost = sHost.lower().strip().strip(".")
    # IPv6 / IPv4: the urlsplit-derived hostname doesn't have brackets,
    # but it can have colons (IPv6) or be all-digits-and-dots (IPv4).
    # Treat both as a single token (no dotted-domain decomposition).
    if ":" in sHost: return sHost
    if all((c.isdigit() or c == ".") for c in sHost): return sHost
    lParts = sHost.split(".")
    if len(lParts) <= 2: return sHost
    sLast = lParts[-1]
    sSecond = lParts[-2]
    # 2-letter country TLD with a known 2-level suffix label:
    #   bbc.co.uk -> bbc.co.uk (3 labels)
    #   news.bbc.co.uk -> bbc.co.uk (last 3)
    if len(sLast) == 2 and sSecond in sSecondLevelPublicSuffixes:
        return ".".join(lParts[-3:])
    # Otherwise registrable domain is the last 2 labels.
    return ".".join(lParts[-2:])


def getDomainForAuth(sUrl):
    """Return the registrable domain of sUrl as a stable lowercase
    key, or None if there is no meaningful host to authenticate
    against (file://, about:, blank, etc.). The caller compares
    results case-insensitively without further work since we
    lowercase here.

    Subdomains collapse to the registrable domain by default, since
    most sites authenticate per registrable domain (cookies set on
    .example.com cover m.example.com, www.example.com, accounts.
    example.com, etc.).
    """
    if not sUrl: return None
    try:
        splitResult = urllib.parse.urlsplit(sUrl)
    except Exception:
        return None
    sScheme = (splitResult.scheme or "").lower()
    if sScheme in ("file", "about", "data", "javascript", ""): return None
    sHost = (splitResult.hostname or "").lower()
    if not sHost: return None
    return getRegistrableDomain(sHost)


def pauseForAuthenticationIfNeeded(page, sUrl, bAuthenticate, bGuiMode):
    """If --authenticate is on and sUrl's registrable domain has not
    been seen this run, prompt the user to interact with the visible
    browser window (sign in, dismiss cookie banners, accept popups,
    etc.) and confirm when ready to continue. The browser was
    launched non-headless because -a is on, so the user can interact
    directly with its window during the pause.

    Returns the page that should be used for the rest of the scan.
    Usually this is the original `page`, but auth flows commonly
    open popups, redirect through OAuth/SSO providers, or land the
    user on a different page than the originally-requested url. In
    those cases the original page reference is stale; we swap to
    the most-recently-active page in the context and return that.

    Two prompt mechanisms:
      - CLI mode: write the prompt to stdout (preceded by a newline
        so it doesn't run into the URL printed inline by the per-url
        progress) and wait for Enter via sys.stdin.readline().
      - GUI mode: show a small MessageBox via System.Windows.Forms,
        owned by an invisible TopMost form so the dialog appears on
        top of the Edge browser window rather than behind it. Without
        the TopMost owner, JAWS/NVDA users can't find the dialog.

    After confirmation, refreshes the page state by waiting for
    domcontentloaded so any user-initiated navigation (sign-in
    redirect, OAuth callback, popup, etc.) is settled before the
    accessibility scan runs against the now-authenticated DOM.

    A no-op when:
      - bAuthenticate is False
      - sUrl has no host (file://, about:, etc.)
      - the registrable domain has already been seen this run
    """
    if not bAuthenticate: return page
    sDomain = getDomainForAuth(sUrl)
    if sDomain is None: return page
    if sDomain in setSeenDomains: return page
    setSeenDomains.add(sDomain)

    sPrompt = ("Authenticate credentials, if needed, to show the "
        "page. To resume automation, press Enter:")
    sPromptGui = ("Authenticate credentials, if needed, to show "
        "the page in the open browser window. When ready, click OK "
        "to resume automation.")

    logger.info(f"Prompting for authentication on registrable domain "
        f"{sDomain} (url={sUrl})")
    try:
        if bGuiMode:
            # GUI mode: pop a MessageBox owned by a hidden TopMost
            # form so the dialog appears on top of the Edge window
            # rather than behind it. Without an owner, MessageBox
            # belongs to whatever process/window happens to have
            # focus -- and Edge typically does, after the user
            # clicked into it to sign in.
            import clr  # type: ignore
            clr.AddReference("System.Windows.Forms")
            clr.AddReference("System.Drawing")
            from System.Windows.Forms import (
                Form, FormBorderStyle, FormStartPosition,
                MessageBox, MessageBoxButtons, MessageBoxIcon,
                MessageBoxDefaultButton)
            from System.Drawing import Size
            owner = Form()
            owner.TopMost = True
            owner.ShowInTaskbar = False
            owner.FormBorderStyle = getattr(FormBorderStyle, "None")
            owner.StartPosition = FormStartPosition.Manual
            owner.Size = Size(1, 1)  # 1x1 pixel; effectively invisible
            try:
                owner.Show()
                MessageBox.Show(
                    owner,
                    sPromptGui,
                    f"{sProgramName} - Authenticate",
                    MessageBoxButtons.OK,
                    MessageBoxIcon.Information,
                    MessageBoxDefaultButton.Button1)
            finally:
                try: owner.Close(); owner.Dispose()
                except Exception: pass
        else:
            # CLI mode: leading "\n" terminates the URL printed inline
            # by the per-url progress, so the prompt appears on its
            # own line. End with a space (no newline) so the user's
            # Enter naturally terminates the prompt line.
            sys.stdout.write("\n" + sPrompt + " ")
            sys.stdout.flush()
            sys.stdin.readline()
        logger.info(f"User confirmed; resuming scan of {sUrl}")
    except Exception as ex:
        logger.warn(f"Authentication prompt failed: {ex}; continuing")

    # After the user has had a chance to interact, the original page
    # may be stale: an OAuth/SSO flow may have opened popups,
    # redirected through other origins, or left the user on a
    # different page than the originally-requested url. Swap to the
    # most-recently-active page in the context if it's different
    # from the original. The original page (if it's still open and
    # different) is closed so it doesn't hold onto resources.
    pageActive = page
    try:
        lPages = list(page.context.pages)
        if lPages:
            pageCandidate = lPages[-1]
            try:
                # is_closed exists in modern Playwright; fall through if not.
                bClosed = bool(pageCandidate.is_closed())
            except Exception:
                bClosed = False
            if (not bClosed) and (pageCandidate is not page):
                logger.info(f"Auth flow created a new page; switching from "
                    f"{page.url!r} to {pageCandidate.url!r}")
                try:
                    if not page.is_closed(): page.close()
                except Exception:
                    pass
                pageActive = pageCandidate
    except Exception as ex:
        logger.warn(f"Could not re-acquire active page after auth: {ex}; "
            f"using original page reference")

    # Re-settle whichever page we now have in case the user is
    # mid-navigation or the page is still loading. domcontentloaded
    # is sufficient -- networkidle would hang on chatty pages.
    try:
        pageActive.wait_for_load_state("domcontentloaded", timeout=15000)
    except Exception:
        pass

    return pageActive


def pauseForAuthenticationWithDisconnect(playwrightCtx, sWsEndpoint,
        browser, context, page, sUrl, bAuthenticate, bGuiMode):
    """Disconnect/reconnect variant of pauseForAuthenticationIfNeeded
    used when --authenticate (-a) is set without --main-profile (-m).
    Differs from the default-profile variant in two ways:

    1. Before prompting the user, calls browser.close() to sever
       Playwright's CDP connection to the launched Edge process.
       The Edge process keeps running (it's a separate subprocess
       owned by main(), not by this client connection), so the user
       can interact with the visible window without any automation
       channel attached -- which lets sites that detect CDP-level
       fingerprinting (WhatsApp Web, some banking sites) complete
       navigations they would otherwise refuse.

    2. After the user confirms, reconnects via
       chromium.connect_over_cdp() to the same CDP HTTP endpoint
       (parameter name kept as sWsEndpoint for legacy reasons),
       finds the existing context and the page the user was last
       on, re-applies the navigator.webdriver init script to the
       new context binding, and returns the new (browser, context,
       page) triple. The caller MUST use the returned triple; the
       input browser/context/page references become stale during
       the disconnect.

    Returns (browser, context, page) -- a tuple of fresh references
    suitable for the rest of the scan. On the no-prompt-needed path
    (domain already authenticated this run), returns the inputs
    unchanged.
    """
    if not bAuthenticate: return (browser, context, page)
    sDomain = getDomainForAuth(sUrl)
    if sDomain is None: return (browser, context, page)
    if sDomain in setSeenDomains: return (browser, context, page)
    setSeenDomains.add(sDomain)

    sPrompt = ("Authenticate credentials, if needed, to show the "
        "page. To resume automation, press Enter:")
    sPromptGui = ("Authenticate credentials, if needed, to show "
        "the page in the open browser window. When ready, click OK "
        "to resume automation.")

    logger.info(f"Disconnecting Playwright before authentication "
        f"prompt on registrable domain {sDomain} (url={sUrl})")
    # Sever the Playwright client connection. The Edge process keeps
    # running -- we'll reconnect after the user confirms.
    try:
        browser.close()
    except Exception as ex:
        logger.warn(f"browser.close() during auth disconnect raised: "
            f"{ex}; continuing anyway")

    try:
        if bGuiMode:
            import clr  # type: ignore
            clr.AddReference("System.Windows.Forms")
            clr.AddReference("System.Drawing")
            from System.Windows.Forms import (
                Form, FormBorderStyle, FormStartPosition,
                MessageBox, MessageBoxButtons, MessageBoxIcon,
                MessageBoxDefaultButton)
            from System.Drawing import Size
            owner = Form()
            owner.TopMost = True
            owner.ShowInTaskbar = False
            owner.FormBorderStyle = getattr(FormBorderStyle, "None")
            owner.StartPosition = FormStartPosition.Manual
            owner.Size = Size(1, 1)
            try:
                owner.Show()
                MessageBox.Show(
                    owner,
                    sPromptGui,
                    f"{sProgramName} - Authenticate",
                    MessageBoxButtons.OK,
                    MessageBoxIcon.Information,
                    MessageBoxDefaultButton.Button1)
            finally:
                try: owner.Close(); owner.Dispose()
                except Exception: pass
        else:
            sys.stdout.write("\n" + sPrompt + " ")
            sys.stdout.flush()
            sys.stdin.readline()
        logger.info(f"User confirmed; reconnecting to Edge at {sWsEndpoint}")
    except Exception as ex:
        logger.warn(f"Authentication prompt failed: {ex}; continuing")

    # Reconnect. The Edge process kept running so all open pages,
    # contexts, and cookies (in our temp --user-data-dir) persist.
    try:
        browserNew = playwrightCtx.chromium.connect_over_cdp(sWsEndpoint)
    except Exception as ex:
        logger.warn(f"Reconnect to {sWsEndpoint} failed: {ex}")
        # Best-effort fallback: return stale references and let the
        # caller surface a per-url error rather than crash the run.
        return (browser, context, page)

    # Get the existing context (the one the user was interacting with
    # post-auth). With connect_over_cdp, the connected browser
    # exposes any contexts in the running Edge process. We expect
    # exactly one context that we created at run start.
    contextNew = None
    try:
        lContexts = list(browserNew.contexts)
        if lContexts: contextNew = lContexts[0]
    except Exception as ex:
        logger.warn(f"Could not enumerate contexts after reconnect: {ex}")
    if contextNew is None:
        try:
            contextNew = browserNew.new_context(
                bypass_csp=True,
                ignore_https_errors=bDefaultIgnoreHttpsErrors,
                user_agent=sUserAgent,
                viewport={"width": iDefaultViewportWidth,
                    "height": iDefaultViewportHeight})
        except Exception as ex:
            logger.warn(f"Could not create new context after reconnect: {ex}")
            return (browser, context, page)

    # Re-apply the navigator.webdriver init script. add_init_script
    # is idempotent in effect -- adding the same script twice means
    # it runs twice on each page load, which is harmless because the
    # second defineProperty() call overwrites with the same value.
    applyWebdriverOverride(contextNew)

    # Find the page the user was last on. The auth flow may have
    # navigated, opened popups, or redirected through OAuth/SSO; the
    # last page in the context is most likely the one with the
    # authenticated post-auth content.
    pageNew = None
    try:
        lPages = list(contextNew.pages)
        if lPages: pageNew = lPages[-1]
    except Exception as ex:
        logger.warn(f"Could not enumerate pages after reconnect: {ex}")
    if pageNew is None:
        try:
            pageNew = contextNew.new_page()
            pageNew.goto(sUrl, timeout=iDefaultNavTimeoutMs,
                wait_until="domcontentloaded")
        except Exception as ex:
            logger.warn(f"Could not create/navigate page after reconnect: {ex}")
            return (browserNew, contextNew, page)

    try:
        pageNew.wait_for_load_state("domcontentloaded", timeout=15000)
    except Exception:
        pass

    return (browserNew, contextNew, pageNew)


def scanUrl(sInput, sNormalizedUrl, browser, context, pathBaseDir, sAxeContent="", bForce=False, bAuthenticate=False, bGuiMode=False, bMainProfile=False, playwrightCtx=None, sWsEndpoint=None, lConnHolder=None):
    """Run a single-url scan.

    Returns:
      - the output directory path string on a successful scan
      - the sentinel string "skipped" if the per-page folder already
        exists and bForce is False (caller decided not to overwrite
        previous results)

    Raises on real errors (navigation failure, axe failure, etc.).
    """
    dMetadata = {}
    dResults = {}
    lArgs = []
    lRows = []
    page = None
    pathOutputDir = None
    sAxeSource = ""
    sPageTitle = ""
    sResultsJson = ""
    sSnapshot = ""

    try:
        page = context.new_page()
        logger.info(f"Navigating to {sNormalizedUrl}")
        # When --authenticate is on, wait only for "domcontentloaded"
        # initially -- enough for the page to be interactive so the
        # user can authenticate, but without paying the long "load"
        # event delay before they see the auth prompt. The full
        # "load" + networkidle settling happens AFTER the user
        # confirms, when the page may have navigated to a post-auth
        # destination that needs its own settling time.
        #
        # When --authenticate is off, keep the original "load"
        # behavior so the rest of the flow is unchanged for non-auth
        # runs.
        sInitialWaitUntil = "domcontentloaded" if bAuthenticate else "load"
        page.goto(sNormalizedUrl, timeout=iDefaultNavTimeoutMs, wait_until=sInitialWaitUntil)

        # If --authenticate is on and this is the first url on this
        # registrable domain, pause now so the user can sign in /
        # accept cookies / dismiss popups / complete 2FA. Returns the
        # page that should be used for the rest of the scan --
        # normally the same `page`, but if the auth flow opened a
        # popup or redirected to another page, the helper swaps in
        # the most-recently-active page so the scan captures the
        # post-auth content.
        if bAuthenticate and not bMainProfile:
            # Temp-profile + disconnect path: sever the CDP
            # connection before prompting so the page can't detect
            # active automation, prompt, reconnect via
            # connect_over_cdp after the user confirms. Returns a
            # fresh (browser, context, page) triple; the input
            # references become stale across the disconnect.
            # Update lConnHolder so the caller can pick up the fresh
            # refs for the next url.
            (browser, context, page) = pauseForAuthenticationWithDisconnect(
                playwrightCtx, sWsEndpoint,
                browser, context, page, sNormalizedUrl,
                bAuthenticate, bGuiMode)
            if lConnHolder is not None:
                lConnHolder[0] = browser
                lConnHolder[1] = context
        else:
            # Main-profile-with-auth, or non-auth runs: prompt (or
            # do nothing) without touching the connection. The
            # disconnect pattern is not available against the real
            # Edge profile because Chrome 136+ refuses
            # --remote-debugging-port against the default user-data
            # directory.
            page = pauseForAuthenticationIfNeeded(page, sNormalizedUrl,
                bAuthenticate, bGuiMode)

        # After the user confirms (or immediately, for non-auth
        # runs), settle the page. For auth runs we add an extra
        # wait_for_load_state("load") because we only waited for
        # domcontentloaded above, AND because the user may have
        # triggered a post-auth navigation that we need to follow.
        # For non-auth runs the goto already waited for "load", so
        # this call returns quickly.
        try:
            page.wait_for_load_state("load", timeout=iDefaultNavTimeoutMs)
        except Exception:
            pass
        # Then wait briefly for network to settle. A short
        # networkidle timeout lets SPA content finish rendering on
        # most sites; sites with persistent connections (e.g.
        # WebSocket-heavy SPAs) simply time out and continue after
        # iNetworkIdleTimeoutMs ms.
        try:
            page.wait_for_load_state("networkidle", timeout=iNetworkIdleTimeoutMs)
        except Exception:
            pass
        # For auth runs, add an extra settling delay because post-
        # auth UIs (Facebook feed, WhatsApp main panel, Slack
        # workspace) often render their main content asynchronously
        # via React/Vue/etc. after networkidle is reported. Without
        # this extra time the accessibility scan can run against a
        # half-rendered DOM and produce a misleadingly empty report.
        if bAuthenticate:
            page.wait_for_timeout(iAuthPostConfirmSettleDelayMs)
        page.wait_for_timeout(iDefaultPostLoadDelayMs)
        # Scroll to bottom and back to trigger lazy-loaded content, then wait briefly
        page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        page.wait_for_timeout(500)
        page.evaluate("window.scrollTo(0, 0)")
        page.wait_for_timeout(500)
        sPageTitle = str(page.title() or sFallbackTitle)
        # Decide the output folder. If chooseOutputDir returns None,
        # the per-page folder already exists and --force is not set,
        # so we skip this url. Skip happens BEFORE the expensive
        # axe-core run, screenshot, snapshot, and report-writing.
        pathOutputDir = chooseOutputDir(pathBaseDir, sPageTitle, bForce=bForce)
        if pathOutputDir is None:
            sExistingDir = str(pathBaseDir / getSafeTitle(sPageTitle))
            logger.info(f"Skipped (output folder exists, no --force): "
                f"{sNormalizedUrl} -> {sExistingDir}")
            return "skipped"
        logger.info(f"Output folder: {pathOutputDir}")
        logger.info("Running axe-core")
        sAxeSource = getAxeScript(page, sAxeContent)
        ensureSuccess(bool(page.evaluate("() => Boolean(window.axe && window.axe.run)")), "axe-core did not load into the page.")
        sResultsJson = page.evaluate("async (opts) => JSON.stringify(await window.axe.run(document, opts))", aAxeRunOptions)
        dResults = json.loads(sResultsJson)
        # browser.version is None when using launch_persistent_context
        # (Playwright doesn't expose a Browser object in that case).
        # Fall back to extracting Edge's version from the page's
        # navigator.userAgent string, which always contains
        # "Edg/<version>" on a Microsoft Edge browser.
        sBrowserVersion = ""
        try:
            if browser is not None:
                sBrowserVersion = str(browser.version or "")
        except Exception:
            pass
        if not sBrowserVersion:
            try:
                sUa = str(page.evaluate("navigator.userAgent") or "")
                match = re.search(r"Edg/(\S+)", sUa)
                if match: sBrowserVersion = match.group(1)
            except Exception:
                pass

        dMetadata = {
            "axeSource": sAxeSource,
            "browserChannel": sBrowserChannel,
            "browserVersion": sBrowserVersion,
            "inputValue": sInput,
            "navTimeoutMs": iDefaultNavTimeoutMs,
            "normalizedUrl": sNormalizedUrl,
            "osVersion": getOsVersion(),
            "pageTitle": sPageTitle,
            "pageUrl": str(page.url or sNormalizedUrl),
            "postLoadDelayMs": iDefaultPostLoadDelayMs,
            "programName": sProgramName,
            "programVersion": sProgramVersion,
            "scanTimestampUtc": datetime.datetime.now(datetime.timezone.utc).replace(microsecond=0).isoformat(),
            "userAgent": str(page.evaluate("() => navigator.userAgent")),
            "viewportHeight": iDefaultViewportHeight,
            "viewportWidth": iDefaultViewportWidth,
        }
        # report.csv was previously written here. As of v1.11.0 it
        # has been removed: the same row-per-violation data is in
        # the Results sheet of report.xlsx with proper formatting,
        # so the CSV duplicated content with inferior presentation.
        # The lRows list is still computed for use by the HTML
        # report and the workbook's Results sheet.
        lRows = buildCsvRows(dResults, dMetadata)
        writeTextFile(pathlib.Path(pathOutputDir, sJsonName), json.dumps({"metadata": dMetadata, "results": dResults}, indent=2, ensure_ascii=False))
        logger.info("Capturing page snapshot and screenshot")
        sSnapshot = getPageSnapshot(page, sNormalizedUrl)
        writeTextFile(pathlib.Path(pathOutputDir, sSourceName), sSnapshot)
        try:
            page.screenshot(path=str(pathlib.Path(pathOutputDir, sScreenshotName)), full_page=bDefaultFullPage, timeout=30000)
        except Exception:
            try:
                page.screenshot(path=str(pathlib.Path(pathOutputDir, sScreenshotName)), full_page=False, timeout=15000)
            except Exception:
                pass
        try:
            locatorBody = page.locator("body")
            sYaml = locatorBody.aria_snapshot()
            writeTextFile(pathlib.Path(pathOutputDir, sAccessibilityYamlName), sYaml)
        except Exception:
            pass
        # Compute the per-page accessibility failure rate so it can
        # be displayed in report.htm and report.xlsx. The denominator
        # is the saved page.htm size, written above. Smaller is
        # better; the goal is to track this metric over time.
        iPageBytes = computePageBytes(pathOutputDir)
        iImpactNumer = computePageImpactNumerator(dResults)
        nPageRate = computeAccessibilityFailureRate(iImpactNumer, iPageBytes)
        # Capture per-violation element screenshots. Saves PNGs to a
        # 'violations/' subfolder and returns a dict mapping each
        # captured (ruleId, nodeIndex) to its relative path. Failures
        # are silent (per project policy) -- nodes whose selectors
        # don't resolve simply have no entry in the dict.
        try:
            dViolationImages = captureViolationImages(page, dResults, pathOutputDir)
        except Exception:
            dViolationImages = {}
        writeTextFile(pathlib.Path(pathOutputDir, sReportName), buildReportHtml(dResults, dMetadata, lRows, nPageRate, iPageBytes, iImpactNumer))
        writeReportWorkbook(pathlib.Path(pathOutputDir, sReportWorkbookName), dResults, dMetadata, lRows, nPageRate, iPageBytes, iImpactNumer, dViolationImages)
        # The CLI parent loop prints the URL inline before calling
        # scanUrl and appends a newline (success), ": skipped" (skip),
        # or ": <reason>" (failure) afterward. So scanUrl itself
        # doesn't print anything user-facing -- we'd duplicate that
        # CLI line. Detailed violation counts go to report.htm /
        # report.xlsx and to urlCheck.log via buildConsoleSummary.
        # The GUI mode never prints inline; the end-of-session
        # summary is built fresh by main() from lScanned / lFailed /
        # lSkippedExisting and shown in the MessageBox.
        logger.info(buildConsoleSummary(dResults, dMetadata, str(pathOutputDir)))
        return (str(pathOutputDir), sPageTitle)
    finally:
        try:
            if page is not None: page.close()
        except Exception:
            pass


def styleWorksheet(worksheet):
    cell = None
    iColumnIndex = 0
    lWidths = []
    sColumnLetter = ""

    worksheet.freeze_panes = "A2"
    for cell in worksheet["1:1"]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(fill_type="solid", fgColor="1F4E78")
        cell.alignment = Alignment(vertical="top", wrap_text=True)
    lWidths = [18, 24, 20, 16, 18, 18, 18, 18, 16, 16, 20, 28, 28, 28, 18, 14, 40, 60, 60, 18]
    for iColumnIndex, iWidth in enumerate(lWidths, start=1):
        sColumnLetter = get_column_letter(iColumnIndex)
        worksheet.column_dimensions[sColumnLetter].width = iWidth
    return worksheet.title


def writeReportWorkbook(pathWorkbook, dResults, dMetadata, lRows, nPageRate=0.0, iPageBytes=0, iImpactNumer=0, dViolationImages=None):
    cell = None
    iRow = 0
    iSheetIndex = 0
    lImpactRows = []
    lRuleRows = []
    lRow = []
    lWcagRows = []
    workbook = None
    worksheet = None
    if dViolationImages is None: dViolationImages = {}

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Metadata"
    worksheet.append(["Field", "Value"])
    for lRow in [
        ["Program", sProgramName],
        ["Version", sProgramVersion],
        ["Input", str(dMetadata.get("inputValue") or "")],
        ["Normalized url", str(dMetadata.get("normalizedUrl") or "")],
        ["Page url", str(dMetadata.get("pageUrl") or "")],
        ["Page title", str(dMetadata.get("pageTitle") or "")],
        ["Scan timestamp (UTC)", str(dMetadata.get("scanTimestampUtc") or "")],
        ["Browser channel", str(dMetadata.get("browserChannel") or "")],
        ["Browser version", str(dMetadata.get("browserVersion") or "")],
        ["User agent", str(dMetadata.get("userAgent") or "")],
        ["Viewport width", int(dMetadata.get("viewportWidth") or 0)],
        ["Viewport height", int(dMetadata.get("viewportHeight") or 0)],
        ["Navigation timeout (ms)", int(dMetadata.get("navTimeoutMs") or 0)],
        ["Post-load delay (ms)", int(dMetadata.get("postLoadDelayMs") or 0)],
        ["axe-core source", str(dMetadata.get("axeSource") or "")],
        ["axe result types", ", ".join(aAxeRunOptions.get("resultTypes", []))],
        ["Screenshot file", sScreenshotName],
        ["JSON file", sJsonName],
        ["ARIA tree file", sAccessibilityYamlName],
        ["Source snapshot file", sSourceName],
        ["HTML report file", sReportName],
        ["Workbook file", sReportWorkbookName],
    ]:
        worksheet.append(lRow)
    worksheet.column_dimensions["A"].width = 28
    worksheet.column_dimensions["B"].width = 90
    worksheet["A1"].font = Font(bold=True, color="FFFFFF")
    worksheet["B1"].font = Font(bold=True, color="FFFFFF")
    worksheet["A1"].fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    worksheet["B1"].fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    worksheet.freeze_panes = "A2"
    for cell in worksheet["B"]: cell.alignment = Alignment(vertical="top", wrap_text=True)

    worksheet = workbook.create_sheet("Summary")
    worksheet.append(["Section", "Name", "Value"])
    lImpactRows = getImpactRows(lRows)
    lRuleRows = getRuleFrequencyRows(lRows)
    lWcagRows = getWcagFrequencyRows(lRows)
    worksheet.append(["Overview", "Page title", str(dMetadata.get("pageTitle") or "")])
    worksheet.append(["Overview", "Page url", str(dMetadata.get("pageUrl") or "")])
    worksheet.append(["Overview", "Violations (rules)", len(dResults.get("violations", []))])
    worksheet.append(["Overview", "Failed instances", sum(len(dRule.get("nodes", [])) for dRule in dResults.get("violations", []))])
    worksheet.append(["Overview", "Needs Review (rules)", len(dResults.get("incomplete", []))])
    worksheet.append(["Overview", "Needs Review (instances)", sum(len(dRule.get("nodes", [])) for dRule in dResults.get("incomplete", []))])
    worksheet.append(["Overview", "Passes (rules)", len(dResults.get("passes", []))])
    worksheet.append(["Overview", "Inapplicable (rules)", len(dResults.get("inapplicable", []))])
    worksheet.append(["Overview", "Page bytes (page.htm)", int(iPageBytes)])
    worksheet.append(["Overview", "Impact-weighted instance count", int(iImpactNumer)])
    worksheet.append(["Overview", "Accessibility failure rate (%)", round(nPageRate, 1)])
    for lRow in lImpactRows: worksheet.append(["Impact (failed instances)", str(lRow[0]), int(lRow[1])])
    for lRow in lRuleRows: worksheet.append(["Top rules by instance count", str(lRow[0]), int(lRow[1])])
    for lRow in lWcagRows:
        sRef = str(lRow[0])
        tInfo = getWcagScInfo(sRef)
        sLabel = f"{sRef} — {tInfo[0]} (Level {tInfo[1]})" if tInfo else sRef
        worksheet.append(["Top WCAG criteria", sLabel, int(lRow[1])])
    worksheet.column_dimensions["A"].width = 28
    worksheet.column_dimensions["B"].width = 56
    worksheet.column_dimensions["C"].width = 16
    styleWorksheet(worksheet)

    worksheet = workbook.create_sheet("Results")
    worksheet.append(["Scan timestamp (UTC)", "Page title", "Page url", "Browser version", "axe-core source", "Outcome", "Rule ID", "Impact", "Description", "Help", "Help url", "Tags", "WCAG criteria", "Standards refs", "Instance count", "Instance index", "Path (CSS selector)", "Snippet (HTML)", "Failure summary", "Image"])
    iImageColIdx = 20  # 1-based: 20th column == "Image"
    for lRow in lRows:
        worksheet.append([lRow["scanTimestampUtc"], lRow["pageTitle"], lRow["pageUrl"], lRow["browserVersion"], lRow["axeSource"], lRow["outcome"], lRow["ruleId"], lRow["impact"], lRow["description"], lRow["help"], lRow["helpUrl"], lRow["tags"], lRow["wcagRefs"], lRow["standardsRefs"], lRow["ruleNodeCount"], lRow["ruleNodeIndex"], lRow["target"], lRow["html"], lRow["failureSummary"], ""])
        # If a screenshot was successfully captured for this
        # violation node, populate the Image cell as a hyperlink to
        # the relative file path. Clicking the cell in Excel opens
        # the PNG in the OS default image viewer. Relative paths
        # (no drive letter, no scheme) are resolved by Excel
        # against the workbook's actual location, so the link
        # survives moving or zipping the page folder.
        if lRow.get("outcome") == "violations":
            sRuleId = str(lRow.get("ruleId") or "")
            iNodeIdx = int(lRow.get("ruleNodeIndex") or 0)
            sRelPath = dViolationImages.get((sRuleId, iNodeIdx))
            if sRelPath:
                # The new row is the last one written; address it
                # by max_row.
                cellImage = worksheet.cell(row=worksheet.max_row, column=iImageColIdx)
                # Display text: just the basename so the user sees
                # 'image-001.png' rather than 'violations/image-001.png'.
                # The hyperlink target IS the relative path, which is
                # what Excel resolves on click.
                cellImage.value = sRelPath.rsplit("/", 1)[-1]
                cellImage.hyperlink = sRelPath
                cellImage.style = "Hyperlink"
    styleWorksheet(worksheet)

    worksheet = workbook.create_sheet("Glossary")
    worksheet.append(["Term", "Definition"])
    for lRow in aGlossaryRows: worksheet.append(lRow)
    worksheet.append(["", ""])
    worksheet.append(["Procedure steps", ""])
    for sProcedure in aProcedures: worksheet.append(["Procedure", sProcedure])
    worksheet.column_dimensions["A"].width = 24
    worksheet.column_dimensions["B"].width = 100
    worksheet["A1"].font = Font(bold=True, color="FFFFFF")
    worksheet["B1"].font = Font(bold=True, color="FFFFFF")
    worksheet["A1"].fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    worksheet["B1"].fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    worksheet.freeze_panes = "A2"
    for cell in worksheet["B"]: cell.alignment = Alignment(vertical="top", wrap_text=True)

    # Finalization: apply xlFormat/xlHeaders learnings to every
    # sheet for consistency with ACR.xlsx. Data-driven column widths
    # (capped at 50 chars, wrap_text on overflow), vertical-align top,
    # bold first row with white-on-blue fill, frozen header row, plus
    # a Title01-style named cell so screen readers announce both row
    # and column headers as the user navigates. AutoFilter is
    # deliberately NOT enabled: in JAWS, the surrounding "filter is
    # off" announcement interferes with reading column headers.
    for iSheetIndex in range(len(workbook.sheetnames)):
        worksheet = workbook[workbook.sheetnames[iSheetIndex]]
        for iRow in range(2, worksheet.max_row + 1): worksheet.row_dimensions[iRow].height = None
        acrBuilder.applyFormatting(worksheet, iHeaderRow=1, iDataStart=2)
        acrBuilder.addNamedRangeForCell(workbook, worksheet, worksheet.cell(row=1, column=1), "Title")
    workbook.save(pathWorkbook)
    return str(pathWorkbook)


# --- GUI launch detection (Windows) ---

def _getParentProcessName():
    """
    Returns the lowercased base filename of this process's parent (e.g.
    'cmd.exe', 'explorer.exe'), or '' if it can't be determined. Uses a
    Toolhelp32 snapshot, which is available on every Windows version this
    program supports.
    """
    if sys.platform != "win32": return ""
    try:
        # We use a couple of structures from kernel32. Define them via ctypes
        # to keep the dependency surface minimal (no pywin32, no psutil).
        import ctypes.wintypes as wt
        iSnapProcess = 0x00000002
        iMaxPath = 260

        class ProcessEntry32(ctypes.Structure):
            _fields_ = [
                ("dwSize", wt.DWORD),
                ("cntUsage", wt.DWORD),
                ("th32ProcessID", wt.DWORD),
                ("th32DefaultHeapID", ctypes.c_void_p),
                ("th32ModuleID", wt.DWORD),
                ("cntThreads", wt.DWORD),
                ("th32ParentProcessID", wt.DWORD),
                ("pcPriClassBase", ctypes.c_long),
                ("dwFlags", wt.DWORD),
                ("szExeFile", ctypes.c_char * iMaxPath),
            ]

        kernel32 = ctypes.windll.kernel32
        kernel32.CreateToolhelp32Snapshot.restype = wt.HANDLE
        kernel32.CreateToolhelp32Snapshot.argtypes = [wt.DWORD, wt.DWORD]
        kernel32.Process32First.argtypes = [wt.HANDLE, ctypes.POINTER(ProcessEntry32)]
        kernel32.Process32Next.argtypes = [wt.HANDLE, ctypes.POINTER(ProcessEntry32)]
        kernel32.CloseHandle.argtypes = [wt.HANDLE]

        iMyPid = int(os.getpid())
        iParentPid = 0
        sParentExe = ""

        hSnap = kernel32.CreateToolhelp32Snapshot(iSnapProcess, 0)
        if not hSnap or hSnap == wt.HANDLE(-1).value: return ""
        try:
            processEntry = ProcessEntry32()
            processEntry.dwSize = ctypes.sizeof(ProcessEntry32)
            # Pass 1: find our own entry to learn the parent PID.
            if not kernel32.Process32First(hSnap, ctypes.byref(processEntry)): return ""
            while True:
                if processEntry.th32ProcessID == iMyPid:
                    iParentPid = int(processEntry.th32ParentProcessID)
                    break
                if not kernel32.Process32Next(hSnap, ctypes.byref(processEntry)): break
            if iParentPid == 0: return ""

            # Pass 2: find the parent entry to learn its exe name. Re-walk
            # because Process32First/Next is unidirectional.
            oEntry2 = ProcessEntry32()
            oEntry2.dwSize = ctypes.sizeof(ProcessEntry32)
            if not kernel32.Process32First(hSnap, ctypes.byref(oEntry2)): return ""
            while True:
                if oEntry2.th32ProcessID == iParentPid:
                    sParentExe = oEntry2.szExeFile.decode("ascii", errors="replace")
                    break
                if not kernel32.Process32Next(hSnap, ctypes.byref(oEntry2)): break
        finally:
            kernel32.CloseHandle(hSnap)

        return os.path.basename(sParentExe).lower()
    except Exception:
        return ""


def isLaunchedFromGui():
    """
    Returns True when this process appears to have been launched by a GUI shell
    rather than from a command-line shell.

    Primary signal: GetConsoleProcessList. The same approach 2htm uses. A
    console-subsystem program inherits its parent shell's console (count >= 2)
    when launched from cmd / PowerShell / Windows Terminal / etc., but is the
    only process on a fresh console (count == 1) when Windows creates a new
    console for a double-clicked exe, a Start-menu shortcut, the Run dialog,
    or a desktop hotkey.

    Secondary signal: parent process name. Used only when the count is
    ambiguous: in particular, count == 0 (truly no console attached, which
    happens in some service / scheduled-task contexts) or when a third party
    has briefly attached a process to our console between launch and the
    moment we make the check (rare but possible -- accessibility tools,
    AV scanners, JAWS / NVDA event hooks, etc.). When count == 1 the answer
    is unambiguous regardless of parent.

    Both signals and the resulting decision are written to the log when -l
    is in effect, so the detection can be diagnosed from the log file.
    """
    bResult = False
    iCount = 0
    iCountErr = 0
    sParent = ""
    sReason = ""

    if sys.platform != "win32":
        sReason = "non-Windows platform; assume CLI"
        bResult = False
    else:
        try:
            aiBuf = (ctypes.c_uint * 16)()
            iCount = int(ctypes.windll.kernel32.GetConsoleProcessList(
                aiBuf, ctypes.c_uint(16)))
        except Exception as ex:
            iCountErr = 1
            iCount = -1
            sReason = f"GetConsoleProcessList failed: {ex}"
        if iCount == 1:
            bResult = True
            sReason = "console process count is 1 (fresh console -> GUI launch)"
        elif iCount >= 2:
            bResult = False
            sReason = f"console process count is {iCount} (sharing parent shell -> CLI)"
        else:
            # iCount == 0 (no console) or iCount < 0 (call failed). Fall
            # back to parent-process inspection.
            sParent = _getParentProcessName()
            sShells = ("cmd.exe", "powershell.exe", "pwsh.exe",
                "windowsterminal.exe", "openconsole.exe", "wt.exe",
                "conemu64.exe", "conemu.exe", "cmder.exe",
                "bash.exe", "wsl.exe", "git-bash.exe", "mintty.exe")
            if sParent in sShells:
                bResult = False
                sReason = f"console count={iCount}, parent={sParent} -> CLI"
            elif sParent:
                bResult = True
                sReason = f"console count={iCount}, parent={sParent} -> GUI"
            else:
                bResult = False
                sReason = f"console count={iCount}, parent unknown -> assume CLI"

    logger.info(f"isLaunchedFromGui: parent='{sParent}' "
        f"consoleCount={iCount} -> bIsGui={bResult} ({sReason})")
    return bResult


def hideOwnConsoleWindow():
    """
    Hides this process's console window. Only safe when isLaunchedFromGui()
    returned True; the console was created by Windows for us in that case and
    no parent shell is sharing it. Non-fatal on failure.
    """
    iSwHide = 0
    if sys.platform != "win32": return
    try:
        hwnd = ctypes.windll.kernel32.GetConsoleWindow()
        if hwnd: ctypes.windll.user32.ShowWindow(hwnd, iSwHide)
    except Exception:
        pass


def openFolderInExplorer(sPath):
    """Open the given folder in Windows Explorer. Non-fatal on failure."""
    try:
        if sys.platform == "win32":
            os.startfile(sPath)
        else:
            subprocess.Popen(["xdg-open", sPath])
    except Exception:
        pass


# --- Logger ---

class logger:
    """
    Tiny diagnostic logger written to urlCheck.log in CWD when -l / --log is
    given. UTF-8 with BOM so Notepad opens it correctly. By default the log
    is OPENED IN APPEND mode so accumulated history across runs is
    preserved -- helpful when diagnosing intermittent issues. Pass
    bReplace=True to start fresh; this is hooked to the -f / --force
    flag so users who want a clean log per run get it deliberately.
    Open is lazy and silent on failure; a logging error must never sink a
    scan.
    """
    fLog = None
    bEnabled = False
    sPath = None  # resolved log file path
    aBuffer = []  # pre-open diagnostic buffer (timestamp, level, msg)
    bBuffering = True  # collect into aBuffer until open() is called or buffer is dropped

    @classmethod
    def open(cls, bReplace=False, pathDir=None):
        try:
            sLogPath = sLogFileName
            if pathDir is not None:
                sLogPath = str(pathlib.Path(pathDir) / sLogFileName)
            cls.sPath = sLogPath
            if bReplace:
                # Replace mode: delete any prior session's log first so
                # the new file contains only this session's output.
                # unlink may fail if the file does not exist
                # (FileNotFoundError) -- that's fine; just continue.
                try:
                    os.unlink(sLogPath)
                except FileNotFoundError:
                    pass
                except Exception:
                    # If unlink fails for some other reason (e.g.,
                    # another process has the file open), fall through
                    # to open() in "w" mode, which will truncate.
                    # open() may also fail, and the outer except will
                    # mark logger disabled -- the right outcome (better
                    # to lose logging than mix old and new content).
                    pass
                cls.fLog = open(sLogPath, "w", encoding="utf-8-sig", newline="\r\n")
                cls.bEnabled = True
                cls.bBuffering = False
                # Flush any pre-open diagnostics, then write the
                # log-opened banner. The banner appears BEFORE pre-
                # open content would be confusing; instead, flush
                # buffer first so the order matches the time the
                # events actually occurred.
                cls._flushBuffer()
                cls.info(f"{sProgramName} {sProgramVersion} log opened (replaced prior {sLogFileName})")
            else:
                # Append mode (default): preserve history across runs.
                # The first time the file is opened it is created fresh;
                # subsequent runs append. A blank line is written before
                # the new session's header to visually separate runs.
                bExistedBefore = os.path.exists(sLogPath)
                cls.fLog = open(sLogPath, "a", encoding="utf-8-sig", newline="\r\n")
                cls.bEnabled = True
                cls.bBuffering = False
                if bExistedBefore: cls.fLog.write("\n")
                cls._flushBuffer()
                cls.info(f"{sProgramName} {sProgramVersion} log opened (appending)")
        except Exception:
            cls.bEnabled = False
            cls.bBuffering = False  # give up; drop buffer
            cls.aBuffer = []

    @classmethod
    def _flushBuffer(cls):
        """Write any pre-open buffered diagnostics to the now-open log
        file. Called by open() after the file is ready. Buffer is
        cleared so subsequent calls don't re-write."""
        if cls.fLog is None: return
        for sStamp, sLevel, sMsg in cls.aBuffer:
            try: cls.fLog.write(f"[{sStamp}] [{sLevel}] {sMsg}\n")
            except Exception: pass
        try: cls.fLog.flush()
        except Exception: pass
        cls.aBuffer = []

    @classmethod
    def discardBuffer(cls):
        """Discard the pre-open buffer without writing it. Called when
        we know logging will never happen for this run, so memory can
        be freed and bBuffering disabled."""
        cls.aBuffer = []
        cls.bBuffering = False

    @classmethod
    def write(cls, sLevel, sMsg):
        # Internal level-tagged writer. Public level methods (info,
        # warn, error, debug) all funnel through here so the format
        # is uniform and a level filter could be added later in one
        # place. Two paths:
        #
        #   1. Logger is enabled (file open): write directly.
        #   2. Logger is buffering (pre-open): append to aBuffer, to
        #      be flushed by open() when the file is ready.
        #
        # The buffering path exists because urlCheck needs to record
        # diagnostics before the output folder is known (e.g., GUI
        # mode hasn't shown the dialog yet, so pathLogDir isn't
        # final). Without buffering, those diagnostics would either
        # be lost or land in the wrong folder.
        sStamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if cls.bEnabled and cls.fLog is not None:
            try:
                cls.fLog.write(f"[{sStamp}] [{sLevel}] {sMsg}\n")
                cls.fLog.flush()
            except Exception:
                pass
        elif cls.bBuffering:
            try:
                cls.aBuffer.append((sStamp, sLevel, sMsg))
                # Cap the buffer to avoid unbounded growth in extreme
                # cases (some unforeseen failure mode where the open
                # never happens). 500 messages is plenty for the
                # pre-dialog diagnostic phase.
                if len(cls.aBuffer) > 500:
                    cls.aBuffer = cls.aBuffer[-500:]
            except Exception:
                pass

    @classmethod
    def info(cls, sMsg):  cls.write("INFO",  sMsg)
    @classmethod
    def warn(cls, sMsg):  cls.write("WARN",  sMsg)
    @classmethod
    def error(cls, sMsg): cls.write("ERROR", sMsg)
    @classmethod
    def debug(cls, sMsg): cls.write("DEBUG", sMsg)

    @classmethod
    def header(cls, sName, sVersion, lParams):
        """Write the run header to the log: program name and version,
        the friendly run-start timestamp, and the resolved parameter
        list. Emits raw lines (no per-line timestamp/level prefix) so
        the header reads as a clean banner. The processing
        notifications that follow use the standard format via
        info/warn/error/debug.

        lParams is a list of (label, value) tuples; the caller controls
        the order in which parameters appear.
        """
        if not cls.bEnabled or cls.fLog is None: return
        try:
            cls.fLog.write(f"=== {sName} {sVersion} ===\n")
            cls.fLog.write(f"Run on {cls.friendlyTime(datetime.datetime.now())}\n")
            if lParams:
                cls.fLog.write("Parameters:\n")
                iPad = max((len(sLabel) for sLabel, _ in lParams), default=0)
                for sLabel, sValue in lParams:
                    cls.fLog.write(f"  {sLabel.ljust(iPad)} : {sValue}\n")
            cls.fLog.write("===\n")
            cls.fLog.flush()
        except Exception:
            pass

    @classmethod
    def friendlyTime(cls, dt):
        """Render a datetime in a friendly form, e.g.,
        'May 1, 2026 at 2:30 PM'. Removes a leading zero from the
        day and the hour for natural reading.
        """
        sMonth = dt.strftime("%B")
        sDay = str(dt.day)
        sYear = str(dt.year)
        sHour12 = str(((dt.hour - 1) % 12) + 1)
        sMin = f"{dt.minute:02d}"
        sAmPm = "AM" if dt.hour < 12 else "PM"
        return f"{sMonth} {sDay}, {sYear} at {sHour12}:{sMin} {sAmPm}"

    @classmethod
    def close(cls):
        try:
            if cls.fLog is not None: cls.fLog.close()
        except Exception:
            pass
        cls.fLog = None
        cls.bEnabled = False


# --- Config manager (opt-in, INI under %LOCALAPPDATA%\urlCheck) ---

class configManager:
    """
    Persists user preferences to %LOCALAPPDATA%\\urlCheck\\urlCheck.ini, but
    only when the user opts in via -u / --use-configuration or via the GUI
    checkbox. Without opt-in urlCheck leaves no filesystem footprint of its
    own beyond the per-scan output folders the user explicitly asks for.
    """

    @staticmethod
    def getConfigDir():
        sLocal = os.environ.get("LOCALAPPDATA", "") or os.path.expanduser("~")
        return os.path.join(sLocal, sConfigDirName)

    @staticmethod
    def getConfigPath():
        return os.path.join(configManager.getConfigDir(), sConfigFileName)

    @staticmethod
    def configExists():
        try: return os.path.isfile(configManager.getConfigPath())
        except Exception: return False

    @staticmethod
    def eraseAll():
        sDir = configManager.getConfigDir()
        sPath = configManager.getConfigPath()
        try:
            if os.path.isfile(sPath):
                os.remove(sPath)
                logger.info(f"Deleted configuration file: {sPath}")
        except Exception as ex:
            logger.info(f"Could not delete configuration file {sPath}: {ex}")
        try:
            if os.path.isdir(sDir) and not os.listdir(sDir):
                os.rmdir(sDir)
                logger.info(f"Removed empty configuration folder: {sDir}")
        except Exception as ex:
            logger.info(f"Could not remove configuration folder {sDir}: {ex}")

    @staticmethod
    def parseFile(sPath):
        d = {}
        sLine = ""
        sLineRaw = ""
        with open(sPath, "r", encoding="utf-8-sig") as fIni:
            for sLineRaw in fIni:
                sLine = sLineRaw.strip()
                if not sLine: continue
                if sLine.startswith(";") or sLine.startswith("#"): continue
                if sLine.startswith("[") and sLine.endswith("]"): continue
                iEq = sLine.find("=")
                if iEq <= 0: continue
                d[sLine[:iEq].strip().lower()] = sLine[iEq + 1:].strip()
        return d

    @staticmethod
    def getBool(d, sKey):
        s = (d.get(sKey, "") or "").strip().lower()
        return s in ("1", "true", "yes", "on")

    @staticmethod
    def loadInto(arguments):
        """
        Load saved values into the parsed-arguments object. CLI-supplied
        values always win: only fields the user did NOT specify on the command
        line are overwritten. The cli-supplied set is reconstructed by
        comparing against parser defaults.
        """
        d = {}
        sPath = ""

        sPath = configManager.getConfigPath()
        if not os.path.isfile(sPath): return
        try:
            d = configManager.parseFile(sPath)
        except Exception as ex:
            print(f"[WARN] Could not read configuration from {sPath}: {ex}")
            return

        # Source: only adopt the saved value if the CLI provided no positional.
        if (not getattr(arguments, "sSource", None)) and d.get("source", ""):
            arguments.sSource = d.get("source", "")

        # Output folder: only adopt if the CLI did not pass -o.
        if not getattr(arguments, "sOutputDir", "") and d.get("output_folder", ""):
            arguments.sOutputDir = d.get("output_folder", "")

        # Booleans: only adopt if the CLI did not pass the flag (i.e., it's
        # currently False, the parser default).
        if not getattr(arguments, "bViewOutput", False):
            arguments.bViewOutput = configManager.getBool(d, "view_output")
        if not getattr(arguments, "bInvisible", False):
            arguments.bInvisible = configManager.getBool(d, "invisible")
        if not getattr(arguments, "bAuthenticate", False):
            arguments.bAuthenticate = configManager.getBool(d, "authenticate")
        if not getattr(arguments, "bMainProfile", False):
            arguments.bMainProfile = configManager.getBool(d, "main_profile")
        if not getattr(arguments, "bForce", False):
            arguments.bForce = configManager.getBool(d, "force_replacements")
        if not getattr(arguments, "bLog", False):
            arguments.bLog = configManager.getBool(d, "log_session")

    @staticmethod
    def save(sSource, sOutputDir, bViewOutput, bInvisible, bForce, bLog, bAuthenticate=False, bMainProfile=False):
        sDir = ""
        sPath = ""

        sDir = configManager.getConfigDir()
        sPath = configManager.getConfigPath()
        try:
            if not os.path.isdir(sDir): os.makedirs(sDir, exist_ok=True)
            with open(sPath, "w", encoding="utf-8-sig", newline="\r\n") as fIni:
                fIni.write("; urlCheck configuration\n")
                fIni.write("; auto-written when Use configuration was checked at OK time.\n")
                fIni.write("; Delete this file to reset, or click Default settings in the\n")
                fIni.write("; GUI, which also deletes the file and the urlCheck folder.\n")
                fIni.write(f"source={sSource or ''}\n")
                fIni.write(f"output_folder={sOutputDir or ''}\n")
                fIni.write(f"view_output={'1' if bViewOutput else '0'}\n")
                fIni.write(f"invisible={'1' if bInvisible else '0'}\n")
                fIni.write(f"authenticate={'1' if bAuthenticate else '0'}\n")
                fIni.write(f"main_profile={'1' if bMainProfile else '0'}\n")
                fIni.write(f"force_replacements={'1' if bForce else '0'}\n")
                fIni.write(f"log_session={'1' if bLog else '0'}\n")
            logger.info(f"Saved configuration to {sPath}")
        except Exception as ex:
            print(f"[WARN] Could not save configuration to {sPath}: {ex}")
            logger.info(f"Could not save configuration: {ex}")




# --- ACR (Accessibility Conformance Report) builder ---
#
# Produces ACR.xlsx in the parent output folder, summarizing axe-core
# results across multiple per-page scans against the WCAG 2.2 success
# criteria. Each WCAG criterion gets a row showing how many distinct
# axe rules failed / passed / were incomplete / were inapplicable for
# that criterion across the included pages, plus a calculated verdict
# (Calc) and an ACR-terminology mapping (Conformance) suitable for
# inclusion in a published Accessibility Conformance Report.
#
# Architecture: walk the parent output folder for subfolders that
# contain a results.json, parse each, accumulate per-criterion counts,
# then render a workbook with one rollup sheet plus one sheet per
# included page.
#
# Source-of-truth rules:
#   --force set: build ACR.xlsx from ONLY the URLs scanned in this
#                session (in-memory tracked list; older subfolders are
#                ignored even if they still contain results.json).
#   --force unset: build ACR.xlsx from ALL subfolders under the parent
#                that contain a results.json. The user curates the
#                workbook by deleting subfolders to exclude their
#                pages.
#
# Calc precedence (per criterion, per page):
#   1. fail+pass on different rules -> partial
#   2. fail with no pass            -> fail
#   3. incomplete or needs-review   -> manual
#   4. pass with none failing       -> pass
#   5. only inapplicable            -> na
#   6. no axe coverage              -> unknown
#
# Rollup precedence (worst-result-wins across pages):
#   any partial OR any (fail+pass) -> partial
#   any fail (no pass anywhere)    -> fail
#   any manual                      -> manual
#   any pass                       -> pass
#   any na                         -> na
#   no coverage                     -> unknown
#
# ACR terminology mapping (per VPAT 2.5):
#   pass     -> Supports
#   partial  -> Partially Supports
#   fail     -> Does Not Support
#   manual   -> Incomplete
#   na       -> Not Applicable
#   unknown  -> Not Evaluated

class acrBuilder:
    """
    Build ACR.xlsx for the urlCheck session.

    Public methods:
      classmethod buildIfApplicable(arguments, lThisRunPaths) -- driver.
        lThisRunPaths is a list of pathlib.Path objects pointing to the
        per-page output folders written during this run. With --force,
        this list is the ONLY source. Without --force, it's ignored and
        the parent folder is walked instead.
    """

    @staticmethod
    def fnWcagSortKey(sId):
        """Numeric sort key so 1.10 follows 1.9, not 1.2."""
        try: return tuple(int(p) for p in sId.split("."))
        except Exception: return (999,)

    @classmethod
    def discoverPageFolders(cls, pathParent, lThisRunPaths, bForceMode):
        """
        Return a list of (pathFolder, dResults, dMetadata) tuples for
        every page that should be included in the workbook.
        """
        lFound = []
        lPaths = []
        if bForceMode:
            # In --force mode the workbook is built strictly from this
            # session's tracked outputs.
            lPaths = list(lThisRunPaths or [])
        else:
            # Append-mode: walk the parent for every subfolder that has
            # a results.json. Sort alphabetically for stable sheet
            # ordering across runs.
            for path in sorted(pathParent.iterdir()):
                if not path.is_dir(): continue
                if (path / sJsonName).is_file():
                    lPaths.append(path)
        for path in lPaths:
            pathJson = path / sJsonName
            if not pathJson.is_file(): continue
            try:
                d = json.loads(pathJson.read_text(encoding="utf-8"))
            except Exception as ex:
                logger.info(f"ACR: skipping {path.name}: cannot parse "
                    f"{sJsonName}: {ex}")
                continue
            dResults = d.get("results") or {}
            dMetadata = d.get("metadata") or {}
            lFound.append((path, dResults, dMetadata))
        return lFound

    @classmethod
    def axeVersionFromBundle(cls, sAxeSource):
        """
        Extract the axe-core version from the bundled axe.min.js source
        text. Falls back to sFallbackAxeVersion if not found. The
        version number is used to construct Deque University rule URLs.
        """
        if not sAxeSource: return sFallbackAxeVersion
        sCheck = sAxeSource[:8000]
        # Common shape: axe.version="4.10.2" or var version="4.10.2"
        m = re.search(r'(?:axe\.version|version)\s*[=:]\s*["\']([0-9]+\.[0-9]+(?:\.[0-9]+)?)["\']', sCheck)
        if m:
            sVer = m.group(1)
            # Major.minor for Deque URLs
            aParts = sVer.split(".")
            if len(aParts) >= 2: return f"{aParts[0]}.{aParts[1]}"
            return sVer
        return sFallbackAxeVersion

    @classmethod
    def criteriaFromTags(cls, lTags):
        """
        Given an axe rule's tags (e.g. ["wcag2a","wcag111","cat.text-alternatives"]),
        return the list of WCAG 2.2 criterion ids the rule applies to
        (e.g. ["1.1.1"]). A single rule can cover multiple criteria.
        Tags use the convention wcag<major><minor><sub> with no dots,
        e.g. wcag111 -> 1.1.1, wcag1410 -> 1.4.10.
        """
        lOut = []
        for sTag in lTags or []:
            sTag = str(sTag)
            if not sTag.startswith("wcag"): continue
            sRest = sTag[4:]
            # Skip level-only tags like wcag2a, wcag2aa, wcag22aa
            if not sRest or not sRest[0].isdigit(): continue
            # Skip combined level tags like wcag22aa
            if any(ch.isalpha() for ch in sRest): continue
            # Now sRest is digits like 111, 1410, 2410.
            # Parse: first digit = principle, second digit = guideline,
            # remaining digits = criterion (1 or 2 digits).
            if len(sRest) < 3: continue
            sCrit = f"{sRest[0]}.{sRest[1]}.{sRest[2:]}"
            if sCrit in dWcag22 and sCrit not in lOut:
                lOut.append(sCrit)
        return lOut

    @classmethod
    def perPageBuckets(cls, dResults):
        """
        Convert one page's axe results into a dict keyed by criterion id
        with values:
          {
            "fail": {rule_id: instance_count, ...},
            "pass": {rule_id: instance_count, ...},
            "incomplete": {rule_id: instance_count, ...},
            "na": {rule_id: instance_count, ...},
          }
        Instance count is the length of the rule's nodes array on this
        page. For inapplicable rules, the instance count is recorded
        as 1 (the rule itself was tested and found not applicable;
        nodes is empty by axe's design).

        For pass / incomplete / inapplicable to have accurate instance
        counts, axe.run must be called with all four resultTypes
        listed (otherwise axe truncates pass and inapplicable nodes
        arrays to one entry per rule). See aAxeRunOptions.
        """
        d = {}
        dKeyToBucket = {
            "violations":   "fail",
            "passes":       "pass",
            "incomplete":   "incomplete",
            "needsReview":  "incomplete",
            "inapplicable": "na",
        }
        for sKey, sBucket in dKeyToBucket.items():
            for dRule in dResults.get(sKey, []) or []:
                sRuleId = dRule.get("id") or ""
                if not sRuleId: continue
                lTags = dRule.get("tags") or []
                lCrits = cls.criteriaFromTags(lTags)
                # Instance count: length of nodes array. For
                # inapplicable, the array is empty by design; count
                # the rule itself once.
                aNodes = dRule.get("nodes") or []
                iCount = len(aNodes) if aNodes else (1 if sBucket == "na" else 0)
                if iCount == 0: continue
                for sCrit in lCrits:
                    if sCrit not in d:
                        d[sCrit] = {"fail": {}, "pass": {},
                                    "incomplete": {}, "na": {}}
                    if sRuleId in d[sCrit][sBucket]:
                        d[sCrit][sBucket][sRuleId] += iCount
                    else:
                        d[sCrit][sBucket][sRuleId] = iCount
        return d

    @classmethod
    def calcVerdict(cls, dBuckets):
        """
        Apply Calc precedence to a per-criterion buckets dict.
        Returns one of: "partial", "fail", "manual", "pass", "na", "unknown".

        Precedence (most specific first):
          1. partial -- at least one fail AND at least one pass (any
                        incompletes are additional unfinished work; the
                        mixed verdict still stands)
          2. fail    -- at least one fail (no pass anywhere)
          3. manual  -- at least one incomplete (no fail, no pass)
          4. pass    -- at least one pass (no fail, no incomplete)
          5. na      -- only inapplicable
          6. unknown -- no axe coverage
        """
        if not dBuckets: return "unknown"
        bHasFail = bool(dBuckets["fail"])
        bHasPass = bool(dBuckets["pass"])
        bHasIncomplete = bool(dBuckets["incomplete"])
        bHasNa = bool(dBuckets["na"])
        if bHasFail and bHasPass: return "partial"
        if bHasFail: return "fail"
        if bHasIncomplete: return "manual"
        if bHasPass: return "pass"
        if bHasNa: return "na"
        return "unknown"

    @classmethod
    def rollupVerdict(cls, lPageBuckets, sCrit):
        """
        Compute the rollup verdict across multiple pages for a given
        criterion. lPageBuckets is a list of per-page bucket dicts (the
        output of perPageBuckets). Returns the calc verdict string.
        Worst-result-wins across pages.
        """
        bAnyFail = False
        bAnyPass = False
        bAnyIncomplete = False
        bAnyNa = False
        bAnyCovered = False
        for d in lPageBuckets:
            if sCrit not in d: continue
            bAnyCovered = True
            b = d[sCrit]
            if b["fail"]:       bAnyFail = True
            if b["pass"]:       bAnyPass = True
            if b["incomplete"]: bAnyIncomplete = True
            if b["na"]:         bAnyNa = True
        if not bAnyCovered: return "unknown"
        if bAnyFail and bAnyPass: return "partial"
        if bAnyFail:              return "fail"
        if bAnyIncomplete:        return "manual"
        if bAnyPass:              return "pass"
        if bAnyNa:                return "na"
        return "unknown"

    @classmethod
    def calcToConformance(cls, sCalc):
        """Map internal Calc value to VPAT 2.5-style ACR terminology.

        The mapping follows ITI guidance and WCAG 2.0 Understanding
        Conformance: when no content on the page is subject to a
        criterion, the criterion IS satisfied (a vacuous truth), so
        the appropriate term is "Supports", not "Not Applicable".
        See https://www.w3.org/TR/UNDERSTANDING-WCAG20/conformance.html

        VPAT 2.5 reserves "Not Evaluated" for AAA criteria, but
        urlCheck uses it more broadly as the honest verdict for
        criteria where automated testing cannot decide. urlCheck
        produces a draft ACR; the user is expected to revisit each
        "Not Evaluated" entry, perform the manual checks listed in
        the Manual column, and rewrite the Result column with their
        final verdict before publishing.
        """
        d = {
            "pass":     "Supports",
            "partial":  "Partially Supports",
            "fail":     "Does Not Support",
            "manual":   "Not Evaluated",
            "na":       "Supports",
            "unknown":  "Not Evaluated",
        }
        return d.get(sCalc, "Not Evaluated")

    @classmethod
    def aggregateInstances(cls, lPageBuckets, sCrit, sBucket):
        """
        Sum instance counts across all pages for the given criterion
        and bucket. Returns total node-instance count.
        """
        iTotal = 0
        for d in lPageBuckets:
            if sCrit not in d: continue
            for sRuleId, iCount in d[sCrit].get(sBucket, {}).items():
                iTotal += iCount
        return iTotal

    @classmethod
    def aggregateAllRules(cls, lPageBuckets, sCrit):
        """All distinct rule ids that touched the criterion across all
        pages, in any bucket."""
        sAll = set()
        for d in lPageBuckets:
            if sCrit not in d: continue
            for sBucket in ("fail", "pass", "incomplete", "na"):
                for sRuleId in d[sCrit].get(sBucket, {}):
                    sAll.add(sRuleId)
        return sorted(sAll)

    @classmethod
    def pagesByVerdict(cls, lPageBuckets, lPageNames, sCrit):
        """
        Return a dict { 'fail': [...], 'partial': [...], 'manual': [...] }
        of sheet-name lists for the criterion, where each list contains
        the page names whose per-page Calc landed in that bucket.

        lPageBuckets and lPageNames are parallel lists. Pages with calc
        of pass / na / unknown are not listed (they're not actionable
        in the rollup's Conformance cell).
        """
        dOut = {"fail": [], "partial": [], "manual": []}
        for dB, sName in zip(lPageBuckets, lPageNames):
            buckets = dB.get(sCrit) or {}
            if not buckets: continue
            sCalc = cls.calcVerdict(buckets)
            if sCalc in dOut:
                dOut[sCalc].append(sName)
        for k in dOut: dOut[k].sort()
        return dOut

    @classmethod
    def captureExistingRemarks(cls, pathWorkbook):
        """
        If ACR.xlsx already exists, capture the user-editable Result
        and Remarks columns from Sheet 1 (keyed by criterion id) so
        they survive regeneration. Returns dict { sCrit: (sResult, sRemarks) }.
        """
        dCaptured = {}
        if not pathWorkbook.is_file(): return dCaptured
        try:
            # Open WITHOUT data_only so HYPERLINK formulas are
            # readable as raw text. data_only=True returns None for
            # formula cells unless Excel has cached results, which
            # openpyxl-written files do not have.
            wb = openpyxl.load_workbook(str(pathWorkbook))
        except Exception as ex:
            logger.info(f"ACR: cannot read existing {pathWorkbook.name} "
                f"to preserve user edits: {ex}")
            return dCaptured
        sName = "Conformance Report"
        if sName not in wb.sheetnames:
            wb.close()
            return dCaptured
        ws = wb[sName]
        # Find Result and Remarks column indices from header row.
        # Also accept the legacy "Notes" header for upgraded workbooks.
        iResult = -1
        iRemarks = -1
        iCrit = -1
        for iCol, cell in enumerate(ws[1], 1):
            sVal = (cell.value or "").strip() if cell.value else ""
            if sVal == "Result": iResult = iCol
            elif sVal in ("Remarks", "Notes"): iRemarks = iCol
            elif sVal == "Criterion": iCrit = iCol
        if iCrit < 0:
            wb.close()
            return dCaptured
        for row in ws.iter_rows(min_row=2):
            sCritCell = ""
            sResult = ""
            sRemarks = ""
            for cell in row:
                if cell.column == iCrit and cell.value:
                    s = str(cell.value).strip()
                    if s.startswith("="):
                        m = re.search(r'HYPERLINK\s*\(\s*"[^"]*"\s*,\s*"([^"]*)"\s*\)', s, re.IGNORECASE)
                        if m: s = m.group(1)
                    sCritCell = s.split()[0] if s else ""
                elif cell.column == iResult and cell.value:
                    sResult = str(cell.value).strip()
                elif cell.column == iRemarks and cell.value:
                    sRemarks = str(cell.value).strip()
            if sCritCell and (sResult or sRemarks):
                dCaptured[sCritCell] = (sResult, sRemarks)
        wb.close()
        return dCaptured

    @classmethod
    def writeGlossarySheet(cls, wb):
        """
        Write a Glossary sheet at the end of the workbook with key
        terms and links the user might want when reading the report.

        Sections (separated by blank rows), ordered axe -> urlCheck ->
        WCAG -> misc:
          1. Outcome Categories (axe-core)
          2. urlCheck Calc Values
          3. VPAT 2.5 Conformance Terms
          4. WCAG Accessibility Principles
          5. WCAG Conformance Levels
          6. Additional Resources

        Each section is its own contiguous data region with two
        columns (Term, Definition). Each section's "Term" header
        cell (the column-header cell of that region) is named
        Title02 / Title03 / etc. so screen readers announce both
        the column header and the row label as the user navigates
        within that section.
        """
        ws = wb.create_sheet(title="Glossary")
        iRow = 1
        aHeaderRows = []  # list of row numbers of "Term" header cells

        def fnSection(sHeader, lEntries):
            nonlocal iRow
            # Section heading (above the data region, decorative).
            ws.cell(row=iRow, column=1, value=sHeader)
            iRow += 1
            # Data region begins here. The first row is the column-
            # header row (Term, Definition). Track this row so the
            # named-range pass can attach a Title<NN> name to it.
            aHeaderRows.append(iRow)
            ws.cell(row=iRow, column=1, value="Term")
            ws.cell(row=iRow, column=2, value="Definition")
            iRow += 1
            for sTerm, sDefn in lEntries:
                ws.cell(row=iRow, column=1, value=sTerm)
                # If the definition is a bare URL, set both the
                # cell text AND the cell's native hyperlink property.
                # A real Excel hyperlink (as opposed to a HYPERLINK
                # formula) is recognized by screen readers as a link
                # element, and renders with the standard link
                # appearance. The native hyperlink approach is more
                # accessible than the formula approach because:
                # (a) JAWS announces native hyperlink cells as links;
                # (b) the link is followable by Enter or Ctrl+Click
                #     without Excel needing to evaluate a formula.
                cell = ws.cell(row=iRow, column=2)
                if (sDefn.startswith("http://") or sDefn.startswith("https://")) and " " not in sDefn:
                    cell.value = sDefn
                    cell.hyperlink = sDefn
                    # Apply Excel's built-in Hyperlink style so the
                    # cell renders blue and underlined, matching the
                    # visual convention users expect.
                    cell.style = "Hyperlink"
                else:
                    cell.value = sDefn
                iRow += 1
            iRow += 1  # blank row between sections

        fnSection("Outcome Categories (axe-core)", [
            ("Fail",
             "An axe rule found one or more failing instances on the page."),
            ("Pass",
             "An axe rule confirmed at least one passing instance."),
            ("Incomplete",
             "An axe rule could not definitively decide; manual review needed."),
            ("Inapplicable",
             "An axe rule was tested but no relevant content existed on the page."),
        ])

        fnSection("urlCheck Calc Values", [
            ("fail",
             "At least one rule instance fails, with no passes on the same criterion."),
            ("pass",
             "At least one rule instance passes, with no fails or incompletes."),
            ("partial",
             "At least one rule instance fails AND at least one passes for the same criterion."),
            ("manual",
             "At least one incomplete result, with no fails or passes; manual testing is required."),
            ("na",
             "All rule instances are inapplicable to the page."),
            ("unknown",
             "No automated rules apply to this criterion. Default before the page has been scanned."),
        ])

        fnSection("VPAT 2.5 Conformance Terms", [
            ("Supports",
             "The functionality of the product has at least one method that meets the criterion without known defects, or meets with equivalent facilitation."),
            ("Partially Supports",
             "Some functionality of the product does not meet the criterion."),
            ("Does Not Support",
             "The majority of product functionality does not meet the criterion."),
            ("Not Applicable",
             "The criterion is not relevant to the product."),
            ("Incomplete",
             "Manual review is required to reach a conformance verdict."),
            ("Not Evaluated",
             "The product has not been evaluated against the criterion. May only be used in WCAG Level AAA criteria."),
        ])

        fnSection("WCAG Accessibility Principles", [
            ("Perceivable",
             "Information and user interface components must be presentable to users in ways they can perceive."),
            ("Operable",
             "User interface components and navigation must be operable."),
            ("Understandable",
             "Information and the operation of the user interface must be understandable."),
            ("Robust",
             "Content must be robust enough that it can be interpreted reliably by a wide variety of user agents, including assistive technology."),
        ])

        fnSection("WCAG Conformance Levels", [
            ("Level A",
             "Minimum level of accessibility conformance. Removes the most significant barriers."),
            ("Level AA",
             "Addresses the major and most common accessibility barriers; commonly required by laws and procurement standards."),
            ("Level AAA",
             "Highest and most comprehensive level. Some content may not be able to satisfy all AAA criteria."),
        ])

        fnSection("Additional Resources", [
            ("WCAG 2.2 Quick Reference",
             "https://www.w3.org/WAI/WCAG22/quickref/"),
            ("Understanding WCAG 2.2",
             "https://www.w3.org/WAI/WCAG22/Understanding/"),
            ("Deque University rule reference",
             "https://dequeuniversity.com/rules/axe/"),
            ("WebAIM",
             "https://webaim.org/"),
            ("ITI VPAT (Voluntary Product Accessibility Template)",
             "https://www.itic.org/policy/accessibility/vpat"),
            ("Section 508",
             "https://www.section508.gov/"),
        ])

        # The Glossary's row 1 is a section heading (not a column
        # header), so applyFormatting with iHeaderRow=1 would bold it
        # white-on-blue along with all the data. We deliberately call
        # applyFormatting only for column-width sizing, then re-style
        # each section's "Term/Definition" header row separately.
        cls.applyFormatting(ws, iHeaderRow=1, iDataStart=2)
        # Apply white-on-blue header styling to each section's header
        # row, and attach a Title<NN> named range to each section's
        # "Term" cell so screen readers correctly announce headers
        # as the user navigates within that section.
        for iHdrRow in aHeaderRows:
            for iCol in (1, 2):
                cell = ws.cell(row=iHdrRow, column=iCol)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(fill_type="solid", fgColor="1F4E78")
                cell.alignment = Alignment(vertical="top", wrap_text=True)
            cls.addNamedRangeForCell(wb, ws, ws.cell(row=iHdrRow, column=1), "Title")
        # Section-heading cells (the rows above each header row, with
        # text like "WCAG Accessibility Principles") should NOT be
        # styled as if they were data; restore plain bold formatting.
        for iHdrRow in aHeaderRows:
            iSectHeadingRow = iHdrRow - 1
            if iSectHeadingRow < 1: continue
            cell = ws.cell(row=iSectHeadingRow, column=1)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(fill_type=None)
            cell.alignment = Alignment(vertical="top")

    @classmethod
    def safeSheetName(cls, sBase, lExisting):
        """Sanitize a folder name for use as a sheet name. Excel sheet
        names are limited to 31 characters and may not contain
        : \\ / ? * [ ]. The folder name (already a sanitized page
        title) is reused directly; if it exceeds 31 chars or if a
        prior sheet has the same name, suffix with _2, _3, etc."""
        sClean = re.sub(r'[:\\/?*\[\]]', '_', sBase or "Page")
        sClean = sClean[:31]
        if sClean and sClean not in lExisting: return sClean
        # Disambiguate
        for iSuffix in range(2, 1000):
            sSuffix = f"_{iSuffix}"
            sTry = (sClean[:31 - len(sSuffix)]) + sSuffix
            if sTry not in lExisting: return sTry
        return (sClean[:30] or "Page")

    @classmethod
    def applyFormatting(cls, ws, iHeaderRow=1, iDataStart=2, bSingleLineRows=False):
        """
        Apply the xlFormat.vbs-derived formatting to a worksheet:
        bold first row with white-on-blue fill, vertical-align top
        everywhere, per-column data-driven width capped at 50, wrap
        text on columns over the cap, frozen header row. AutoFilter
        is deliberately not enabled: in JAWS, the surrounding "filter
        is off" announcement on every cell interferes with reading
        column headers.

        When bSingleLineRows is True, each data row is given an
        explicit fixed height of one line (~15 points) so that multi-
        line cells show only their first line by default. Excel
        honors the explicit height even when wrap_text is True;
        wrapped content remains in the cell but overflows below the
        visible area until the user expands the row. This is the
        idiom used for Sheet 1 (Conformance Report) and per-URL
        sheets so the user sees one row = one line of summary.

        Idempotent: safe to call after data is written.
        """
        if ws.max_row < iHeaderRow: return
        iCols = ws.max_column
        if iCols < 1: return
        # Header row: bold, white-on-blue
        for iCol in range(1, iCols + 1):
            cell = ws.cell(row=iHeaderRow, column=iCol)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(fill_type="solid", fgColor="1F4E78")
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        # Compute per-column max content length and cap widths.
        iCapWidth = 50
        for iCol in range(1, iCols + 1):
            iMaxLen = 0
            for iRow in range(iHeaderRow, ws.max_row + 1):
                v = ws.cell(row=iRow, column=iCol).value
                if v is None: continue
                # Multi-line cells: longest line determines width.
                aLines = str(v).split("\n")
                for sLine in aLines:
                    if len(sLine) > iMaxLen: iMaxLen = len(sLine)
            iWidth = min(iMaxLen + 2, iCapWidth)
            if iWidth < 6: iWidth = 6
            ws.column_dimensions[get_column_letter(iCol)].width = iWidth
            # Wrap text on columns at the cap (where overflow likely)
            bWrap = (iMaxLen + 2) > iCapWidth
            for iRow in range(iDataStart, ws.max_row + 1):
                cell = ws.cell(row=iRow, column=iCol)
                cell.alignment = Alignment(vertical="top", wrap_text=bWrap)
        # Single-line row-height enforcement: each data row is fixed
        # to one line of content; multi-line cells overflow
        # invisibly below until the user expands the row.
        if bSingleLineRows:
            for iRow in range(iDataStart, ws.max_row + 1):
                ws.row_dimensions[iRow].height = 15
        # Freeze header
        ws.freeze_panes = ws.cell(row=iDataStart, column=1).coordinate

    @classmethod
    def addNamedRangeForCell(cls, wb, ws, cell, sStem):
        """
        Apply the xlHeaders.vbs technique: name a specific column-
        header cell with a unique <stem>NN identifier so screen
        readers (JAWS in particular) announce both the column header
        and the row label as the user navigates within the data
        region. Pass the actual cell that is the column header (the
        top-left cell of a contiguous data region's header row),
        not always row 1 column 1 of the sheet.

        Stem is "Title" when both column headers (row 1) and row
        labels (column 1) are present. The Glossary's multi-section
        layout uses one named cell per section; the per-URL and
        rollup sheets each have a single region.
        """
        from openpyxl.workbook.defined_name import DefinedName
        # Find next free name with this stem.
        iSeq = 1
        while True:
            sName = f"{sStem}{iSeq:02d}"
            if sName not in wb.defined_names: break
            iSeq += 1
        sRef = f"'{ws.title}'!${cell.column_letter}${cell.row}"
        wb.defined_names[sName] = DefinedName(name=sName, attr_text=sRef)

    @classmethod
    def writePerUrlSheet(cls, wb, sSheetName, dPageBuckets, dMetadata,
                         sAxeVer, lOrderedCrits):
        """
        Write a per-URL diagnostic sheet.

        Columns:
          1. Criterion        (HYPERLINK; format "1.1.1 Non-text Content (A)")
          2. Summary          (≤50 chars one-sentence summary)
          3. Calc             (pass / fail / partial / manual / na / unknown)
          4. Fail             (multi-line: "rule-id N" per line, alpha order)
          5. Pass             (same)
          6. Incomplete       (same)
          7. Inapplicable     (same)

        N is the node-instance count for that rule on this page.
        Per-URL sheets are diagnostic; no Conformance / Result / Notes
        columns. The page identification (title and URL) is in rows
        above the header row.
        """
        ws = wb.create_sheet(title=sSheetName)
        sPageTitle = str(dMetadata.get("pageTitle") or "")
        sPageUrl = str(dMetadata.get("pageUrl") or "")
        ws.cell(row=1, column=1, value=f"Page: {sPageTitle}")
        ws.cell(row=2, column=1, value=f"URL: {sPageUrl}")
        iHeaderRow = 4
        iDataStart = iHeaderRow + 1
        aHeaders = ["Criterion", "Summary", "Calc",
                    "Fail", "Pass", "Incomplete", "Inapplicable"]
        for iCol, sH in enumerate(aHeaders, 1):
            ws.cell(row=iHeaderRow, column=iCol, value=sH)
        iRow = iDataStart
        # Mapping from header column to bucket key.
        aBucketCols = [("fail", 4), ("pass", 5), ("incomplete", 6), ("na", 7)]
        for sCrit in lOrderedCrits:
            d = dWcag22[sCrit]
            sCritText = f"{sCrit} {d['name']} ({d['level']})"
            sQuickRef = sWcagQuickRefBase + "#" + d["fragment"]
            cell = ws.cell(row=iRow, column=1)
            cell.value = f'=HYPERLINK("{sQuickRef}","{sCritText}")'
            ws.cell(row=iRow, column=2, value=dWcag22Summary.get(sCrit, ""))
            buckets = dPageBuckets.get(sCrit) or {}
            sCalc = cls.calcVerdict(buckets) if buckets else "unknown"
            ws.cell(row=iRow, column=3, value=sCalc)
            for sBucket, iCol in aBucketCols:
                d2 = buckets.get(sBucket, {}) if buckets else {}
                aLines = [f"{sRid} {iCount}" for sRid, iCount in sorted(d2.items())]
                ws.cell(row=iRow, column=iCol, value="\n".join(aLines))
            iRow += 1
        cls.applyFormatting(ws, iHeaderRow=iHeaderRow, iDataStart=iDataStart, bSingleLineRows=True)
        ws.freeze_panes = ws.cell(row=iDataStart, column=1).coordinate
        cls.addNamedRangeForCell(wb, ws, ws.cell(row=iHeaderRow, column=1), "Title")

    @classmethod
    def writeRollupSheet(cls, wb, lPageBuckets, lPageNames, lOrderedCrits, dCapturedRemarks):
        """
        Write Sheet 1 (the Accessibility Conformance Report rollup).

        Columns:
          1. Criterion    HYPERLINK; format "1.1.1 Non-text Content (A)"
          2. Summary      single-sentence summary
          3. Conformance  multi-line:
                          line 1: VPAT 2.5 verdict (Supports/Partially
                          Supports/Does Not Support/Not Evaluated).
                          When the verdict is not "Supports",
                          subsequent lines list the per-page sheet
                          names where this criterion was not supported
                          ("Not supported:" header) and where it could
                          not be evaluated ("Not evaluated:" header).
          4. Manual       numbered manual-test steps
          5. Result       user-editable; blank by default
          6. Remarks      user-editable; first line auto-generated
                          "Axe: fail N, pass N, incomplete N, inapplicable N"

        lPageNames is parallel to lPageBuckets and contains the
        sheet names (truncated and disambiguated to <=31 chars) that
        the per-URL sheets will be given. Used to build the page
        lists in the Conformance cell.
        """
        sName = "Conformance Report"
        if wb.worksheets and wb.active.title in ("Sheet", "Sheet1"):
            ws = wb.active
            ws.title = sName
        else:
            ws = wb.create_sheet(title=sName, index=0)
        iHeaderRow = 1
        iDataStart = 2
        aHeaders = ["Criterion", "Summary", "Conformance",
                    "Manual", "Result", "Remarks"]
        for iCol, sH in enumerate(aHeaders, 1):
            ws.cell(row=iHeaderRow, column=iCol, value=sH)
        iRow = iDataStart
        for sCrit in lOrderedCrits:
            d = dWcag22[sCrit]
            sCritText = f"{sCrit} {d['name']} ({d['level']})"
            sQuickRef = sWcagQuickRefBase + "#" + d["fragment"]
            cell = ws.cell(row=iRow, column=1)
            cell.value = f'=HYPERLINK("{sQuickRef}","{sCritText}")'
            ws.cell(row=iRow, column=2, value=dWcag22Summary.get(sCrit, ""))

            sCalc = cls.rollupVerdict(lPageBuckets, sCrit)
            sConformance = cls.calcToConformance(sCalc)
            # Build the multi-line Conformance cell. Line 1 is the
            # verdict; if any pages have Calc fail/partial they are
            # listed under "Not supported:"; if any have Calc manual
            # they are listed under "Not evaluated:". Pages with pass
            # / na / unknown are not listed (not actionable).
            dPages = cls.pagesByVerdict(lPageBuckets, lPageNames, sCrit)
            aLines = [sConformance]
            aNotSupported = sorted(set(dPages["fail"] + dPages["partial"]))
            aNotEvaluated = dPages["manual"]
            if aNotSupported:
                aLines.append("Not supported:")
                aLines.extend(aNotSupported)
            if aNotEvaluated:
                aLines.append("Not evaluated:")
                aLines.extend(aNotEvaluated)
            ws.cell(row=iRow, column=3, value="\n".join(aLines))

            sChecks = "\n".join(f"{i}. {sStep}" for i, sStep in enumerate(d["checks"], 1))
            ws.cell(row=iRow, column=4, value=sChecks)
            # Restore captured user edits if present
            sResult, sRemarks = dCapturedRemarks.get(sCrit, ("", ""))
            ws.cell(row=iRow, column=5, value=sResult)
            # Remarks first line auto-generated: instance counts in
            # axe-context labeling. Append captured user remarks below.
            iFail = cls.aggregateInstances(lPageBuckets, sCrit, "fail")
            iPass = cls.aggregateInstances(lPageBuckets, sCrit, "pass")
            iInc = cls.aggregateInstances(lPageBuckets, sCrit, "incomplete")
            iNa = cls.aggregateInstances(lPageBuckets, sCrit, "na")
            sAuto = f"Axe: fail {iFail}, pass {iPass}, incomplete {iInc}, inapplicable {iNa}"
            sRemarksCell = sAuto
            if sRemarks:
                sExisting = sRemarks
                m = re.match(r'^Axe: fail \d+, pass \d+, incomplete \d+, (?:in|not )?applicable \d+\n?', sExisting)
                if not m: m = re.match(r'^fail \d+, pass \d+, incomplete \d+, (?:in|not )?applicable \d+\n?', sExisting)
                if m: sExisting = sExisting[m.end():]
                if sExisting.strip():
                    sRemarksCell = sAuto + "\n" + sExisting
            ws.cell(row=iRow, column=6, value=sRemarksCell)
            iRow += 1
        cls.applyFormatting(ws, iHeaderRow=iHeaderRow, iDataStart=iDataStart, bSingleLineRows=True)
        cls.addNamedRangeForCell(wb, ws, ws.cell(row=iHeaderRow, column=1), "Title")
        return ws

    @classmethod
    def writeDocx(cls, pathParent, lFolders, lPageBuckets, lPageNames, lOrderedCrits, dtRunStart=None):
        """
        Generate ACR.docx, a narrative companion to ACR.xlsx.

        The DOCX is informed by report.htm's information architecture
        but adapted for ACR purposes: per-criterion focus rather than
        per-rule focus. Sections:

          1. Title and metadata (date, run start time, axe version,
             pages count, aggregate accessibility failure rate)
          2. Overview
          3. Pages Analyzed (with per-page accessibility failure rate)
          4. Conformance Summary (counts by verdict)
          5. Criteria Requiring Attention (every non-Supports row)
          6. Methodology
          7. Resources

        Uses python-docx (pure-Python, bundles cleanly with PyInstaller).
        Word is NOT required to be installed. Generated fresh on every
        run; user edits to a prior DOCX are not preserved.

        dtRunStart is the wall-clock run-start datetime; if None the
        current time is used. Falling back to current time means an
        ACR regenerated from a prior session won't have the original
        scan time, just when the rebuild ran -- but that's the
        honest answer in the regen case.
        """
        try:
            import docx
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.info("ACR: python-docx not available; skipping ACR.docx")
            return

        # Extract the axe-core version, browser version, and OS
        # version from per-page metadata. Each results.json
        # includes axeSource (e.g. "axe-core 4.10.3"), browserVersion
        # (Edge build number), and osVersion (e.g. "Windows 11"). We
        # take the most-recent values; if the user scanned across
        # different machines or browser updates mid-session, the
        # values reported reflect the LAST page scanned. If no
        # pages have a particular field yet, fall back to safe
        # defaults.
        sAxeVersion = ""
        sBrowserVersion = ""
        sOsVersion = ""
        for path, dResults, dMetadata in lFolders:
            sAxe = str(dMetadata.get("axeSource") or "")
            if sAxe: sAxeVersion = sAxe
            sBrV = str(dMetadata.get("browserVersion") or "")
            if sBrV: sBrowserVersion = sBrV
            sOsV = str(dMetadata.get("osVersion") or "")
            if sOsV: sOsVersion = sOsV
        if not sAxeVersion: sAxeVersion = f"axe-core {sFallbackAxeVersion}"
        if not sOsVersion: sOsVersion = getOsVersion()

        # Compute aggregate accessibility failure rate across all
        # included pages. Sum-of-numerators / sum-of-denominators is
        # mathematically equivalent to a size-weighted mean of per-
        # page rates: bigger pages get proportionally more weight.
        iTotalNumer = 0
        iTotalBytes = 0
        lPerPageRates = []  # parallel to lFolders, holds (sName, nRate)
        for (path, dResults, dMetadata), sSheetName in zip(lFolders, lPageNames):
            iN = computePageImpactNumerator(dResults)
            iB = computePageBytes(path)
            iTotalNumer += iN
            iTotalBytes += iB
            lPerPageRates.append((sSheetName,
                                   computeAccessibilityFailureRate(iN, iB)))
        nAggregateRate = computeAccessibilityFailureRate(iTotalNumer, iTotalBytes)

        d = docx.Document()
        # Tighten default style: smaller font, single-spacing
        styleNormal = d.styles["Normal"]
        styleNormal.font.name = "Calibri"
        styleNormal.font.size = Pt(11)

        # ---- Section 1: Title and metadata ----
        h = d.add_heading(level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.add_run("Accessibility Conformance Report")
        sStart = (dtRunStart or datetime.datetime.now()).strftime("%Y-%m-%d %H:%M:%S")
        d.add_paragraph(
            f"Generated by {sProgramName} {sProgramVersion}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.add_paragraph(
            f"Run started: {sStart}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.add_paragraph(
            f"Testing engine: {sAxeVersion}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sBrowserVersion:
            d.add_paragraph(
                f"Browser: Microsoft Edge {sBrowserVersion}"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sOsVersion:
            d.add_paragraph(
                f"Operating system: {sOsVersion}"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.add_paragraph(
            f"Pages analyzed: {len(lFolders)}"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lFolders and iTotalBytes > 0:
            d.add_paragraph(
                f"Aggregate accessibility failure rate: {nAggregateRate:.1f}%"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- Section 2: Overview ----
        d.add_heading("Overview", level=1)
        if lFolders:
            d.add_paragraph(
                f"This report summarizes accessibility conformance of "
                f"{len(lFolders)} web {('page' if len(lFolders) == 1 else 'pages')} "
                f"to WCAG 2.2, based on automated testing with axe-core and "
                f"manual-test guidance. The report follows the conventions "
                f"of the ITI Voluntary Product Accessibility Template (VPAT) "
                f"version 2.5, where each WCAG success criterion is assigned "
                f"one of: Supports, Partially Supports, Does Not Support, "
                f"Not Applicable, or Not Evaluated."
            )
            d.add_paragraph(
                f"This d is a draft generated by automation. It is "
                f"intended as a starting point for the human reviewer, who "
                f"is expected to verify each criterion's verdict (especially "
                f"those marked Not Evaluated), perform the manual checks "
                f"recommended in ACR.xlsx, and produce a final published "
                f"ACR in a separate working folder."
            )
        else:
            d.add_paragraph(
                f"No pages have been scanned yet. ACR.xlsx contains the "
                f"manual-test instructions for all 86 WCAG 2.2 success "
                f"criteria, ready for the user to fill in once pages are "
                f"scanned."
            )

        # ---- Section 3: Pages Analyzed ----
        if lFolders:
            d.add_heading("Pages Analyzed", level=1)
            d.add_paragraph(
                f"The following {len(lFolders)} {('page was' if len(lFolders) == 1 else 'pages were')} "
                f"included in this report. Each page has its own diagnostic "
                f"sheet in ACR.xlsx; the sheet name is a sanitized form of "
                f"the title shortened to fit Excel's 31-character tab limit. "
                f"The accessibility failure rate after each entry is the "
                f"page's impact-weighted violation density (lower is better; "
                f"see the Methodology section for how it is computed)."
            )
            dRateBySheet = dict(lPerPageRates)
            for (path, dResults, dMetadata), sSheetName in zip(lFolders, lPageNames):
                sPageTitle = str(dMetadata.get("pageTitle") or "(untitled)")
                sPageUrl = str(dMetadata.get("pageUrl") or "")
                nRate = dRateBySheet.get(sSheetName, 0.0)
                p = d.add_paragraph(style="List Bullet")
                p.add_run(f"{sPageTitle}").bold = True
                if sPageUrl:
                    p.add_run(f" ({sPageUrl})")
                p.add_run(f" — rate: {nRate:.1f}%")

        # ---- Section 4: Conformance Summary ----
        d.add_heading("Conformance Summary", level=1)
        # Tally each criterion's rollup verdict
        dCounts = {"Supports": 0, "Partially Supports": 0,
                   "Does Not Support": 0, "Not Evaluated": 0}
        for sCrit in lOrderedCrits:
            sCalc = cls.rollupVerdict(lPageBuckets, sCrit)
            sConf = cls.calcToConformance(sCalc)
            if sConf in dCounts: dCounts[sConf] += 1
            else: dCounts[sConf] = 1
        d.add_paragraph(
            f"Of the {len(lOrderedCrits)} WCAG 2.2 success criteria evaluated:"
        )
        for sTerm in ("Supports", "Partially Supports",
                      "Does Not Support", "Not Evaluated"):
            iN = dCounts.get(sTerm, 0)
            p = d.add_paragraph(style="List Bullet")
            p.add_run(f"{sTerm}: ").bold = True
            p.add_run(f"{iN} {('criterion' if iN == 1 else 'criteria')}")

        # ---- Section 5: Criteria Requiring Attention ----
        # Every criterion whose Conformance is not "Supports".
        d.add_heading("Criteria Requiring Attention", level=1)
        aAttention = []
        for sCrit in lOrderedCrits:
            sCalc = cls.rollupVerdict(lPageBuckets, sCrit)
            sConf = cls.calcToConformance(sCalc)
            if sConf != "Supports":
                aAttention.append((sCrit, sConf, sCalc))
        if not aAttention:
            d.add_paragraph(
                "All evaluated criteria are marked Supports. Note that some "
                "criteria may be marked Supports because the criterion is "
                "satisfied vacuously when no relevant content exists on the "
                "page. See per-criterion details in ACR.xlsx."
            )
        else:
            d.add_paragraph(
                f"{len(aAttention)} {'criterion needs' if len(aAttention) == 1 else 'criteria need'} "
                f"further attention. For each, see the corresponding row in "
                f"ACR.xlsx for the full Calc precedence, contributing axe "
                f"rules, manual-test steps, and per-page diagnostic results."
            )
            for sCrit, sConf, sCalc in aAttention:
                dC = dWcag22[sCrit]
                p = d.add_paragraph()
                p.add_run(f"{sCrit} {dC['name']} ({dC['level']}): ").bold = True
                p.add_run(f"{sConf}.")
                # Add the page lists
                dPages = cls.pagesByVerdict(lPageBuckets, lPageNames, sCrit)
                aNotSupp = sorted(set(dPages["fail"] + dPages["partial"]))
                aNotEval = dPages["manual"]
                if aNotSupp:
                    pp = d.add_paragraph()
                    pp.add_run("Not supported: ").italic = True
                    pp.add_run(", ".join(aNotSupp))
                if aNotEval:
                    pp = d.add_paragraph()
                    pp.add_run("Not evaluated: ").italic = True
                    pp.add_run(", ".join(aNotEval))

        # ---- Section 6: Methodology ----
        d.add_heading("Methodology", level=1)
        d.add_paragraph(
            f"{sProgramName} runs axe-core in Microsoft Edge against each "
            f"target page, then aggregates the raw results into ACR.xlsx "
            f"using a layered Calc precedence:"
        )
        for sLine in [
            "partial: at least one rule instance fails AND at least one passes for the same criterion",
            "fail: at least one rule instance fails (no pass)",
            "manual: at least one incomplete result (no fail or pass)",
            "pass: at least one pass (no fail or incomplete)",
            "na: only inapplicable rules touched the criterion",
            "unknown: no axe rules apply to the criterion",
        ]:
            p = d.add_paragraph(style="List Bullet")
            p.add_run(sLine)
        d.add_paragraph(
            "Calc values are then mapped to ACR conformance terms: pass and "
            "na become Supports (the criterion is satisfied, vacuously when "
            "no relevant content exists, per WCAG 2.0 Understanding "
            "Conformance); partial becomes Partially Supports; fail becomes "
            "Does Not Support; manual and unknown become Not Evaluated. The "
            "VPAT 2.5 standard reserves Not Evaluated for AAA criteria; "
            f"{sProgramName} uses it more broadly because automated testing "
            f"cannot decide outside its scope, and the user is expected to "
            f"complete manual testing before publishing the ACR."
        )
        d.add_paragraph(
            "Counts in the Remarks column are by node instance: one DOM "
            "element flagged by an axe rule equals one instance. So three "
            "image elements lacking alt text produce 'fail 3' for criterion "
            "1.1.1, even though all three failures come from the same "
            "image-alt rule."
        )
        d.add_heading("Accessibility failure rate", level=2)
        d.add_paragraph(
            f"For each scanned page, {sProgramName} computes an "
            f"impact-weighted violation density called the accessibility "
            f"failure rate. The numerator weights each violation instance "
            f"by impact level (minor 1, moderate 2, serious 3, critical 4) "
            f"and sums across all rules. The denominator is the byte size "
            f"of the saved page source (page.htm). The result is multiplied "
            f"by a tuning constant ({iAccessibilityRateScale}) so the "
            f"output reads naturally as a percent: typical pages land below "
            f"100%, with most well under that, while a heavily problematic "
            f"page can exceed 100%. The percent framing is purely for "
            f"display readability; the underlying quantity is "
            f"impact-weighted violation instances per byte of page source."
        )
        d.add_paragraph(
            f"For a page set, the aggregate rate is computed by summing "
            f"per-page numerators and per-page denominators before the "
            f"final division. This is a size-weighted view of the page "
            f"set: bigger pages contribute proportionally more, which "
            f"correctly reflects that they have more content with more "
            f"chances for users to encounter violations."
        )
        d.add_paragraph(
            f"Lower is better. The metric is intended to be tracked over "
            f"time: as the page owner remediates accessibility issues, "
            f"the rate should drop from one scan session to the next. "
            f"Accessibility is a journey, not a destination."
        )

        # ---- Section 7: Resources ----
        d.add_heading("Resources", level=1)
        for sLabel, sUrl in [
            ("WCAG 2.2 Quick Reference", "https://www.w3.org/WAI/WCAG22/quickref/"),
            ("Understanding WCAG 2.2", "https://www.w3.org/WAI/WCAG22/Understanding/"),
            ("Deque University rule reference", "https://dequeuniversity.com/rules/axe/"),
            ("WebAIM", "https://webaim.org/"),
            ("ITI VPAT", "https://www.itic.org/policy/accessibility/vpat"),
            ("Section 508", "https://www.section508.gov/"),
        ]:
            p = d.add_paragraph(style="List Bullet")
            p.add_run(f"{sLabel}: ").bold = True
            p.add_run(sUrl)

        # Save. Same case-enforcement as the XLSX.
        pathDocx = pathParent / sAcrDocxName
        for child in pathParent.iterdir():
            if child.is_file() and child.name.lower() == sAcrDocxName.lower() and child.name != sAcrDocxName:
                try: child.unlink()
                except Exception as ex: logger.info(f"ACR: cannot remove case-variant {child.name}: {ex}")
        d.save(str(pathDocx))
        logger.info(f"ACR: wrote {pathDocx}")

    @classmethod
    def buildIfApplicable(cls, arguments, lThisRunPaths, dtRunStart=None):
        """
        Driver: build ACR.xlsx in the parent output folder if there is
        anything to report. Called at end of session from main().

        dtRunStart is the wall-clock time the run was launched
        (datetime). It's surfaced in ACR.docx metadata header so
        users can correlate a draft ACR with a specific scan session.
        """
        try:
            pathParent = pathlib.Path(arguments.sOutputDir).resolve() if arguments.sOutputDir else pathlib.Path.cwd()
            pathParent.mkdir(parents=True, exist_ok=True)
            pathWorkbook = pathParent / sAcrWorkbookName
            bForceMode = bool(getattr(arguments, "bForce", False))
            lFolders = cls.discoverPageFolders(pathParent, lThisRunPaths, bForceMode)
            # Capture user edits before regeneration (only meaningful
            # in append mode; --force is the explicit "start fresh"
            # gesture and discards prior edits)
            dCaptured = cls.captureExistingRemarks(pathWorkbook) if not bForceMode else {}
            # Compute per-page buckets (empty when no pages)
            lPageBuckets = []
            for path, dResults, dMetadata in lFolders:
                lPageBuckets.append(cls.perPageBuckets(dResults))
            # Order criteria by numeric sort
            lOrderedCrits = sorted(dWcag22.keys(), key=cls.fnWcagSortKey)
            # Build workbook. Even with zero pages, Sheet 1
            # (Conformance Report) is written with all 86 criteria
            # showing "Not Evaluated" so the user has manual-test
            # checks to start with, and the Glossary follows. This
            # makes ACR.xlsx useful as a blank starting template even
            # before any page has been scanned.
            wb = openpyxl.Workbook()
            # Pre-compute the per-URL sheet names so the rollup's
            # Conformance cell can reference them in its "Not
            # supported" / "Not evaluated" page lists. The names are
            # the same ones the per-URL sheets will get (truncated and
            # disambiguated to fit Excel's 31-char limit), so the user
            # can match a name in a Conformance cell to a sheet tab.
            lUsedNames = ["Conformance Report"]
            lPageNames = []
            for path, dResults, dMetadata in lFolders:
                sName = cls.safeSheetName(path.name, lUsedNames)
                lUsedNames.append(sName)
                lPageNames.append(sName)
            cls.writeRollupSheet(wb, lPageBuckets, lPageNames, lOrderedCrits, dCaptured)
            # Per-URL sheets, in the same order as lPageNames
            for ((path, dResults, dMetadata), dBuckets, sName) in zip(lFolders, lPageBuckets, lPageNames):
                cls.writePerUrlSheet(wb, sName, dBuckets, dMetadata,
                                     "", lOrderedCrits)
            # Glossary
            cls.writeGlossarySheet(wb)
            # Delete any prior workbook with a case-variant name on
            # Windows. The file system is case-insensitive but case-
            # preserving: if "acr.xlsx" already exists from a prior
            # run, wb.save("ACR.xlsx") writes to the existing file and
            # keeps the lowercase casing. Removing it first lets the
            # new write create the file with the desired casing.
            for child in pathParent.iterdir():
                if child.is_file() and child.name.lower() == sAcrWorkbookName.lower() and child.name != sAcrWorkbookName:
                    try: child.unlink()
                    except Exception as ex: logger.info(f"ACR: cannot remove case-variant {child.name}: {ex}")
            wb.save(str(pathWorkbook))
            logger.info(f"ACR: wrote {pathWorkbook} ({len(lFolders)} pages, "
                f"{len(lOrderedCrits)} criteria)")
            # Generate the companion narrative DOCX. Same scope rules
            # as the XLSX (this-session vs walk-parent depending on
            # --force). Same case-enforcement strategy.
            try:
                cls.writeDocx(pathParent, lFolders, lPageBuckets, lPageNames, lOrderedCrits, dtRunStart=dtRunStart)
            except Exception as ex:
                logger.info(f"ACR: DOCX generation failed: {ex}")
                try: logger.info(traceback.format_exc())
                except Exception: pass
        except Exception as ex:
            logger.info(f"ACR: failed to write {sAcrWorkbookName}: {ex}")
            try: logger.info(traceback.format_exc())
            except Exception: pass


# --- GUI dialog (Python.NET / WinForms) ---
#
# Uses pythonnet's `clr` to load the .NET Framework 4.8 System.Windows.Forms
# assembly that ships with every supported version of Windows 10 and 11. This
# gives us the *same* widget set 2htm uses (System.Windows.Forms.Form,
# Label, TextBox, CheckBox, Button, FolderBrowserDialog, OpenFileDialog,
# MessageBox), so screen-reader behavior in JAWS and NVDA is identical
# between the two tools: labels announced via Label-to-control association,
# ampersand access keys, AcceptButton/CancelButton wiring, predictable tab
# order. No extra runtime to install -- the .NET Framework 4.8 is built
# into Windows 10 (since 1903) and Windows 11.
#
# pythonnet (the `clr` module) is imported lazily, so the CLI path does not
# require it to be installed at all. PyInstaller bundles pythonnet's tiny
# loader (clr.pyd, Python.Runtime.dll) when the build script runs; at
# runtime that loader bridges into the in-box .NET Framework.

def _loadDotNetForms():
    """
    Import pythonnet and the System.Windows.Forms / System.Drawing namespaces.
    Returns a dict of the names this file uses, or None on failure (e.g.
    pythonnet not installed). Keeps the import-mess in one place so the
    dialog code below reads cleanly.
    """
    try:
        import clr  # provided by the pythonnet package
        clr.AddReference("System.Windows.Forms")
        clr.AddReference("System.Drawing")
        from System import EventHandler, Environment as DotNetEnvironment
        from System.Drawing import ContentAlignment, Font, FontFamily, Point, Size, SystemFonts
        from System.Threading import ApartmentState, Thread
        from System.Windows.Forms import (
            AnchorStyles, Application, Button, CheckBox, DialogResult,
            DockStyle, FolderBrowserDialog, Form, FormBorderStyle,
            FormStartPosition, Keys, Label, MessageBox, MessageBoxButtons,
            MessageBoxDefaultButton, MessageBoxIcon, OpenFileDialog, Panel,
            ScrollBars, TextBox)
        return {
            "AnchorStyles": AnchorStyles, "ApartmentState": ApartmentState,
            "Application": Application,
            "Button": Button, "CheckBox": CheckBox,
            "ContentAlignment": ContentAlignment, "DialogResult": DialogResult,
            "DockStyle": DockStyle, "DotNetEnvironment": DotNetEnvironment,
            "EventHandler": EventHandler,
            "FolderBrowserDialog": FolderBrowserDialog,
            "Font": Font, "FontFamily": FontFamily, "Form": Form,
            "FormBorderStyle": FormBorderStyle,
            "FormStartPosition": FormStartPosition,
            "Keys": Keys, "Label": Label, "MessageBox": MessageBox,
            "MessageBoxButtons": MessageBoxButtons,
            "MessageBoxDefaultButton": MessageBoxDefaultButton,
            "MessageBoxIcon": MessageBoxIcon, "OpenFileDialog": OpenFileDialog,
            "Panel": Panel, "Point": Point, "ScrollBars": ScrollBars,
            "Size": Size, "SystemFonts": SystemFonts,
            "TextBox": TextBox, "Thread": Thread,
        }
    except Exception as ex:
        print(f"[ERROR] GUI mode requires pythonnet (the `clr` module) and the .NET Framework: {ex}")
        print("        pip install pythonnet")
        return None


def browseForFolderViaShell(sTitle, sInitialPath=""):
    """
    Native folder picker via SHBrowseForFolderW + SHGetPathFromIDListW from
    shell32.dll. Used in place of WinForms' FolderBrowserDialog because the
    latter deadlocks under pythonnet (issue #657).

    We deliberately use the OLDER ("classic") browse-for-folder dialog by
    setting only BIF_RETURNONLYFSDIRS in ulFlags. The newer dialog style
    (BIF_NEWDIALOGSTYLE / BIF_USENEWUI) requires OleInitialize and a
    single-threaded apartment (COINIT_APARTMENTTHREADED); per the Microsoft
    documentation, "If COM is initialized through CoInitializeEx with the
    COINIT_MULTITHREADED flag set, SHBrowseForFolder fails if
    BIF_NEWDIALOGSTYLE is passed." Pythonnet's COM initialization state can
    leave the thread in MTA, so the new dialog style hangs. The classic
    dialog has no such requirement and works reliably.

    The trade-off is cosmetic: the classic dialog has no resize handle, no
    edit field for typing a path, and no "create new folder" button. It
    does have a complete file system tree, which is enough to choose an
    output directory.

    sInitialPath: if non-empty and an existing directory, the dialog opens
    with that folder pre-selected and expanded in the tree. We honor this
    by registering a small BFFCALLBACK that posts BFFM_SETSELECTION on
    BFFM_INITIALIZED. The callback is a plain Win32 function pointer and
    introduces no COM dependencies.

    Returns the chosen folder path as a string, or "" if the user cancelled.
    sTitle is shown above the folder tree.
    """
    import ctypes.wintypes as wt
    iBifReturnOnlyFsDirs = 0x00000001
    iBffmInitialized = 1
    iBffmSetSelection = 0x400 + 103  # WM_USER + 103 == BFFM_SETSELECTIONW
    iMaxPath = 260

    class BrowseInfoW(ctypes.Structure):
        _fields_ = [
            ("hwndOwner", wt.HWND),
            ("pidlRoot", ctypes.c_void_p),
            ("pszDisplayName", wt.LPWSTR),
            ("lpszTitle", wt.LPCWSTR),
            ("ulFlags", wt.UINT),
            ("lpfn", ctypes.c_void_p),
            ("lParam", wt.LPARAM),
            ("iImage", ctypes.c_int),
        ]

    shell32 = ctypes.windll.shell32
    user32 = ctypes.windll.user32
    ole32 = ctypes.windll.ole32

    # Display-name buffer required by the API; we discard it and use
    # SHGetPathFromIDListW for the actual file-system path.
    bufDisplay = ctypes.create_unicode_buffer(iMaxPath)

    # Build the BFFCALLBACK that selects the initial path when the dialog
    # finishes initializing. SendMessage with BFFM_SETSELECTIONW takes a
    # wide-string path in lParam (and TRUE in wParam to indicate it's a
    # path string, not a PIDL).
    user32.SendMessageW.restype = ctypes.c_long
    user32.SendMessageW.argtypes = [wt.HWND, ctypes.c_uint, ctypes.c_void_p, ctypes.c_void_p]

    # Validate and normalize the seed path. If it's empty or doesn't exist,
    # we don't install a callback at all -- the dialog opens at its default
    # (a file-system root view).
    sNormalizedInitial = ""
    if sInitialPath:
        try:
            pathSeed = pathlib.Path(sInitialPath).expanduser().resolve()
            if pathSeed.is_dir():
                sNormalizedInitial = str(pathSeed)
        except Exception:
            sNormalizedInitial = ""

    pCallback = None
    callbackProto = None
    if sNormalizedInitial:
        # Keep the path string alive for the lifetime of the dialog by
        # closing over it in the callback. Without this, ctypes would
        # garbage-collect the wide-string before the dialog reads it.
        sCallbackPath = sNormalizedInitial
        # BFFCALLBACK signature: int CALLBACK Fn(HWND, UINT, LPARAM, LPARAM).
        # We use ctypes.WINFUNCTYPE for the stdcall convention.
        callbackProto = ctypes.WINFUNCTYPE(
            ctypes.c_int, wt.HWND, ctypes.c_uint, wt.LPARAM, wt.LPARAM)
        def fnBrowseCallback(hwnd, iMsg, lParam, lpData):
            if iMsg == iBffmInitialized:
                try:
                    user32.SendMessageW(hwnd, iBffmSetSelection,
                        ctypes.c_void_p(1),  # TRUE: lParam is a string path
                        ctypes.c_wchar_p(sCallbackPath))
                except Exception:
                    pass
            return 0
        pCallback = callbackProto(fnBrowseCallback)

    bi = BrowseInfoW()
    bi.hwndOwner = user32.GetActiveWindow()  # may be NULL; that's OK
    bi.pidlRoot = None
    bi.pszDisplayName = ctypes.cast(bufDisplay, wt.LPWSTR)
    bi.lpszTitle = sTitle
    bi.ulFlags = iBifReturnOnlyFsDirs  # classic dialog only
    bi.lpfn = ctypes.cast(pCallback, ctypes.c_void_p) if pCallback else None
    bi.lParam = 0
    bi.iImage = 0

    # Initialize COM as STA. CoInitialize is shorthand for
    # CoInitializeEx(NULL, COINIT_APARTMENTTHREADED). It returns S_FALSE if
    # the thread already had a compatible apartment; we treat that as fine.
    # If the thread was previously in MTA, this call will return RPC_E_CHANGED_MODE
    # and the classic dialog will still work because it does not require STA.
    shell32.SHBrowseForFolderW.restype = ctypes.c_void_p
    shell32.SHBrowseForFolderW.argtypes = [ctypes.POINTER(BrowseInfoW)]
    shell32.SHGetPathFromIDListW.restype = wt.BOOL
    shell32.SHGetPathFromIDListW.argtypes = [ctypes.c_void_p, wt.LPWSTR]
    ole32.CoTaskMemFree.argtypes = [ctypes.c_void_p]
    try:
        ole32.CoInitialize(None)
    except Exception: pass

    pidl = shell32.SHBrowseForFolderW(ctypes.byref(bi))
    # Keep callbackProto and pCallback alive until SHBrowseForFolderW returns.
    # This line is a no-op at runtime but documents the lifetime requirement
    # so a future reader doesn't accidentally drop the reference.
    _ = (pCallback, callbackProto)
    if not pidl: return ""
    try:
        bufPath = ctypes.create_unicode_buffer(iMaxPath)
        if not shell32.SHGetPathFromIDListW(pidl, bufPath):
            return ""
        return bufPath.value
    finally:
        try: ole32.CoTaskMemFree(pidl)
        except Exception: pass


def showGuiDialog(arguments):
    """
    Show the urlCheck parameter dialog. Mutates `arguments` in place with the
    user's chosen values. Returns True on OK, False on Cancel.

    Layout mirrors 2htm:
      Row 1: "Source urls:"                  [textbox]   [Browse source]
      Row 2: "Output directory:"             [textbox]   [Choose output]
      Row 3: [x] Invisible mode             [x] View output
      Row 4: [x] Log session                [x] Use configuration
      Row 5: [Help] [Default settings]                  [OK] [Cancel]
    """
    d = _loadDotNetForms()
    logger.info(f"showGuiDialog: _loadDotNetForms returned: {'ok' if d is not None else 'NONE (pythonnet missing?)'}")
    if d is None: return False

    # Pull names out of the namespace dict for readability.
    Application = d["Application"]
    Button = d["Button"]
    CheckBox = d["CheckBox"]
    ContentAlignment = d["ContentAlignment"]
    DialogResult = d["DialogResult"]
    EventHandler = d["EventHandler"]
    FolderBrowserDialog = d["FolderBrowserDialog"]
    Form = d["Form"]
    FormBorderStyle = d["FormBorderStyle"]
    FormStartPosition = d["FormStartPosition"]
    Keys = d["Keys"]
    Label = d["Label"]
    MessageBox = d["MessageBox"]
    MessageBoxButtons = d["MessageBoxButtons"]
    MessageBoxDefaultButton = d["MessageBoxDefaultButton"]
    MessageBoxIcon = d["MessageBoxIcon"]
    OpenFileDialog = d["OpenFileDialog"]
    Point = d["Point"]
    Size = d["Size"]
    SystemFonts = d["SystemFonts"]
    TextBox = d["TextBox"]

    # Log .NET / pythonnet diagnostics so the GUI environment is
    # self-documenting in the log file.
    try:
        environment = d["DotNetEnvironment"]
        logger.info(f".NET runtime: Version={environment.Version} "
            f"OSVersion={environment.OSVersion} Is64Bit={environment.Is64BitProcess}")
    except Exception as ex:
        logger.info(f".NET diagnostics unavailable: {ex}")

    # Set the calling thread to the single-threaded apartment (STA) COM
    # model. WinForms common dialogs -- OpenFileDialog, FolderBrowserDialog,
    # and the various shell extensions they delegate to -- require an STA
    # thread. C# WinForms apps get this for free via [STAThread] on Main();
    # in pythonnet we have to set it explicitly. If this is omitted, the
    # main dialog usually opens fine, but clicking Browse source or Choose
    # output deadlocks the COM marshaler and the program appears to lock up.
    #
    # SetApartmentState fails with InvalidOperationException if the thread
    # has already started COM in MTA mode (e.g. by a prior dialog call in
    # the same process). The wrap-and-log pattern below makes that case
    # diagnosable in the log without crashing.
    try:
        thread = d["Thread"].CurrentThread
        sBefore = str(thread.GetApartmentState())
        thread.SetApartmentState(d["ApartmentState"].STA)
        sAfter = str(thread.GetApartmentState())
        logger.info(f"Thread apartment state: {sBefore} -> {sAfter}")
    except Exception as ex:
        logger.info(f"SetApartmentState(STA) failed (continuing): {ex}")

    # Enable modern Windows visual styles (Common Controls 6 themed widgets:
    # rounded buttons, themed scroll bars, etc.) and the GDI+ TextRenderer
    # for crisper text. Both calls MUST happen before any Form, Button, or
    # other control is constructed in this AppDomain. WinForms ignores them
    # otherwise. EnableVisualStyles is also harmless to call repeatedly, so
    # it is safe even if showGuiDialog is invoked more than once.
    #
    # Accessibility note: visual styles are purely cosmetic. The underlying
    # HWNDs and MSAA / UI Automation properties (control type, name, role,
    # state) are unchanged, so JAWS and NVDA see the exact same accessibility
    # tree they would see with the classic theme. If a regression is
    # observed, comment out these two lines and rebuild.
    try:
        d["Application"].EnableVisualStyles()
        d["Application"].SetCompatibleTextRenderingDefault(False)
    except Exception: pass

    # Initial values from arguments (which may have been pre-loaded from
    # saved config and/or the command line).
    sInitTarget = getattr(arguments, "sSource", "") or ""
    sInitOutDir = getattr(arguments, "sOutputDir", "") or ""
    bInitView = bool(getattr(arguments, "bViewOutput", False))
    bInitInvisible = bool(getattr(arguments, "bInvisible", False))
    bInitAuth = bool(getattr(arguments, "bAuthenticate", False))
    bInitMain = bool(getattr(arguments, "bMainProfile", False))
    bInitForce = bool(getattr(arguments, "bForce", False))
    bInitLog = bool(getattr(arguments, "bLog", False))
    bInitUseCfg = bool(getattr(arguments, "bUseConfig", False))

    # Layout constants live at module scope as iLayout* (see top of file).
    # Match extCheck and 2htm exactly so the three dialogs feel like
    # siblings.
    iFormW = iLayoutFormWidth
    iTextX = iLayoutLeft + iLayoutLabelWidth + iLayoutGap
    iTextW = iFormW - iTextX - iLayoutGap - iLayoutButtonWidth - iLayoutRight
    iBtnX = iFormW - iLayoutRight - iLayoutButtonWidth

    # Build the form.
    frm = Form()
    frm.Text = sProgramName
    frm.FormBorderStyle = FormBorderStyle.FixedDialog
    frm.StartPosition = FormStartPosition.CenterScreen
    frm.MaximizeBox = False
    frm.MinimizeBox = False
    frm.ShowInTaskbar = True
    frm.ClientSize = Size(iFormW, 280)
    frm.Font = SystemFonts.MessageBoxFont

    # F1 -> Help. KeyPreview lets the form see the keystroke before child
    # controls consume it. F1 is the standard Windows help shortcut and is
    # expected by keyboard-driven and screen-reader users.
    frm.KeyPreview = True

    # --- Row 1: Target ---
    y = iLayoutTop
    lblTarget = Label()
    lblTarget.Text = "&Source urls:"
    lblTarget.AutoSize = False
    lblTarget.Location = Point(iLayoutLeft, y + 3)
    lblTarget.Size = Size(iLayoutLabelWidth, iLayoutTextHeight)
    lblTarget.TextAlign = ContentAlignment.MiddleLeft
    frm.Controls.Add(lblTarget)

    txtTarget = TextBox()
    txtTarget.Text = sInitTarget
    txtTarget.Location = Point(iTextX, y)
    txtTarget.Size = Size(iTextW, iLayoutTextHeight)
    txtTarget.TabIndex = 0
    # Explicit AccessibleName so JAWS/NVDA announce the field by its
    # label even when the visual layout doesn't auto-associate.
    txtTarget.AccessibleName = "Source urls"
    frm.Controls.Add(txtTarget)

    btnBrowseTarget = Button()
    btnBrowseTarget.Text = "&Browse source..."
    btnBrowseTarget.Location = Point(iBtnX, y - 1)
    btnBrowseTarget.Size = Size(iLayoutButtonWidth, iLayoutButtonHeight)
    btnBrowseTarget.TabIndex = 1
    btnBrowseTarget.UseVisualStyleBackColor = True
    frm.Controls.Add(btnBrowseTarget)

    # --- Row 2: Output directory ---
    y += iLayoutTextHeight + iLayoutRowGap
    lblOut = Label()
    lblOut.Text = "&Output folder:"
    lblOut.AutoSize = False
    lblOut.Location = Point(iLayoutLeft, y + 3)
    lblOut.Size = Size(iLayoutLabelWidth, iLayoutTextHeight)
    lblOut.TextAlign = ContentAlignment.MiddleLeft
    frm.Controls.Add(lblOut)

    txtOut = TextBox()
    txtOut.Text = sInitOutDir
    txtOut.Location = Point(iTextX, y)
    txtOut.Size = Size(iTextW, iLayoutTextHeight)
    txtOut.TabIndex = 2
    txtOut.AccessibleName = "Output folder"
    frm.Controls.Add(txtOut)

    btnBrowseOut = Button()
    btnBrowseOut.Text = "&Choose output..."
    btnBrowseOut.Location = Point(iBtnX, y - 1)
    btnBrowseOut.Size = Size(iLayoutButtonWidth, iLayoutButtonHeight)
    btnBrowseOut.TabIndex = 3
    btnBrowseOut.UseVisualStyleBackColor = True
    frm.Controls.Add(btnBrowseOut)

    # --- Row 3: program-specific option row (urlCheck has Invisible) ---
    # The program-specific option appears alone above the common
    # option grid so the layout reads consistently with 2htm and
    # extCheck: program-specific first, common options after.
    #
    # Row order (top to bottom, left then right): Authenticate
    # credentials + Main profile (browser-session pair),
    # Invisible mode + Force replacements (run-mode pair),
    # View output + Log session (output pair), Use configuration
    # alone (settings persistence). Tab order matches reading order.
    y += iLayoutTextHeight + iLayoutRowGap * 2
    iChkW = (iFormW - iLayoutLeft - iLayoutRight) // 2

    # --- Row 3: Authenticate credentials + Main profile (browser-session pair) ---
    chkAuth = CheckBox()
    chkAuth.Text = "&Authenticate credentials"
    chkAuth.Checked = bInitAuth
    chkAuth.Location = Point(iLayoutLeft, y)
    chkAuth.Size = Size(iChkW, iLayoutTextHeight)
    chkAuth.TabIndex = 4
    frm.Controls.Add(chkAuth)

    chkMain = CheckBox()
    chkMain.Text = "&Main profile"
    chkMain.Checked = bInitMain
    chkMain.Location = Point(iLayoutLeft + iChkW, y)
    chkMain.Size = Size(iChkW, iLayoutTextHeight)
    chkMain.TabIndex = 5
    frm.Controls.Add(chkMain)

    # --- Row 4: Invisible mode + Force replacements (run-mode pair) ---
    y += iLayoutTextHeight + iLayoutRowGap
    chkInvisible = CheckBox()
    chkInvisible.Text = "&Invisible mode"
    # Both Invisible mode and Authenticate credentials are
    # toggleable independently. If both end up checked at OK time,
    # urlCheck treats --authenticate as overriding --invisible (an
    # auth prompt requires a visible browser); the override is
    # logged so the user sees the resolution. Matches CLI behavior.
    chkInvisible.Checked = bInitInvisible
    chkInvisible.Location = Point(iLayoutLeft, y)
    chkInvisible.Size = Size(iChkW, iLayoutTextHeight)
    chkInvisible.TabIndex = 6
    frm.Controls.Add(chkInvisible)

    chkForce = CheckBox()
    chkForce.Text = "&Force replacements"
    chkForce.Checked = bInitForce
    chkForce.Location = Point(iLayoutLeft + iChkW, y)
    chkForce.Size = Size(iChkW, iLayoutTextHeight)
    chkForce.TabIndex = 7
    frm.Controls.Add(chkForce)

    # --- Row 5: View output + Log session (output pair) ---
    y += iLayoutTextHeight + iLayoutRowGap
    chkView = CheckBox()
    chkView.Text = "&View output"
    chkView.Checked = bInitView
    chkView.Location = Point(iLayoutLeft, y)
    chkView.Size = Size(iChkW, iLayoutTextHeight)
    chkView.TabIndex = 8
    frm.Controls.Add(chkView)

    chkLog = CheckBox()
    chkLog.Text = "&Log session"
    chkLog.Checked = bInitLog
    chkLog.Location = Point(iLayoutLeft + iChkW, y)
    chkLog.Size = Size(iChkW, iLayoutTextHeight)
    chkLog.TabIndex = 9
    frm.Controls.Add(chkLog)

    # --- Row 6: Use configuration alone (settings persistence) ---
    y += iLayoutTextHeight + iLayoutRowGap
    chkUseCfg = CheckBox()
    chkUseCfg.Text = "&Use configuration"
    chkUseCfg.Checked = bInitUseCfg
    chkUseCfg.Location = Point(iLayoutLeft, y)
    chkUseCfg.Size = Size(iChkW, iLayoutTextHeight)
    chkUseCfg.TabIndex = 10
    frm.Controls.Add(chkUseCfg)

    # --- Bottom row: Help, Defaults on the left; OK, Cancel on the right ---
    y += iLayoutTextHeight + iLayoutRowGap * 2
    btnHelp = Button()
    btnHelp.Text = "&Help"
    btnHelp.Location = Point(iLayoutLeft, y)
    btnHelp.Size = Size(iLayoutButtonWidth, iLayoutButtonHeight)
    btnHelp.TabIndex = 11
    btnHelp.UseVisualStyleBackColor = True
    frm.Controls.Add(btnHelp)

    btnDefaults = Button()
    btnDefaults.Text = "&Default settings"
    btnDefaults.Location = Point(iLayoutLeft + iLayoutButtonWidth + iLayoutGap, y)
    btnDefaults.Size = Size(iLayoutButtonWidth, iLayoutButtonHeight)
    btnDefaults.TabIndex = 12
    btnDefaults.UseVisualStyleBackColor = True
    frm.Controls.Add(btnDefaults)

    btnOk = Button()
    btnOk.Text = "OK"
    btnOk.DialogResult = DialogResult.OK
    btnOk.Location = Point(iFormW - iLayoutRight - 2 * iLayoutButtonWidth - iLayoutGap, y)
    btnOk.Size = Size(iLayoutButtonWidth, iLayoutButtonHeight)
    btnOk.TabIndex = 13
    btnOk.UseVisualStyleBackColor = True
    frm.Controls.Add(btnOk)

    btnCancel = Button()
    btnCancel.Text = "Cancel"
    btnCancel.DialogResult = DialogResult.Cancel
    btnCancel.Location = Point(iBtnX, y)
    btnCancel.Size = Size(iLayoutButtonWidth, iLayoutButtonHeight)
    btnCancel.TabIndex = 14
    btnCancel.UseVisualStyleBackColor = True
    frm.Controls.Add(btnCancel)

    # Wire defaults: Enter -> OK, Esc -> Cancel.
    frm.AcceptButton = btnOk
    frm.CancelButton = btnCancel

    # Adjust form height to accommodate the last control plus a margin.
    frm.ClientSize = Size(iFormW, y + iLayoutButtonHeight + iLayoutTop)

    # --- Event handlers (defined as nested functions so they close over the
    #     control variables above) ---

    def fnShowHelp():
        sMsg = (
            f"{sProgramName} {sProgramVersion} checks one or more web pages "
            f"for accessibility problems and saves a set of output files in "
            f"a folder named after each page title.\r\n\r\n"
            f"Source files: enter one url (https://example.com), or a domain "
            f"(microsoft.com), or several of either separated by spaces, or "
            f"the path to a single plain text file that lists urls, domains, "
            f"or local file paths one per line. The list file may have any "
            f"extension; urlCheck verifies it is plain text by inspecting "
            f"its contents.\r\n\r\n"
            f"Output folder: parent folder under which the per-scan "
            f"folders are written. Blank means the current working "
            f"folder.\r\n\r\n"
            f"Options:\r\n"
            f"  Invisible mode - run Edge with no visible browser window. If both Invisible mode and Authenticate credentials are checked, Authenticate credentials wins (an auth prompt requires a visible browser); the override is logged.\r\n"
            f"  Authenticate credentials - when a url's domain is "
            f"encountered for the first time in this run, pause after "
            f"the page loads so the user can sign in / accept cookies / "
            f"dismiss popups, then press Enter (or click OK) to resume. "
            f"Overrides Invisible mode if both are set.\r\n"
            f"  Main profile - launch Edge with your real (default) "
            f"profile so saved logins, cookies, and session state are "
            f"available. Without it, urlCheck uses a fresh temporary "
            f"profile so the scan is anonymous. Requires that no "
            f"Microsoft Edge process is already running; if Edge is "
            f"running, the dialog shows a message asking you to close "
            f"Edge and submit again.\r\n"
            f"  Force replacements - reuse an existing per-page output "
            f"folder (emptying its contents and writing a fresh set of "
            f"files) instead of skipping the url\r\n"
            f"  View output - open the parent output folder in File "
            f"Explorer when all scans are done\r\n"
            f"  Log session - write urlCheck.log (replacing any prior log) "
            f"in the current working folder\r\n"
            f"  Use configuration - remember these settings for next time "
            f"in %LOCALAPPDATA%\\urlCheck\\urlCheck.ini\r\n\r\n"
            f"Press Cancel to exit without scanning.\r\n\r\n"
            f"Open the full README in your browser?")
        dialogResult = MessageBox.Show(sMsg, f"{sProgramName} - Help",
            MessageBoxButtons.YesNo, MessageBoxIcon.Information,
            MessageBoxDefaultButton.Button2)
        if dialogResult == DialogResult.Yes: launchReadMe()

    def fnOnKeyDown(sender, args):
        if args.KeyCode == Keys.F1:
            args.Handled = True
            args.SuppressKeyPress = True
            fnShowHelp()

    def fnPickFile(sender, args):
        # Pythonnet has a long-standing deadlock when it invokes the modern
        # (Vista+) IFileOpenDialog COM-shell file picker -- documented in
        # pythonnet issues #657 and #1286, both unresolved. The fix is to
        # set AutoUpgradeEnabled = False, which forces the legacy Win32
        # GetOpenFileName common dialog (comdlg32.dll). It has the older
        # Windows look but actually works under pythonnet.
        sCurrent = (txtTarget.Text or "").strip()
        sInitial = getInitialBrowseDir(sCurrent)
        logger.info(f"Browse source clicked; opening OpenFileDialog (legacy) at {sInitial!r}")
        try:
            dialog = OpenFileDialog()
            dialog.AutoUpgradeEnabled = False
            dialog.Title = "Choose a plain text url list"
            dialog.Filter = ("Plain text files (*.txt;*.lst;*.md)|*.txt;*.lst;*.md|"
                           "All files (*.*)|*.*")
            dialog.FilterIndex = 2  # default to All files; user may have any extension
            dialog.CheckFileExists = True
            dialog.RestoreDirectory = True
            try:
                dialog.InitialDirectory = sInitial
            except Exception: pass
            dialogResult = dialog.ShowDialog()
            logger.info(f"OpenFileDialog returned: {dialogResult}")
            if dialogResult == DialogResult.OK:
                txtTarget.Text = dialog.FileName
                logger.info(f"Source set to: {dialog.FileName}")
        except Exception as ex:
            logger.info(f"OpenFileDialog raised: {ex}")
            MessageBox.Show(f"Browse source failed: {ex}",
                f"{sProgramName} - Browse error",
                MessageBoxButtons.OK, MessageBoxIcon.Warning)

    def fnPickFolder(sender, args):
        # FolderBrowserDialog has no AutoUpgradeEnabled property, and the
        # default shell-COM machinery deadlocks under pythonnet exactly like
        # the modern OpenFileDialog (issue #657). We sidestep it by calling
        # the older Win32 SHBrowseForFolder API directly via ctypes -- which
        # bypasses pythonnet entirely for this dialog and uses the simpler
        # folder picker that doesn't need the same COM-marshaled callbacks.
        sCurrent = (txtOut.Text or "").strip()
        sInitial = getInitialBrowseDir(sCurrent)
        sChosen = ""
        logger.info(f"Choose output clicked; calling SHBrowseForFolder at {sInitial!r}")
        try:
            sChosen = browseForFolderViaShell(
                "Choose the parent folder under which the per-scan "
                "output folder will be created.",
                sInitial)
        except Exception as ex:
            logger.info(f"SHBrowseForFolder raised: {ex}")
            MessageBox.Show(f"Choose output failed: {ex}",
                f"{sProgramName} - Browse error",
                MessageBoxButtons.OK, MessageBoxIcon.Warning)
            return
        logger.info(f"SHBrowseForFolder returned: {sChosen!r}")
        if sChosen:
            txtOut.Text = sChosen
            logger.info(f"Output dir set to: {sChosen}")

    def fnDefaults(sender, args):
        txtTarget.Text = ""
        txtOut.Text = ""
        chkInvisible.Checked = False
        chkAuth.Checked = False
        chkMain.Checked = False
        chkForce.Checked = False
        chkView.Checked = False
        chkLog.Checked = False
        chkUseCfg.Checked = False
        configManager.eraseAll()

    def fnHelpClick(sender, args):
        fnShowHelp()

    def fnOkClick(sender, args):
        sCurrent = (txtTarget.Text or "").strip()
        if not sCurrent:
            MessageBox.Show(
                "Please enter one or more urls separated by spaces, "
                "or the path to a plain text file of urls.",
                f"{sProgramName} - Missing source",
                MessageBoxButtons.OK, MessageBoxIcon.Warning)
            txtTarget.Focus()
            frm.DialogResult = DialogResult.None_
            return
        sKind, vDetail = classifyInput(sCurrent)
        if sKind == "error":
            MessageBox.Show(str(vDetail),
                f"{sProgramName} - Invalid source",
                MessageBoxButtons.OK, MessageBoxIcon.Warning)
            txtTarget.Focus()
            frm.DialogResult = DialogResult.None_
            return
        # Output directory validation: if the user specified a directory
        # that does not exist, prompt to create it (default Yes). On No
        # or creation failure, keep the dialog open.
        sOutCandidate = (txtOut.Text or "").strip()
        if len(sOutCandidate) >= 2 and sOutCandidate[0] == '"' and sOutCandidate[-1] == '"':
            sOutCandidate = sOutCandidate[1:-1].strip()
        if sOutCandidate and not os.path.isdir(sOutCandidate):
            dr = MessageBox.Show(
                f"Create {sOutCandidate}?",
                sProgramName,
                MessageBoxButtons.YesNo, MessageBoxIcon.Question,
                MessageBoxDefaultButton.Button1)
            if dr != DialogResult.Yes:
                frm.DialogResult = DialogResult.None_
                txtOut.Focus()
                return
            try:
                os.makedirs(sOutCandidate, exist_ok=True)
            except Exception as ex:
                MessageBox.Show(
                    f"Could not create folder:\r\n{sOutCandidate}\r\n\r\n{ex}",
                    sProgramName,
                    MessageBoxButtons.OK, MessageBoxIcon.Warning)
                frm.DialogResult = DialogResult.None_
                txtOut.Focus()
                return
        # If Main profile is checked, refuse to proceed when another
        # Edge process is running. Edge cannot share its profile
        # directory across two processes; attempting to launch
        # against an in-use profile produces a confusing error
        # downstream. Better to surface this here, in the dialog,
        # so the user can close all Edge windows and re-submit the
        # dialog -- without restarting urlCheck. We capture the
        # control that had focus when OK was triggered (which may
        # have been the OK button, or any other dialog control,
        # since pressing Enter on most controls triggers the
        # AcceptButton) and restore focus there after the message
        # box is dismissed -- the user can press Enter again from
        # wherever they were once Edge is closed.
        if bool(chkMain.Checked) and isEdgeRunning():
            controlPriorFocus = None
            try: controlPriorFocus = frm.ActiveControl
            except Exception: pass
            MessageBox.Show(
                "Microsoft Edge is currently running. urlCheck "
                "cannot proceed because Main profile requires "
                "exclusive access to your Edge profile, which Edge "
                "does not share across processes. Please close all "
                "Edge windows (right-click the Edge taskbar icon "
                "and choose Close window, or quit Edge from its "
                "menu), then submit this dialog again to retry.",
                f"{sProgramName} - Edge is running",
                MessageBoxButtons.OK, MessageBoxIcon.Warning)
            frm.DialogResult = DialogResult.None_
            try:
                if controlPriorFocus is not None: controlPriorFocus.Focus()
            except Exception: pass
            return
        # sKind is 'urls' or 'listfile' -- both are valid; let the dialog close.

    btnBrowseTarget.Click += EventHandler(fnPickFile)
    btnBrowseOut.Click += EventHandler(fnPickFolder)
    btnDefaults.Click += EventHandler(fnDefaults)
    btnHelp.Click += EventHandler(fnHelpClick)
    btnOk.Click += EventHandler(fnOkClick)
    # KeyDown takes a different EventHandler<KeyEventArgs>; pythonnet handles
    # the conversion when the function signature matches. Pass the function
    # directly rather than wrapping in EventHandler.
    frm.KeyDown += fnOnKeyDown

    txtTarget.Select()
    logger.info("Showing dialog (frm.ShowDialog)")
    dialogResult = frm.ShowDialog()
    logger.info(f"Dialog returned: {dialogResult}")
    if dialogResult != DialogResult.OK:
        frm.Dispose()
        return False

    # Hand values back into the arguments namespace.
    arguments.sSource = (txtTarget.Text or "").strip()
    arguments.sOutputDir = (txtOut.Text or "").strip()
    arguments.bInvisible = bool(chkInvisible.Checked)
    arguments.bAuthenticate = bool(chkAuth.Checked)
    arguments.bMainProfile = bool(chkMain.Checked)
    arguments.bForce = bool(chkForce.Checked)
    arguments.bViewOutput = bool(chkView.Checked)
    arguments.bLog = bool(chkLog.Checked)
    arguments.bUseConfig = bool(chkUseCfg.Checked)
    frm.Dispose()
    return True


def launchReadMe():
    """
    Opens README.htm next to urlCheck.exe in the user's default browser.
    Falls back to README.md, then to a polite notice if neither is present.
    """
    sExeDir = ""
    sHtm = ""
    sMd = ""
    sTarget = ""

    sExeDir = os.path.dirname(sys.executable if getattr(sys, "frozen", False) else os.path.abspath(__file__))
    sHtm = os.path.join(sExeDir, "README.htm")
    sMd = os.path.join(sExeDir, "README.md")
    if os.path.isfile(sHtm): sTarget = sHtm
    elif os.path.isfile(sMd): sTarget = sMd
    if not sTarget:
        try:
            d = _loadDotNetForms()
            if d is not None:
                d["MessageBox"].Show(
                    "Documentation (README.htm or README.md) was not found in:\r\n\r\n" + sExeDir,
                    f"{sProgramName} - Documentation not found",
                    d["MessageBoxButtons"].OK, d["MessageBoxIcon"].Warning)
                return
        except Exception:
            pass
        print(f"[WARN] No README.htm or README.md in {sExeDir}")
        return
    try: os.startfile(sTarget)
    except Exception as ex: print(f"[WARN] Could not open {sTarget}: {ex}")


def showFinalGuiMessage(sText, sTitle):
    """
    Shows a captured-stdout message after a GUI-mode scan completes. Short
    text uses the native MessageBox; long text uses a scrollable read-only
    multi-line TextBox in a small Form. Mirrors 2htm's showFinalMessage.
    """
    bLong = False
    d = None

    d = _loadDotNetForms()
    if d is None:
        print(sText)
        return
    bLong = len(sText) > 800 or sText.count("\n") > 15
    if not bLong:
        d["MessageBox"].Show(sText or "Done. No output.", sTitle,
            d["MessageBoxButtons"].OK, d["MessageBoxIcon"].Information)
        return

    # Long output: scrollable read-only TextBox in a resizable form.
    AnchorStyles = d["AnchorStyles"]
    Button = d["Button"]
    DialogResult = d["DialogResult"]
    DockStyle = d["DockStyle"]
    Font = d["Font"]
    FontFamily = d["FontFamily"]
    Form = d["Form"]
    FormBorderStyle = d["FormBorderStyle"]
    FormStartPosition = d["FormStartPosition"]
    Panel = d["Panel"]
    Point = d["Point"]
    ScrollBars = d["ScrollBars"]
    Size = d["Size"]
    SystemFonts = d["SystemFonts"]
    TextBox = d["TextBox"]

    frm = Form()
    frm.Text = sTitle
    frm.StartPosition = FormStartPosition.CenterScreen
    frm.ClientSize = Size(700, 480)
    frm.FormBorderStyle = FormBorderStyle.Sizable
    frm.MinimizeBox = False
    frm.MaximizeBox = True
    frm.ShowInTaskbar = False
    frm.Font = SystemFonts.MessageBoxFont

    txt = TextBox()
    txt.Multiline = True
    txt.ReadOnly = True
    txt.ScrollBars = ScrollBars.Vertical
    txt.WordWrap = False
    txt.Text = sText
    txt.Dock = DockStyle.Fill
    try: txt.Font = Font(FontFamily.GenericMonospace, 9.0)
    except Exception: pass
    frm.Controls.Add(txt)

    pnl = Panel()
    pnl.Height = 40
    pnl.Dock = DockStyle.Bottom
    frm.Controls.Add(pnl)

    btn = Button()
    btn.Text = "OK"
    btn.DialogResult = DialogResult.OK
    btn.Size = Size(100, 26)
    btn.Anchor = AnchorStyles.Top | AnchorStyles.Right
    btn.Location = Point(pnl.ClientSize.Width - btn.Width - 12, 7)
    pnl.Controls.Add(btn)
    frm.AcceptButton = btn
    frm.CancelButton = btn

    frm.ShowDialog()
    frm.Dispose()


# --- Entry point ---

def main():
    bAutoLaunchedGui = False
    bGuiMode = False
    bMultiUrl = False
    bOwnConsole = False
    iErrorCount = 0
    iSkippedCount = 0
    iUrlIndex = 0
    iUrlTotal = 0
    lUrls = []
    sInput = ""
    sNormalizedUrl = ""
    # Capture wall-clock start time for the run, including a friendly
    # display string. Used in ACR.docx metadata header so the user
    # can correlate a generated report with a specific scan session.
    dtRunStart = datetime.datetime.now()

    arguments = None
    browser = None
    browserType = None
    context = None
    lArgs = []
    capture = None
    edgeProcess = None
    streamOriginalErr = sys.stderr
    streamOriginalOut = sys.stdout
    pathBaseDir = pathlib.Path.cwd()
    pathOutputDir = None
    sTempUserDataDir = None
    sWsEndpoint = None
    playwrightCtx = None

    # Playwright suppresses the default SIGINT handler. Restore it so Ctrl+C works.
    signal.signal(signal.SIGINT, signal.SIG_DFL)

    # On Windows, PyInstaller sets the DLL search path to the _MEI temporary folder
    # via SetDllDirectoryW. Child processes (Playwright driver, Edge) inherit this,
    # acquiring handles into _MEI that prevent the bootloader from deleting it on exit.
    # Resetting the DLL directory to NULL before launching Playwright prevents this.
    if sys.platform == "win32":
        ctypes.windll.kernel32.SetDllDirectoryW(None)

    cleanPreviousTempDirs()

    arguments = parseArguments()

    # nargs="*" returns a list. Collapse to a single space-joined string so
    # the rest of main() can treat the field uniformly with the GUI dialog's
    # text-field semantics (one url, several space-separated urls, or a path
    # to a url list file). Empty list -> empty string -> "no input given."
    if isinstance(arguments.sSource, list):
        arguments.sSource = " ".join(arguments.sSource).strip()

    # Logging strategy: the logger buffers all info/warn/error/debug
    # messages until logger.open() is called with the final pathLogDir.
    # We open the log AFTER config-load and (for GUI mode) AFTER the
    # dialog has set arguments.sOutputDir, so the log file lands in
    # the correct output folder. The buffer holds early diagnostics
    # (Python version, GUI auto-detection, etc.) which are then
    # flushed to the log file in the order they occurred.
    logger.info(f"{sProgramName} {sProgramVersion} starting")
    logger.info(f"Python: {sys.version.split(chr(10))[0]}")
    logger.info(f"Architecture: {struct.calcsize('P') * 8}-bit "
        f"({platform.machine()}); platform={platform.platform()}")
    logger.info(f"Frozen exe: {bool(getattr(sys, 'frozen', False))}; "
        f"executable={sys.executable}")
    if getattr(sys, 'frozen', False):
        # When PyInstaller-bundled, _MEIPASS exposes the temp extraction
        # directory and we can record the bundle's pid for cross-reference.
        logger.info(f"Bundle: _MEIPASS={getattr(sys, '_MEIPASS', '')}; "
            f"pid={os.getpid()}")
    logger.info(f"Working folder: {os.getcwd()}")
    logger.info(f"argv: {sys.argv}")

    # Auto-detect GUI launch via GetConsoleProcessList (primary) with parent-
    # process name as fallback (see isLaunchedFromGui). When invoked with no
    # arguments from a GUI shell -- File Explorer double-click, Start-menu
    # shortcut, desktop hotkey, Run dialog, etc. -- we fall through to GUI
    # mode and hide the otherwise-visible blank console. When invoked from
    # cmd.exe / PowerShell with no arguments we keep CLI behavior and print
    # usage. The -g flag forces GUI mode unconditionally regardless of how
    # the program was launched.
    bOwnConsole = isLaunchedFromGui()
    logger.info(f"bOwnConsole={bOwnConsole}, "
        f"explicit -g={bool(getattr(arguments, 'bGuiMode', False))}, "
        f"sSource={arguments.sSource!r}")
    if not getattr(arguments, "bGuiMode", False):
        if not arguments.sSource and bOwnConsole:
            arguments.bGuiMode = True
            bAutoLaunchedGui = True
            hideOwnConsoleWindow()
    bGuiMode = bool(getattr(arguments, "bGuiMode", False))
    logger.info(f"bGuiMode={bGuiMode}, bAutoLaunchedGui={bAutoLaunchedGui}")

    # Configuration file: opt-in for CLI (-u required), implicit for GUI mode
    # when an existing config file is present. This matches 2htm's asymmetry.
    if bGuiMode and not arguments.bUseConfig and configManager.configExists():
        arguments.bUseConfig = True
    if arguments.bUseConfig:
        configManager.loadInto(arguments)

    # In GUI mode, present the dialog before any other work. The user's
    # choices replace whatever came from CLI and config.
    if bGuiMode:
        if not showGuiDialog(arguments):
            logger.info("User cancelled the dialog")
            logger.close()
            return 0
        # If the user left Use configuration checked, persist their values.
        if arguments.bUseConfig:
            configManager.save(
                arguments.sSource or "",
                arguments.sOutputDir or "",
                arguments.bViewOutput,
                arguments.bInvisible,
                arguments.bForce,
                arguments.bLog,
                bAuthenticate=bool(getattr(arguments, "bAuthenticate", False)),
                bMainProfile=bool(getattr(arguments, "bMainProfile", False)))
        logger.info(f"After dialog: input={arguments.sSource!r} "
            f"outputDir={arguments.sOutputDir!r} "
            f"invisible={arguments.bInvisible} "
            f"authenticate={bool(getattr(arguments, 'bAuthenticate', False))} "
            f"mainProfile={bool(getattr(arguments, 'bMainProfile', False))} "
            f"force={arguments.bForce} "
            f"viewOutput={arguments.bViewOutput} "
            f"log={arguments.bLog} useConfig={arguments.bUseConfig}")

    if not arguments.sSource:
        print("No urls to scan.")
        return 1

    # Resolve the base output directory. Default = CWD.
    if arguments.sOutputDir:
        try:
            pathBaseDir = pathlib.Path(arguments.sOutputDir).expanduser()
            pathBaseDir.mkdir(parents=True, exist_ok=True)
        except Exception as ex:
            sErr = f"Output folder '{arguments.sOutputDir}' could not be created: {ex}"
            print(sErr)
            if bGuiMode: showFinalGuiMessage(sErr, f"{sProgramName} - Error")
            return 1
    else:
        pathBaseDir = pathlib.Path.cwd()

    # Open the log file now that pathBaseDir is final. This is the
    # ONLY logger.open call site in main(); buffered diagnostics from
    # earlier in the run are flushed to the file in their original
    # order. The log file lands in pathBaseDir alongside ACR.xlsx
    # and the per-page subfolders. If bLog is False, discard the
    # buffer to free memory and stop further buffering.
    if arguments.bLog:
        logger.open(bReplace=bool(getattr(arguments, "bForce", False)), pathDir=pathBaseDir)
    else:
        logger.discardBuffer()

    logger.info(f"Output base: {pathBaseDir}")
    logger.info(f"Target: {arguments.sSource}")

    # Write the run header to the log: program version, friendly start
    # time, and the resolved parameter list (showing both explicit and
    # defaulted values). The parameter labels mirror the GUI dialog
    # controls so the user can map a logged run to the dialog state.
    lHeaderParams = [
        ("Source urls",        str(arguments.sSource or "(none)")),
        ("Output folder",      str(pathBaseDir)),
        ("Force replacements", str(bool(arguments.bForce)).lower()),
        ("Invisible mode",     str(bool(arguments.bInvisible)).lower()),
        ("Authenticate credentials", str(bool(getattr(arguments, "bAuthenticate", False))).lower()),
        ("Main profile",  str(bool(getattr(arguments, "bMainProfile", False))).lower()),
        ("View output",        str(bool(arguments.bViewOutput)).lower()),
        ("Use configuration",  str(bool(arguments.bUseConfig)).lower()),
        ("Log session",        str(bool(arguments.bLog)).lower()),
        ("GUI mode",           str(bool(bGuiMode)).lower()),
        ("Working folder",     os.getcwd()),
        ("Command line",       " ".join(sys.argv)),
    ]
    logger.header(sProgramName, sProgramVersion, lHeaderParams)

    # In GUI mode, capture stdout/stderr so we can present the run summary in
    # a final dialog rather than scrolling past in a console the user can't
    # see. The CLI path leaves stdout/stderr alone.
    if bGuiMode:
        capture = io.StringIO()
        sys.stdout = capture
        sys.stderr = capture

    sInput = str(arguments.sSource).strip()

    # Three-case input dispatch:
    #
    #   (1) sInput is a path to an existing file that is NOT an allowed HTML
    #       extension -> treat as a url list file (one url per line).
    #
    #   (2) sInput contains internal whitespace and case (1) did not match
    #       -> treat as a list of urls separated by spaces. This is what
    #       lets the GUI Source urls field accept "https://a.com https://b.com".
    #
    #   (3) Otherwise -> single url / domain / local HTML file path.
    #
    # All three cases produce the same lUrls list, after which the per-url
    # scan loop below treats every entry uniformly. A multi-url run (lUrls
    # with more than one element) suppresses the auto-launch of report.htm
    # the same way a url list file does.
    # Three-case input dispatch using classifyInput():
    #
    #   ('listfile', sPath)  -> read urls from sPath (.txt only). The list
    #                           file's lines may be urls, domains, or local
    #                           HTML file paths.
    #   ('urls', lTokens)    -> sInput is one or more space-separated urls/
    #                           domains. Local HTML file paths are NOT valid
    #                           on direct input and have already been
    #                           rejected by classifyInput.
    #   ('error', sReason)   -> bail with a clear message.
    bMultiUrl = False
    sKind, vDetail = classifyInput(sInput)
    if sKind == "error":
        print(f"[ERROR] {vDetail}")
        logger.info(f"Input error: {vDetail}")
        if bGuiMode:
            sys.stdout = streamOriginalOut
            sys.stderr = streamOriginalErr
            showFinalGuiMessage(capture.getvalue(), f"{sProgramName} - Error")
        logger.close()
        return 1
    if sKind == "listfile":
        try:
            lUrls = getUrlsFromFile(vDetail)
        except Exception as ex:
            print(f"[ERROR] Could not read url list file: {ex}")
            if bGuiMode:
                sys.stdout = streamOriginalOut
                sys.stderr = streamOriginalErr
                showFinalGuiMessage(capture.getvalue(), f"{sProgramName} - Error")
            logger.close()
            return 1
        bMultiUrl = True
        iUrlTotal = len(lUrls)
        logger.info(f"url list: {vDetail} ({iUrlTotal} url(s))")
    else:
        # sKind == "urls"
        lUrls = [getNormalizedUrl(sToken) for sToken in vDetail if sToken]
        iUrlTotal = len(lUrls)
        bMultiUrl = iUrlTotal > 1
        if bMultiUrl: logger.info(f"url set: {iUrlTotal} url(s) from the source field")

    try:
        # Surface a fast user-visible message before starting the
        # Playwright stack. Spinning up sync_playwright(), then
        # locating msedge.exe, then opening the browser window can
        # take several seconds on a cold start; without this
        # message the user is left wondering whether urlCheck has
        # hung. Print to stdout in CLI mode; GUI mode already shows
        # its own progress dialog so an additional console line
        # would be invisible there.
        # The user-visible "Launching Edge..." message is printed by
        # urlCheck.cmd before invoking the .exe; we keep it in the log
        # for diagnostic purposes only.
        logger.info("Launching Edge...")
        with sync_playwright() as playwrightCtx:
            browserType = playwrightCtx.chromium
            lArgs = [
                "--mute-audio",
                "--no-default-browser-check",
                "--no-first-run",
                f"--window-size={iDefaultViewportWidth},{iDefaultViewportHeight}",
            ]
            # Playwright still calls this "headless"; we expose it to the
            # user as "Invisible" because that wording is clearer and not
            # tied to the implementation. The mapping is a 1:1 boolean.
            # When --authenticate is set the browser MUST be visible
            # so the user can interact with it (sign in, dismiss
            # banners, etc.); we override --invisible in that case
            # and log the override.
            #
            # urlCheck drives the system-installed Microsoft Edge through
            # channel="msedge"; we never download a Playwright-bundled
            # Chromium. On modern Windows 10/11, Edge ships in-box, so
            # this almost always succeeds. If the user has somehow removed
            # Edge or is on an unusual configuration where Playwright
            # cannot find it, surface a friendly message rather than a
            # raw Python traceback.
            bAuthEnabled = bool(getattr(arguments, "bAuthenticate", False))
            bMainProfile = bool(getattr(arguments, "bMainProfile", False))
            bHeadless = bool(arguments.bInvisible)
            if bAuthEnabled and bHeadless:
                logger.info("--authenticate overrides --invisible; "
                    "launching Edge with a visible window.")
                bHeadless = False

            browser = None
            context = None

            if bMainProfile:
                # When -m is set, launch Edge in PERSISTENT-CONTEXT mode
                # against the user's real Edge profile so saved logins,
                # cookies, and session state are available. This is what
                # makes "I'm already logged into Facebook in my normal
                # Edge" actually work in the urlCheck-driven session --
                # an ephemeral Playwright launch would start with an
                # empty profile and trigger anti-automation heuristics
                # (often presenting a blank page after the user signs in).
                #
                # Edge cannot run two instances against the same profile
                # at once, so we check first and surface a clear error
                # before trying. This applies whether or not -a is also
                # set: -m's defining property is that the real profile
                # is used, which on Windows requires no other msedge.exe
                # process to be active.
                if isEdgeRunning():
                    sErr = (
                        "Microsoft Edge is currently running. urlCheck "
                        "needs exclusive access to your Edge profile when "
                        "--main-profile is set so it can use your saved "
                        "logins. Please close all Edge windows (right-"
                        "click the Edge taskbar icon and choose Close "
                        "window, or quit Edge from its menu) and try "
                        "again.")
                    print(sErr)
                    logger.info("Edge already running; aborting -m launch.")
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Edge is running")
                    logger.close()
                    return 1
                sUserDataDir = getEdgeUserDataDir()
                if not sUserDataDir:
                    sErr = ("Could not locate the Microsoft Edge user-"
                        "data folder under "
                        "%LOCALAPPDATA%\\Microsoft\\Edge\\User Data. "
                        "urlCheck cannot use your Edge profile for "
                        "authenticated scans.")
                    print(sErr)
                    logger.info("Edge user-data dir not found; "
                        "aborting -a launch.")
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Edge profile not found")
                    logger.close()
                    return 1
                logger.info(f"Launching Edge with persistent context: "
                    f"user_data_dir={sUserDataDir}")
                # Suppress Playwright's --enable-automation default
                # switch. That switch triggers the visible "Microsoft
                # Edge is being controlled by automated test
                # software" infobar AND sets navigator.webdriver=true.
                # Removing it via ignore_default_args silences the
                # infobar.
                #
                # Pass chromium_sandbox=True so Playwright does NOT
                # add --no-sandbox to the command line. By default
                # Playwright disables the sandbox (chromium_sandbox=
                # False), which causes Edge to display an "unsupported
                # command-line flag: --no-sandbox" warning bar. With
                # the sandbox enabled (the same default Edge uses for
                # normal user browsing), no warning bar appears AND
                # the browser is more secure.
                #
                # Override navigator.webdriver via add_init_script
                # after the context launches. The override runs on
                # every new document before any site JavaScript can
                # read it, producing navigator.webdriver=undefined
                # without using a command-line flag (which would
                # itself trigger the "unsupported flag" warning).
                # This is needed because Playwright still passes
                # --remote-debugging-port internally (its IPC
                # channel), and per the MDN spec, navigator.web
                # driver is set whenever --enable-automation,
                # --headless, or --remote-debugging-port is in
                # effect.
                #
                # Note: as of Edge/Chromium 136 (April 2025) the
                # security model refuses --remote-debugging-port
                # against the default user-data directory, so a
                # literal "disconnect from CDP and reconnect" pattern
                # is not available when running against the user's
                # real Edge profile. The init-script approach
                # achieves the practical goal (sites stop refusing
                # the session because of automation fingerprinting)
                # without needing CDP disconnect/reconnect mechanics.
                try:
                    context = browserType.launch_persistent_context(
                        sUserDataDir,
                        channel=sBrowserChannel,
                        headless=bHeadless,
                        args=lArgs,
                        ignore_default_args=["--enable-automation"],
                        chromium_sandbox=True,
                        bypass_csp=True,
                        ignore_https_errors=bDefaultIgnoreHttpsErrors,
                        user_agent=sUserAgent,
                        viewport={"width": iDefaultViewportWidth,
                            "height": iDefaultViewportHeight})
                    # Silently override navigator.webdriver before any
                    # site JS runs. The 'configurable: true' allows
                    # later test pages (including the user's actual
                    # target site) to redefine it without TypeError.
                    context.add_init_script(
                        "Object.defineProperty(navigator, 'webdriver', "
                        "{get: () => undefined, configurable: true});")
                except Exception as ex:
                    sErr = (
                        f"Could not launch Microsoft Edge with your "
                        f"profile: {ex}\n\n"
                        "If Edge is still running, close all Edge "
                        "windows and try again. If the problem persists, "
                        "you can run urlCheck without --authenticate to "
                        "use a fresh Edge profile (no saved logins).")
                    print(sErr)
                    logger.info(f"Persistent-context launch failed: {ex}")
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Edge launch failed")
                    logger.close()
                    return 1
            elif bAuthEnabled:
                # When -a -t is set, launch msedge.exe ourselves as a
                # subprocess against a fresh temporary profile and use
                # Playwright's chromium.connect_over_cdp() to attach.
                # Then for each url's auth pause, call browser.close()
                # to sever Playwright's CDP connection (the Edge
                # subprocess keeps running because we, not Playwright,
                # own its lifecycle), prompt the user, and connect_
                # over_cdp() again to resume.
                #
                # Why subprocess + connect_over_cdp instead of launch_
                # server: playwright-python (unlike playwright-node)
                # does NOT expose BrowserType.launch_server. The
                # documented Python API for "launch a browser
                # ourselves and have Playwright attach" is connect_
                # over_cdp.
                #
                # Why not launch_persistent_context: it doesn't
                # support the disconnect/reconnect cycle. close() on
                # a launch_persistent_context-owned context kills
                # the browser process. We need the browser process
                # to OUTLIVE Playwright disconnect/reconnect cycles.
                #
                # Why a temp user-data-dir: Chrome 136+ refuses
                # --remote-debugging-port against the default user-
                # data-dir for security. A fresh temp directory is
                # non-default and so satisfies that restriction.
                #
                # Cookies and session state set by the user during
                # authentication persist within the Edge process via
                # its in-memory cookie jar (and on disk in the temp
                # user-data-dir), so subsequent urls on the same
                # domain reuse the established session without
                # re-prompting -- exactly like normal browsing.
                sEdgeExe = getEdgeExecutablePath()
                if not sEdgeExe:
                    sErr = ("Could not find msedge.exe at the standard "
                        "Microsoft Edge install locations. urlCheck -t "
                        "requires Edge to be installed.")
                    print(sErr)
                    logger.info(sErr)
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Edge not found")
                    logger.close()
                    return 1
                try:
                    sTempUserDataDir = tempfile.mkdtemp(prefix="urlCheck-tmp-profile-")
                except Exception as ex:
                    sErr = (f"Could not create temporary profile "
                        f"folder: {ex}")
                    print(sErr)
                    logger.info(sErr)
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Temp profile error")
                    logger.close()
                    return 1
                logger.info(f"Launching msedge.exe directly with "
                    f"--remote-debugging-port=0 and "
                    f"--user-data-dir={sTempUserDataDir}")
                # Build msedge.exe command line. Mirror the same args
                # we pass via Playwright in other paths, plus
                # --remote-debugging-port=0 (let OS pick a free port,
                # which Edge writes to DevToolsActivePort) and the
                # temp user-data-dir.
                lEdgeCmd = [sEdgeExe] + lArgs + [
                    "--remote-debugging-port=0",
                    f"--user-data-dir={sTempUserDataDir}",
                    "about:blank",
                ]
                if bHeadless: lEdgeCmd.append("--headless=new")
                try:
                    iCreateNoWindow = 0x08000000
                    edgeProcess = subprocess.Popen(
                        lEdgeCmd,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                        creationflags=iCreateNoWindow)
                except Exception as ex:
                    sErr = (f"Could not launch msedge.exe: {ex}")
                    print(sErr)
                    logger.info(sErr)
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Edge launch failed")
                    logger.close()
                    return 1
                # Wait for Edge to write its chosen port to disk.
                sPort = waitForDevToolsPort(sTempUserDataDir, iTimeoutSeconds=30)
                if not sPort:
                    sErr = ("msedge.exe did not write DevToolsActivePort "
                        "within 30 seconds. The browser may have failed "
                        "to start. urlCheck cannot continue with -t.")
                    print(sErr)
                    logger.info(sErr)
                    try: edgeProcess.kill()
                    except Exception: pass
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - Edge did not start")
                    logger.close()
                    return 1
                sCdpEndpoint = f"http://localhost:{sPort}"
                sWsEndpoint = sCdpEndpoint  # Reuse same field name elsewhere
                logger.info(f"Edge listening on CDP at {sCdpEndpoint}")
                try:
                    browser = browserType.connect_over_cdp(sCdpEndpoint)
                    # connect_over_cdp returns a Browser whose existing
                    # contexts include the default browser context that
                    # opened with about:blank. Use that one.
                    lContexts = list(browser.contexts)
                    if lContexts:
                        context = lContexts[0]
                    else:
                        context = browser.new_context(
                            bypass_csp=True,
                            ignore_https_errors=bDefaultIgnoreHttpsErrors,
                            user_agent=sUserAgent,
                            viewport={"width": iDefaultViewportWidth,
                                "height": iDefaultViewportHeight})
                    applyWebdriverOverride(context)
                except Exception as ex:
                    sErr = (f"Could not connect Playwright to Edge "
                        f"via CDP at {sCdpEndpoint}: {ex}")
                    print(sErr)
                    logger.info(sErr)
                    try: edgeProcess.kill()
                    except Exception: pass
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(),
                            f"{sProgramName} - CDP connect failed")
                    logger.close()
                    return 1
            else:
                # Non-auth runs: ephemeral launch with a fresh profile.
                # This is the original behavior for -a-not-set runs and
                # is unchanged.
                #
                # urlCheck drives the system-installed Microsoft Edge through
                # channel="msedge"; we never download a Playwright-bundled
                # Chromium. On modern Windows 10/11, Edge ships in-box, so
                # this almost always succeeds. If the user has somehow removed
                # Edge or is on an unusual configuration where Playwright
                # cannot find it, surface a friendly message rather than a
                # raw Python traceback.
                try:
                    browser = browserType.launch(channel=sBrowserChannel, headless=bHeadless, args=lArgs)
                except Exception as ex:
                    sErr = (
                        f"Could not launch Microsoft Edge: {ex}\n\n"
                        "urlCheck requires Microsoft Edge, which ships with "
                        "Windows 10 and 11 by default. If Edge has been "
                        "removed or is unavailable, install or repair it "
                        "from https://www.microsoft.com/edge and try again."
                    )
                    print(sErr)
                    logger.info(f"Browser launch failed: {ex}")
                    if bGuiMode:
                        sys.stdout = streamOriginalOut
                        sys.stderr = streamOriginalErr
                        showFinalGuiMessage(capture.getvalue(), f"{sProgramName} - Edge not available")
                    logger.close()
                    return 1
                context = browser.new_context(bypass_csp=True, ignore_https_errors=bDefaultIgnoreHttpsErrors, user_agent=sUserAgent, viewport={"width": iDefaultViewportWidth, "height": iDefaultViewportHeight})
            # Pre-fetch axe-core content once so CSP-restricted sites can still be scanned.
            sAxeContent = ""
            for sAxeUrl in aAxeCdnUrls:
                try:
                    sAxeContent = fetchText(sAxeUrl)
                    logger.info(f"axe-core pre-fetched from {sAxeUrl}")
                    break
                except Exception:
                    continue
            iUrlIndex = 0
            lScanned = []        # URLs that scanned successfully
            lFailed = []         # list of (sUrl, sReason)
            lSkippedExisting = []  # URLs whose output folder already exists
            lThisRunPaths = []   # per-page output folder paths (used by ACR builder)
            for sUrl in lUrls:
                iUrlIndex += 1
                # lUrls is already a list of normalized targets regardless of
                # which dispatch case produced it; we pass each entry through
                # scanUrl unchanged. (We re-run getNormalizedUrl idempotently
                # for the listfile case where lines were read raw from the
                # file without prior normalization.)
                sNormalizedUrl = getNormalizedUrl(sUrl) if sKind == "listfile" else sUrl
                if iUrlTotal > 1: logger.info(f"[{iUrlIndex}/{iUrlTotal}] {sNormalizedUrl}")
                # CLI mode: print the URL inline as work begins; on
                # success terminate with newline; on failure or skip
                # append ": <reason>" on the same line. GUI mode:
                # captured stdout becomes the final MessageBox -- we
                # want only the structured summary there, so no inline
                # writes. urlCheck has no per-page status form (unlike
                # 2htm and extCheck), so no progress UI to update.
                if not bGuiMode: print(sNormalizedUrl, end="", flush=True)
                try:
                    lConnHolder = [browser, context]
                    vResult = scanUrl(
                        sUrl, sNormalizedUrl, browser, context, pathBaseDir,
                        sAxeContent, bForce=bool(arguments.bForce),
                        bAuthenticate=bool(getattr(arguments, "bAuthenticate", False)),
                        bGuiMode=bGuiMode,
                        bMainProfile=bMainProfile,
                        playwrightCtx=playwrightCtx,
                        sWsEndpoint=sWsEndpoint,
                        lConnHolder=lConnHolder)
                    # If the disconnect/reconnect path replaced our
                    # browser/context inside scanUrl, pick up the
                    # fresh references for the next url. The
                    # disconnect path runs only when -a is set
                    # without -m (temp-profile auth runs).
                    if bAuthEnabled and not bMainProfile:
                        browser = lConnHolder[0]
                        context = lConnHolder[1]
                    if vResult == "skipped":
                        lSkippedExisting.append(sNormalizedUrl)
                        if not bGuiMode: print(": skipped (output folder exists, use -f to overwrite)")
                    else:
                        # scanUrl returns (pathStr, sPageTitle) on success.
                        sPathScanned = ""
                        sTitleScanned = ""
                        if isinstance(vResult, tuple) and len(vResult) == 2:
                            sPathScanned, sTitleScanned = vResult
                        elif isinstance(vResult, str):
                            sPathScanned = vResult
                        lScanned.append((sNormalizedUrl, sTitleScanned))
                        if sPathScanned:
                            try: lThisRunPaths.append(pathlib.Path(sPathScanned))
                            except Exception: pass
                        if not bGuiMode: print()
                except Exception as ex:
                    sReason = firstLine(str(ex))
                    lFailed.append((sNormalizedUrl, sReason))
                    if not bGuiMode: print(f": {sReason}")
                    logger.info(f"Error scanning {sNormalizedUrl}: {ex}")
                    logger.info(f"Traceback:\n{traceback.format_exc()}")

            iScanned = len(lScanned)
            iFailed = len(lFailed)
            iSkipped = len(lSkippedExisting)
            iErrorCount = iFailed     # legacy variable retained for
            iSkippedCount = iSkipped  # the existing exit-code logic

            # ---- Structured results summary ----
            #
            # Three sections, each printed only when non-zero. In CLI
            # mode the URLs were already printed inline during the
            # loop, so the per-URL list under each header is omitted
            # (it would just repeat). In GUI mode the captured stdout
            # becomes the final MessageBox, so include the lists
            # there. The format used in GUI mode for the scanned
            # list co-locates page title and URL on each line, so
            # the user can see what was scanned at a glance:
            #
            #   Checked 2 urls:
            #   Home Page - American Council of the Blind -- http://acb.org
            #   Home Page - National Federation of the Blind -- http://nfb.org

            if iScanned > 0:
                print()
                sScannedHeader = "Checked" if bGuiMode else "Scanned"
                print(f"{sScannedHeader} {iScanned} {'url' if iScanned == 1 else 'urls'}:")
                if bGuiMode:
                    for sUrlS, sTitleS in lScanned:
                        if sTitleS: print(f"{sTitleS} -- {sUrlS}")
                        else:       print(sUrlS)
            if iFailed > 0:
                print()
                print(f"Failed to scan {iFailed} {'url' if iFailed == 1 else 'urls'}:")
                if bGuiMode:
                    for sUrlF, sReason in lFailed:
                        if sReason: print(f"{sUrlF}: {sReason}")
                        else:       print(sUrlF)
            if iSkipped > 0:
                print()
                print(f"Skipped {iSkipped} {'url' if iSkipped == 1 else 'urls'}. "
                      f"Check \"Force replacements\" to overwrite.")
            if iScanned == 0 and iFailed == 0 and iSkipped == 0:
                print()
                print("No urls scanned.")

            if iFailed > 0 and not arguments.bLog and not bGuiMode:
                print("Re-run with -l to log full error tracebacks to urlCheck.log.")

            # Open the parent output directory once at the end of the run, if
            # requested. This shows the user all per-page subdirectories at
            # once rather than focusing on any single page's folder.
            if arguments.bViewOutput: openFolderInExplorer(str(pathBaseDir))
    finally:
        try:
            if context is not None: context.close()
        except Exception:
            pass
        try:
            if browser is not None: browser.close()
        except Exception:
            pass
        # When -t was used, we own the Edge subprocess directly.
        # Terminate it so it doesn't outlive urlCheck. kill() is a
        # hard SIGKILL-equivalent on Windows; that's fine because
        # we'll discard the temp profile next.
        try:
            if edgeProcess is not None:
                try: edgeProcess.kill()
                except Exception: pass
                try: edgeProcess.wait(timeout=5)
                except Exception: pass
        except Exception:
            pass
        # Clean up the temporary user-data directory. ignore_errors
        # because Edge sometimes holds a transient file lock for a
        # moment after exit; better to leak a temp dir than crash
        # the cleanup path.
        try:
            if sTempUserDataDir and os.path.isdir(sTempUserDataDir):
                shutil.rmtree(sTempUserDataDir, ignore_errors=True)
        except Exception:
            pass

        # Restore stdout/stderr and surface captured output in a GUI dialog.
        if bGuiMode and capture is not None:
            sys.stdout = streamOriginalOut
            sys.stderr = streamOriginalErr
            sCaptured = capture.getvalue()
            sTitle = f"{sProgramName} - Results" if iErrorCount == 0 else f"{sProgramName} - Completed with errors"
            showFinalGuiMessage(sCaptured if sCaptured else "Done. No output.", sTitle)

        # Write the session-level Accessibility Conformance Report.
        # See acrBuilder for the source-of-truth rules: --force uses
        # only this run's output paths; otherwise walk the parent for
        # all subfolders containing a results.json.
        try: acrBuilder.buildIfApplicable(arguments, lThisRunPaths, dtRunStart=dtRunStart)
        except NameError: pass  # lThisRunPaths may not be defined on early-exit paths

        logger.info(f"Done. {iUrlTotal - iErrorCount} of {iUrlTotal} scanned. {iErrorCount} error(s).")
        logger.close()

    return 0 if iErrorCount == 0 else 1


if __name__ == "__main__": raise SystemExit(main())
